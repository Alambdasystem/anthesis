from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from pydub import AudioSegment
from io import BytesIO  # Import BytesIO for handling byte streams
import requests
from requests.adapters import HTTPAdapter
from requests.packages.urllib3.util.retry import Retry
import json
import os
import logging
from threading import Thread
import re  # Import the re module for regular expressions
import base64
from midiutil import MIDIFile  # For creating MIDI files
from pydub.generators import Sine  # For generating audio tones
from ftplib import FTP  # <-- Add this import
import ftplib  # <-- Add this import for ftplib.error_perm usage
import uuid  # <-- Add this import for uuid.uuid4()
import jwt
from jwt.exceptions import ExpiredSignatureError, InvalidTokenError
from werkzeug.security import generate_password_hash, check_password_hash
import zipfile
import tempfile
import shutil
from werkzeug.utils import secure_filename

# Environment configuration
from dotenv import load_dotenv
load_dotenv()  # loads .env into os.environ

# Optional imports for document processing
try:
    import docx
except ImportError:
    docx = None
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

# Import agent routes
from agents import agents_bp

SECRET_KEY = os.getenv('SECRET_KEY', 'your_secret_here')

from datetime import datetime, timedelta, timezone
from functools import wraps

DEFAULT_MODEL = "llama3.2"  # Define DEFAULT_MODEL before usage

from flask_cors import CORS
from flask import Flask, request, jsonify

ENROLLMENT_FILE = "enrollment.json"

def load_enrollments():
    if os.path.exists(ENROLLMENT_FILE):
        with open(ENROLLMENT_FILE, "r") as f:
            try:
                data = json.load(f)
                return data.get("enrollments", {})
            except Exception:
                return {}
    return {}

def save_enrollments(enrollments):
    # Save with schema for future extensibility
    schema = {
        "firstName": "string",
        "lastName": "string",
        "email": "string",
        "phone": "string",
        "dob": "string",
        "gender": "string",
        "address": "string",
        "city": "string",
        "state": "string",
        "zip": "string",
        "education": "string",
        "employment": "string",
        "income": "string",
        "agreedToTerms": "boolean"
    }
    with open(ENROLLMENT_FILE, "w") as f:
        json.dump({"schema": schema, "enrollments": enrollments}, f, indent=2)

app = Flask(__name__)

# Configure app from environment variables
app.config.from_mapping(
    SECRET_KEY=SECRET_KEY,
    SMTP_SERVER=os.getenv('SMTP_SERVER', 'smtp.gmail.com'),
    SMTP_PORT=int(os.getenv('SMTP_PORT', '587')),
    SMTP_USER=os.getenv('SMTP_USER', ''),
    SMTP_PASS=os.getenv('SMTP_PASS', ''),
    OLLAMA_URL=os.getenv('OLLAMA_URL', 'http://localhost:11434/api/chat'),
    DEFAULT_MODEL=os.getenv('DEFAULT_MODEL', 'llama3.2'),
    MAX_CONTENT_LENGTH=int(os.getenv('MAX_CONTENT_LENGTH', '16777216')),  # 16MB
    UPLOAD_FOLDER=os.getenv('UPLOAD_FOLDER', 'uploads')
)

# Register agent blueprint
app.register_blueprint(agents_bp)

# Register chat blueprint
from agents.chat_routes import chat_bp
app.register_blueprint(chat_bp)

# Initialize all agents on startup
from agents.base import initialize_default_agents
try:
    initialize_default_agents()
    app.logger.info("✅ Agents initialized successfully")
except Exception as e:
    app.logger.error(f"❌ Agent initialization failed: {e}")

CORS(app, resources={
    r"/*": {
        "origins": ["https://alambda.systems", "http://localhost:3000"],
        "methods": ["GET", "POST", "PUT", "DELETE", "OPTIONS"],
        "allow_headers": ["Content-Type", "Authorization", "X-Requested-With"],
        "supports_credentials": True,
        "expose_headers": ["Content-Type", "Authorization", "Access-Control-Allow-Origin"],
        "max_age": 86400
    }
})

# Handle preflight requests
@app.before_request
def handle_preflight():
    if request.method == "OPTIONS":
        return '', 204

@app.route('/gps', methods=['GET'])
def get_shuttle_gps():
    zonar_url = "https://omi.zonarsystems.net/interface.php"
    params = {
        "action": "showposition",
        "operation": "current",
        "format": "json",
        "username": "api",
        "password": "Christian123!",
        "customer": "STA10021",
        "logvers": "3.0"  # Add the required logvers parameter
    }

    try:
        response = requests.get(zonar_url, params=params, timeout=10)
        data = response.json()
        if "data" in data and isinstance(data["data"], list):
            gps_data = [
                {
                    "name": p["asset"].get("fleet", f"Vehicle {i+1}"),
                    "lat": p["location"].get("lat", "N/A"),
                    "long": p["location"].get("lng", "N/A"),
                    "speed": p["location"].get("speed", "N/A"),
                    "heading": p["location"].get("heading", "N/A"),
                    "date": p["location"].get("date", "N/A"),
                    "power": p["location"].get("power", "N/A")
                }
                for i, p in enumerate(data["data"])
            ]
            return jsonify({"success": True, "vehicles": gps_data})
        else:
            return jsonify({"success": False, "error": "No GPS positions found"})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})
# Handle CORS errors
@app.after_request
def add_cors_headers(response):
    response.headers['Access-Control-Allow-Origin'] = '*'
    response.headers['Access-Control-Allow-Methods'] = 'GET, POST, PUT, DELETE, OPTIONS'
    response.headers['Access-Control-Allow-Headers'] = 'Content-Type, Authorization'
    response.headers['Access-Control-Allow-Credentials'] = 'true'
    return response
app.config['DEFAULT_MODEL'] = DEFAULT_MODEL
# ─── CORS ────────────────────────────────────────────────────────────────────
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')
logging.getLogger("werkzeug").setLevel(logging.ERROR)  # Suppress Flask's default logging
logging.getLogger("requests").setLevel(logging.ERROR)  # Suppress requests logging

# ─── Auth decorator ───────────────────────────────────────────────────────────
def token_required(f):
    from functools import wraps
    @wraps(f)
    def wrapper(*args, **kwargs):
        auth = request.headers.get("Authorization", "")
        if not auth.startswith("Bearer "):
            logging.warning(f"Token required: Missing or malformed token. Auth header: {auth}")
            return jsonify({"error": "Missing or malformed token"}), 403
        token = auth.split(None, 1)[1]
        try:
            payload = jwt.decode(token, SECRET_KEY, algorithms=["HS256"])
            request.user = payload["username"]
            logging.debug(f"Token decoded successfully for user: {request.user}")
        except jwt.ExpiredSignatureError:
            logging.warning("Token expired")
            return jsonify({"error": "Token expired"}), 401
        except jwt.InvalidTokenError as e:
            logging.warning(f"Invalid token: {e}")
            return jsonify({"error": "Invalid token"}), 403
        return f(*args, **kwargs)
    return wrapper

# ─── Enrollment API ─────────────────────────────────────────────────────────
@app.route('/api/enrollment', methods=['POST'])
@token_required
def submit_enrollment():
    try:
        username = request.user
        data = request.get_json() or {}
        required_fields = [
            'firstName', 'lastName', 'email', 'phone', 'dob', 'address', 'city', 'state', 'zip', 'education', 'employment', 'income', 'agreedToTerms'
        ]
        for field in required_fields:
            if field not in data or (isinstance(data[field], str) and not data[field].strip()):
                return jsonify({"error": f"Missing or empty required field: {field}"}), 400
        if not data.get("agreedToTerms", False):
            return jsonify({"error": "You must agree to the terms and conditions."}), 400
        enrollments = load_enrollments()
        enrollments[username] = data
        save_enrollments(enrollments)
        return jsonify({"success": True, "message": "Enrollment submitted."}), 200
    except Exception as e:
        logging.exception("Error in submit_enrollment:")
        return jsonify({"error": str(e)}), 500

@app.route('/api/enrollment/status', methods=['GET'])
@token_required
def enrollment_status():
    try:
        username = request.user
        enrollments = load_enrollments()
        enrolled = username in enrollments
        return jsonify({"enrolled": enrolled, "enrollment": enrollments.get(username) if enrolled else None}), 200
    except Exception as e:
        logging.exception("Error in enrollment_status:")
        return jsonify({"error": str(e)}), 500

@app.route('/cadet/overview', methods=['OPTIONS'])
def cadet_overview_preflight():
    response = jsonify({'status': 'ok'})
    response.headers.add('Access-Control-Allow-Origin', 'https://alambda.systems')
    response.headers.add('Access-Control-Allow-Methods', 'GET, OPTIONS')
    response.headers.add('Access-Control-Allow-Headers', 'Authorization, Content-Type')
    return response

from flask import Flask, request, jsonify
import requests

@app.route('/geocode', methods=['GET'])
def reverse_geocode():
    lat = request.args.get('lat')
    lon = request.args.get('lon')

    if not lat or not lon:
        return jsonify({'error': 'Missing lat/lon'}), 400

    try:
        url = f"https://nominatim.openstreetmap.org/reverse?format=jsonv2&lat={lat}&lon={lon}"
        headers = {'User-Agent': 'AlambdaShuttle/1.0 (contact@alambda.com)'}
        r = requests.get(url, headers=headers)
        data = r.json()
        return jsonify({'address': data.get('display_name')})
    except Exception as e:
        logging.error(f"Reverse geocoding failed: {e}")
        return jsonify({'error': 'Reverse geocoding failed'}), 500

# ─── CORS Preflight ──────────────────────────────────────────────────────────

# Enable debug logging
logging.basicConfig(level=logging.DEBUG)

OLLAMA_URL = "http://localhost:11434/api/chat"
DEFAULT_MODEL = "llama3.2"
SAMPLE_DIR = "samples"
GENERATED_PATH = "static/generated_track.wav"  # Changed to WAV

session = requests.Session()


retries = Retry(total=3, backoff_factor=2, status_forcelist=[500, 502, 503, 504])
session.mount('http://', HTTPAdapter(max_retries=retries))

# Pre-generated music pool
PRE_GENERATED_MUSIC = [
    {"genre": "lofi", "path": os.path.join(SAMPLE_DIR, "lofi_pre_generated.wav")},
    {"genre": "ambient", "path": os.path.join(SAMPLE_DIR, "ambient_pre_generated.wav")},
    {"genre": "chill", "path": os.path.join(SAMPLE_DIR, "chill_pre_generated.wav")},
]

# ─── 1. Add a JSON “database” for drafts ────────────────────────────────────
DRAFTS_FILE = "drafts.json"

def load_all_drafts():
    if os.path.exists(DRAFTS_FILE):
        with open(DRAFTS_FILE, "r") as f:
            return json.load(f)
    return []

def save_all_drafts(drafts):
    with open(DRAFTS_FILE, "w") as f:
        json.dump(drafts, f, indent=2, default=str)

def next_draft_id(drafts):
    if not drafts:
        return 1
    return max(d["id"] for d in drafts) + 1
@app.route('/')
def home():
    return '🧠 Ollama + Flask Chatbot + Music Generator is live!'

@app.route('/health')
def health():
    return jsonify({'status': 'ok'})


def extract_text_from_file(filepath):
    """Extract text from various file formats: txt, docx, pdf, doc, rtf, odt, html, md"""
    try:
        ext = os.path.splitext(filepath)[1].lower()
        
        if ext == '.txt':
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
                
        elif ext == '.docx':
            if not docx:
                return "Error: python-docx not installed"
            doc = docx.Document(filepath)
            return '\n'.join([p.text for p in doc.paragraphs])
            
        elif ext == '.pdf':
            try:
                # Try PyMuPDF (faster and better extraction) first if available
                import fitz
                doc = fitz.open(filepath)
                text = ""
                for page in doc:
                    text += page.get_text()
                return text
            except ImportError:
                # Fall back to PyPDF2
                if not PyPDF2:
                    return "Error: No PDF library installed"
                reader = PyPDF2.PdfReader(filepath)
                text = ''
                for page in reader.pages:
                    text += page.extract_text() or ''
                return text
                
        elif ext == '.doc':
            try:
                import textract
                return textract.process(filepath).decode('utf-8', errors='ignore')
            except ImportError:
                try:
                    import subprocess
                    return subprocess.check_output(['antiword', filepath]).decode('utf-8', errors='ignore')
                except (ImportError, FileNotFoundError):
                    return "Error: Neither textract nor antiword is available for DOC files"
                    
        elif ext == '.rtf':
            try:
                from striprtf.striprtf import rtf_to_text
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    rtf = f.read()
                return rtf_to_text(rtf)
            except ImportError:
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()  # Basic fallback, won't properly handle RTF formatting
                    
        elif ext == '.odt':
            try:
                import odf.opendocument
                import odf.text
                
                doc = odf.opendocument.load(filepath)
                paragraphs = doc.getElementsByType(odf.text.P)
                return '\n'.join([p.plainText() for p in paragraphs])
            except ImportError:
                try:
                    import textract
                    return textract.process(filepath).decode('utf-8', errors='ignore')
                except ImportError:
                    return "Error: No ODT library installed"
                    
        elif ext in ['.html', '.htm']:
            try:
                from bs4 import BeautifulSoup
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    soup = BeautifulSoup(f.read(), 'html.parser')
                return soup.get_text(separator='\n')
            except ImportError:
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()  # Raw HTML as fallback
                    
        elif ext in ['.md', '.markdown']:
            try:
                import markdown
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    md_text = f.read()
                html = markdown.markdown(md_text)
                try:
                    from bs4 import BeautifulSoup
                    return BeautifulSoup(html, 'html.parser').get_text()
                except ImportError:
                    return md_text  # Return raw markdown as fallback
            except ImportError:
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()  # Return raw markdown
                    
        else:
            return f"Unsupported file type: {ext}"
    except Exception as e:
        return f"Error extracting text: {str(e)}"
    

    try:
        # Check if this is a file upload request
        if 'file' in request.files:
            # This is a file upload - handle with process_zip_for_team logic
            file = request.files['file']
            prompt = request.form.get('prompt', '')
            
            if not file or file.filename == '':
                return jsonify({"error": "No file selected"}), 400
                
            # Verify it's a ZIP file
            if not file.filename.lower().endswith('.zip'):
                return jsonify({"error": "Please upload a ZIP file"}), 400

            # Create temp directory to extract files
            with tempfile.TemporaryDirectory() as temp_dir:
                # Save the zip file
                zip_path = os.path.join(temp_dir, secure_filename(file.filename))
                file.save(zip_path)
                
                # Extract all files
                extract_dir = os.path.join(temp_dir, "extracted")
                os.makedirs(extract_dir, exist_ok=True)
                
                with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                    zip_ref.extractall(extract_dir)
                
                # Process all files and extract text
                all_texts = []
                
                # Create extraction log
                extraction_log = {
                    "timestamp": datetime.now(timezone.utc).isoformat(),
                    "zip_filename": file.filename,
                    "files_processed": 0,
                    "files_skipped": 0,
                    "total_files": 0,
                    "file_details": []
                }
                
                for root, dirs, files in os.walk(extract_dir):
                    for filename in files:
                        extraction_log["total_files"] += 1
                        file_path = os.path.join(root, filename)
                        ext = os.path.splitext(filename)[1].lower()
                        
                        file_log = {
                            "filename": filename,
                            "path": os.path.relpath(file_path, extract_dir),
                            "extension": ext,
                            "size_bytes": os.path.getsize(file_path),
                            "status": "skipped",  # default status
                            "reason": ""
                        }
                        
                        # Skip non-document files
                        if ext not in ['.txt', '.pdf', '.docx']:
                            extraction_log["files_skipped"] += 1
                            file_log["reason"] = f"Unsupported file type: {ext}"
                            extraction_log["file_details"].append(file_log)
                            continue
                        
                        # Extract text and add filename as context
                        try:
                            text = extract_text_from_file(file_path)
                            
                            # Check if extraction was successful
                            if text and not text.startswith("Error:"):
                                all_texts.append(f"--- {filename} ---\n{text}\n\n")
                                extraction_log["files_processed"] += 1
                                file_log["status"] = "success"
                                file_log["extracted_chars"] = len(text)
                            else:
                                extraction_log["files_skipped"] += 1
                                file_log["status"] = "failed"
                                file_log["reason"] = text if text.startswith("Error:") else "Empty text extracted"
                        except Exception as e:
                            extraction_log["files_skipped"] += 1
                            file_log["status"] = "error"
                            file_log["reason"] = str(e)
                        
                        extraction_log["file_details"].append(file_log)
                
                # Save extraction log to file (optional)
                log_dir = "extraction_logs"
                os.makedirs(log_dir, exist_ok=True)
                log_filename = f"extraction_log_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}.json"
                log_path = os.path.join(log_dir, log_filename)
                with open(log_path, 'w') as f:
                    json.dump(extraction_log, f, indent=2)
                
                # Combine all texts
                combined_text = "\n".join(all_texts)
                
                # Prepare the payload for Ollama with the combined text and prompt
                full_prompt = f"{prompt}\n\nRESUME CONTENTS:\n{combined_text}"
                
                # Limit text size if needed
                if len(full_prompt) > 30000:
                    full_prompt = full_prompt[:30000] + "...\n[Content truncated due to size limits]"
                
                # Call Ollama API
                payload = {
                    "model": "llama3.2",
                    "messages": [
                        {"role": "system", "content": "You are a team-building AI assistant that analyzes resumes."},
                        {"role": "user", "content": full_prompt}
                    ],
                    "temperature": 0.7,
                    "max_tokens": 2000,
                    "stream": False
                }
                
                response = session.post(OLLAMA_URL, json=payload, timeout=120)
                
                if response.status_code != 200:
                    return jsonify({"error": f"Ollama error {response.status_code}"}), 500
                    
                # Extract and return the content
                content = response.json().get("message", {}).get("content", "").strip()
                
                # Try to parse as JSON if possible
                try:
                    parsed_json = json.loads(content)
                    return jsonify({
                        "response": parsed_json,
                        "extraction_log": extraction_log
                    })
                except json.JSONDecodeError:
                    # If not valid JSON, return as text
                    return jsonify({
                        "response": content,
                        "extraction_log": extraction_log
                    })
    except Exception as e:
        logging.exception("Error in predict endpoint:")
        return jsonify({"error": str(e)}), 500
# ─── Contact Management ──────────────────────────────────────────────────────
@app.route('/contacts', methods=['GET'])
@token_required
def get_contacts():
    try:
        if os.path.exists('contacts.json'):
            with open('contacts.json', 'r') as f:
                contacts = json.load(f)
        else:
            contacts = []
        return jsonify(contacts)
    except Exception as e:
        logging.exception("Error in get_contacts:")
        return jsonify({"error": str(e)}), 500

@app.route('/contacts', methods=['POST'])
@token_required
def create_contact():
    try:
        data = request.get_json()
        required_fields = ['name', 'email', 'agency', 'focus', 'why_hot']
        if not all(field in data for field in required_fields):
            return jsonify({"error": "Missing required fields"}), 400

        if os.path.exists('contacts.json'):
            with open('contacts.json', 'r') as f:
                contacts = json.load(f)
        else:
            contacts = []

        new_contact = {
            'id': str(uuid.uuid4()),
            'name': data['name'],
            'email': data['email'],
            'agency': data['agency'],
            'focus': data['focus'],
            'why_hot': data['why_hot'],
            'created_at': datetime.now(timezone.utc).isoformat(),
            'updated_at': datetime.now(timezone.utc).isoformat()
        }

        contacts.append(new_contact)
        with open('contacts.json', 'w') as f:
            json.dump(contacts, f, indent=2)

        return jsonify(new_contact), 201
    except Exception as e:
        logging.exception("Error in create_contact:")
        return jsonify({"error": str(e)}), 500

@app.route('/contacts/<contact_id>', methods=['PUT'])
@token_required
def update_contact(contact_id):
    try:
        data = request.get_json()
        if os.path.exists('contacts.json'):
            with open('contacts.json', 'r') as f:
                contacts = json.load(f)
        else:
            return jsonify({"error": "No contacts found"}), 404

        contact = next((c for c in contacts if c['id'] == contact_id), None)
        if not contact:
            return jsonify({"error": "Contact not found"}), 404

        for key, value in data.items():
            if key in contact:
                contact[key] = value
        contact['updated_at'] = datetime.now(timezone.utc).isoformat()

        with open('contacts.json', 'w') as f:
            json.dump(contacts, f, indent=2)

        return jsonify(contact)
    except Exception as e:
        logging.exception("Error in update_contact:")
        return jsonify({"error": str(e)}), 500

@app.route('/contacts/<contact_id>', methods=['DELETE'])
@token_required
def delete_contact(contact_id):
    try:
        if os.path.exists('contacts.json'):
            with open('contacts.json', 'r') as f:
                contacts = json.load(f)
        else:
            return jsonify({"error": "No contacts found"}), 404

        contacts = [c for c in contacts if c['id'] != contact_id]
        with open('contacts.json', 'w') as f:
            json.dump(contacts, f, indent=2)

        return jsonify({"message": "Contact deleted successfully"})
    except Exception as e:
        logging.exception("Error in delete_contact:")
        return jsonify({"error": str(e)}), 500

# ─── Analytics ────────────────────────────────────────────────────────────────
@app.route('/analytics/<contact_id>', methods=['GET'])
@token_required
def get_contact_analytics(contact_id):
    try:
        # Load contacts from file
        if os.path.exists('contacts.json'):
            with open('contacts.json', 'r') as f:
                contacts = json.load(f)
        else:
            contacts = []

        # Find the contact by ID
        contact = next((c for c in contacts if c['id'] == contact_id), None)
        if not contact:
            return jsonify({"error": "Contact not found"}), 404

        # Extract fields for the prompt
        name = contact.get('name', '')
        agency = contact.get('agency', '')
        focus = contact.get('focus', '')
        why_hot = contact.get('why_hot', '')

        prompt = f"""
        Analyze the following contact information and provide detailed insights:
        Name: {name}
        Agency: {agency}
        Focus: {focus}
        Why Hot: {why_hot}
        
        Provide insights on:
        1. Potential opportunities
        2. Key decision-making factors
        3. Communication preferences
        4. Potential challenges
        """

        response = session.post(OLLAMA_URL, json={
            "model": "llama3.2",
            "messages": [
                {"role": "system", "content": "You are a CRM analytics expert."},
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.7,
            "max_tokens": 500
        })

        if response.status_code != 200:
            return jsonify({"error": f"Ollama error {response.status_code}"}), 500

        analytics = response.json().get("message", {}).get("content", "").strip()
        return jsonify({"analytics": analytics})
    except Exception as e:
        logging.exception("Error in get_contact_analytics:")
        return jsonify({"error": str(e)}), 500

# ─── Email Templates ──────────────────────────────────────────────────────────
@app.route('/email-templates', methods=['GET'])
@token_required
def get_email_templates():
    try:
        if os.path.exists('email_templates.json'):
            with open('email_templates.json', 'r') as f:
                templates = json.load(f)
        else:
            templates = []
        return jsonify(templates)
    except Exception as e:
        logging.exception("Error in get_email_templates:")
        return jsonify({"error": str(e)}), 500

@app.route('/email-templates', methods=['POST'])
@token_required
def create_email_template():
    try:
        data = request.get_json()
        required_fields = ['name', 'subject', 'body']
        if not all(field in data for field in required_fields):
            return jsonify({"error": "Missing required fields"}), 400

        if os.path.exists('email_templates.json'):
            with open('email_templates.json', 'r') as f:
                templates = json.load(f)
        else:
            templates = []

        new_template = {
            'id': str(uuid.uuid4()),
            'name': data['name'],
            'subject': data['subject'],
            'body': data['body'],
            'created_at': datetime.now(timezone.utc).isoformat()
        }

        templates.append(new_template)
        with open('email_templates.json', 'w') as f:
            json.dump(templates, f, indent=2)

        return jsonify(new_template), 201
    except Exception as e:
        logging.exception("Error in create_email_template:")
        return jsonify({"error": str(e)}), 500

# ─── Profile ──────────────────────────────────────────────────────────────────
@app.route('/profile')
@token_required
def profile():
    try:
        # Get user data from token
        user = request.user
        
        # For now, return a sample profile
        # In production, this would fetch actual user data from your database
        return jsonify({
            "name": user,
            "email": f"{user}@example.com",
            "bio": "AI CRM Specialist"
        })
    except Exception as e:
        logging.exception("Error in profile endpoint:")
        return jsonify({"error": str(e)}), 500

@app.route('/register', methods=['POST'])
def register():
    logging.info("Register endpoint called.")
    try:
        data = request.get_json()
        logging.debug(f"Received registration data: {data}")

        username = data.get("username")
        email = data.get("email")
        password = data.get("password")
        track = data.get("track", "")  # Optional, default empty

        if not username or not email or not password:
            logging.warning("Missing required fields in registration.")
            return jsonify({"error": "Username, email, and password are required."}), 400

        if os.path.exists("users.json"):
            with open("users.json", "r") as f:
                users = json.load(f)
        else:
            users = {}

        if username in users:
            logging.warning(f"User '{username}' already exists.")
            return jsonify({"error": "User already exists."}), 400

        # Hash the password before saving
        hashed_password = generate_password_hash(password)
        logging.debug(f"Generated hashed password for user '{username}'.")

        users[username] = {
            "email": email,
            "password": hashed_password,  # Save hashed password
            "track": track
        }

        with open("users.json", "w") as f:
            json.dump(users, f, indent=2)
        logging.info(f"User '{username}' registered successfully.")

        token = jwt.encode(
            {"username": username, "exp": datetime.now(timezone.utc) + timedelta(hours=24)},
            SECRET_KEY,
            algorithm="HS256"
        )
        if isinstance(token, bytes):
            token = token.decode('utf-8')

        return jsonify({"success": True, "message": "Registration successful", "token": token}), 200
    except Exception as e:
        logging.exception("Error in register endpoint:")
        return jsonify({"error": str(e)}), 500


@app.route('/login', methods=['POST'])
def login():
    logging.info("Login endpoint called.")
    try:
        data = request.get_json()
        logging.debug(f"Received login data: {data}")

        username = data.get("username")
        password = data.get("password")

        if not username or not password:
            return jsonify({"error": "Username and password are required."}), 400

        if os.path.exists("users.json"):
            with open("users.json", "r") as f:
                users = json.load(f)
        else:
            return jsonify({"error": "No users registered."}), 400

        user = users.get(username)
        if not user or not check_password_hash(user["password"], password):
            return jsonify({"error": "Invalid username or password."}), 401

        token = jwt.encode(
            {"username": username, "exp": datetime.now(timezone.utc) + timedelta(hours=24)},
            SECRET_KEY,
            algorithm="HS256"
        )
        if isinstance(token, bytes):
            token = token.decode('utf-8')

        return jsonify({"success": True, "message": "Login successful", "token": token}), 200
    except Exception as e:
        logging.exception("Error in login endpoint:")
        return jsonify({"error": str(e)}), 500


@app.route('/logout', methods=['POST'])
@token_required
def logout():
    try:
        # Clear the token from localStorage on the frontend
        return jsonify({"success": True, "message": "Logged out successfully"}), 200
    except Exception as e:
        logging.exception("Error in logout endpoint:")
        return jsonify({"error": str(e)}), 500
    logging.info("Login endpoint called.")
    try:
        data = request.get_json()
        logging.debug(f"Received login data: {data}")

        username = data.get("username")
        password = data.get("password")

        if not username or not password:
            logging.warning("Missing username or password in login.")
            return jsonify({"error": "Username and password are required."}), 400

        if os.path.exists("users.json"):
            with open("users.json", "r") as f:
                users = json.load(f)
        else:
            users = {}

        logging.debug(f"Loaded users: {users}")

        user = users.get(username)
        if not user:
            logging.warning(f"User '{username}' not found.")
            return jsonify({"error": "Invalid credentials"}), 401

        # Check if the password matches
        if not check_password_hash(user.get("password"), password):
            logging.warning(f"Password mismatch for user '{username}'.")
            return jsonify({"error": "Invalid credentials"}), 401

        logging.info(f"User '{username}' logged in successfully.")

        token = jwt.encode(
            {"username": username, "exp": datetime.now(timezone.utc) + timedelta(hours=24)},
            SECRET_KEY,
            algorithm="HS256"
        )
        if isinstance(token, bytes):
            token = token.decode('utf-8')

        response = jsonify({"token": token})
        return response
    except Exception as e:
        logging.exception("Error in login endpoint:")
        return jsonify({"error": str(e)}), 500

def generate_music_in_background(prompt, model, system, temperature, max_tokens):
    """Generate music in the background, now saving as WAV instead of MP3."""
    try:
        ollama_payload = {
            "model": model,
            "messages": [
                {"role": "system", "content": system},
                {"role": "user", "content": f"{prompt}. Respond with JSON: {{\"genre\": \"lofi\", \"drums\": \"lofi1\", \"melody\": \"melody1\"}}"}
            ],
            "temperature": temperature,
            "max_tokens": max_tokens,
            "stream": False
        }

        logging.debug("Sending music prompt to Ollama: %s", json.dumps(ollama_payload, indent=2))
        response = session.post(OLLAMA_URL, json=ollama_payload, timeout=60)

        if response.status_code != 200:
            logging.error(f"Ollama returned status {response.status_code}")
            return

        content = response.json().get("message", {}).get("content", "{}")
        logging.debug("LLM music output: %s", content)

        # Extract JSON part from the response content
        start_index = content.find("{")
        end_index = content.rfind("}")
        if start_index == -1 or end_index == -1:
            logging.error("No valid JSON found in the response content.")
            return

        json_content = content[start_index:end_index + 1]
        music_info = json.loads(json_content)
        notes = music_info.get("notes", [])
        if not notes:
            logging.warning("No notes retrieved for WAV generation.")
            return

        # Make the final track from notes
        create_wave_from_notes(notes, GENERATED_PATH)
        logging.info("Generated track saved to %s", GENERATED_PATH)

    except Exception as e:
        logging.exception("Music generation error:")

@app.route('/ftp-upload', methods=['POST'])
@token_required
def ftp_upload():
    try:
        uploaded_file = request.files['file']
        if not uploaded_file:
            return jsonify({"error": "No file uploaded."}), 400

        ftp_host = "107.180.25.236"
        ftp_user = "mrworker@alambda.systems"
        ftp_pass = "Christian123!"

        ftp = FTP()
        ftp.connect(ftp_host, 21, timeout=10)
        ftp.login(ftp_user, ftp_pass)

        # 🧠 New: Get username from token
        username = request.user  
        user_folder = f"{username}/"

        # 🛠️ Create user folder if it doesn't exist
        try:
            ftp.cwd(user_folder)
        except Exception:
            ftp.mkd(user_folder)
            ftp.cwd(user_folder)

        # ✅ Upload file into the user's folder
        ftp.storbinary(f"STOR {uploaded_file.filename}", uploaded_file.stream)
        ftp.quit()

        return jsonify({"success": True, "message": f"File '{uploaded_file.filename}' uploaded successfully into {user_folder}."})

    except Exception as e:
        logging.exception("FTP Upload Error:")
        return jsonify({"error": str(e)}), 500


@app.route('/ftp-browse', methods=['GET', 'OPTIONS'])
@token_required
def ftp_browse():
    if request.method == 'OPTIONS':
        response = jsonify({'status': 'ok'})
        response.headers.add('Access-Control-Allow-Origin', 'https://alambda.systems')
        response.headers.add('Access-Control-Allow-Methods', 'GET, OPTIONS')
        response.headers.add('Access-Control-Allow-Headers', 'Authorization, Content-Type')
        return response

    try:
        path = request.args.get('path', '/')
        
        ftp_host = "107.180.25.236"
        ftp_user = "mrworker@alambda.systems"
        ftp_pass = "Christian123!"

        ftp = FTP()
        ftp.connect(ftp_host, 21, timeout=10)
        ftp.login(ftp_user, ftp_pass)

        # Create necessary directories if they don't exist
        parts = path.strip('/').split('/')
        current_path = ''
        for part in parts:
            if part:
                current_path = f"{current_path}/{part}" if current_path else part
                try:
                    # Try to change to the directory
                    ftp.cwd(current_path)
                except ftplib.error_perm as e:
                    if "550" in str(e):  # Directory does not exist
                        try:
                            ftp.mkd(current_path)  # Create the directory
                            ftp.cwd(current_path)  # Change to the newly created directory
                        except ftplib.error_perm as mkd_error:
                            print(f"Failed to create directory: {current_path}")
                            raise mkd_error
                    else:
                        raise e

        # Get directory listing
        contents = []
        ftp.dir(contents.append)
        
        # Parse directory listing
        files = []
        for line in contents:
            parts = line.split()
            if len(parts) >= 9:
                name = ' '.join(parts[8:])
                size = parts[4]
                is_dir = parts[0].startswith('d')
                files.append({
                    'name': name,
                    'size': size if not is_dir else None,
                    'is_dir': is_dir
                })
        
        ftp.quit()
        return jsonify({"success": True, "path": path, "contents": files})

    except Exception as e:
        logging.exception("FTP Browse Error:")
        return jsonify({"error": str(e)}), 500




@app.route('/ftp', methods=['POST', 'OPTIONS'])
@token_required
def ftp_operations():
    if request.method == 'OPTIONS':
        response = jsonify({'status': 'ok'})
        response.headers.add('Access-Control-Allow-Origin', 'https://alambda.systems')
        response.headers.add('Access-Control-Allow-Methods', 'POST, OPTIONS')
        response.headers.add('Access-Control-Allow-Headers', 'Authorization, Content-Type')
        return response

    try:
        operation = request.form.get('operation') or request.json.get('operation')
        if not operation:
            return jsonify({"error": "Operation is required."}), 400

        ftp_host = "107.180.25.236"
        ftp_user = "mrworker@alambda.systems"
        ftp_pass = "Christian123!"

        ftp = FTP()
        ftp.connect(ftp_host, 21, timeout=10)
        ftp.login(ftp_user, ftp_pass)

        if operation == 'read':
            filename = request.form.get('filename') or request.json.get('filename')
            if not filename:
                return jsonify({"error": "Filename is required."}), 400

            try:
                # Read file content
                content = BytesIO()
                ftp.retrbinary(f"RETR {filename}", content.write)
                content.seek(0)
                
                # Determine content type based on file extension
                import mimetypes
                content_type, _ = mimetypes.guess_type(filename)
                
                return send_file(
                    content,
                    mimetype=content_type or 'application/octet-stream',
                    as_attachment=False
                )
            except Exception as e:
                return jsonify({"error": f"Failed to read file: {str(e)}"}), 500

        elif operation == 'download':
            filename = request.form.get('filename') or request.json.get('filename')
            if not filename:
                return jsonify({"error": "Filename is required."}), 400

            try:
                # Read file content
                content = BytesIO()
                ftp.retrbinary(f"RETR {filename}", content.write)
                content.seek(0)
                
                # Get filename from path
                import os
                filename = os.path.basename(filename)
                
                return send_file(
                    content,
                    as_attachment=True,
                    download_name=filename
                )
            except Exception as e:
                return jsonify({"error": f"Failed to download file: {str(e)}"}), 500

        elif operation == 'upload':
            file = request.files.get('file')
            if not file:
                return jsonify({"error": "No file provided."}), 400

            try:
                # Create necessary directories if they don't exist
                parts = filename.strip('/').split('/')
                current_path = ''
                for part in parts[:-1]:  # Exclude the filename itself
                    if part:
                        current_path = f"{current_path}/{part}" if current_path else part
                        try:
                            ftp.cwd(current_path)
                        except ftplib.error_perm:
                            ftp.mkd(current_path)
                            ftp.cwd(current_path)

                # Upload the file
                ftp.storbinary(f"STOR {filename}", file)
                return jsonify({"success": True, "message": f"File '{file.filename}' uploaded successfully"})
            except Exception as e:
                return jsonify({"error": f"Failed to upload file: {str(e)}"}), 500

        else:
            return jsonify({"error": "Invalid operation."}), 400

    except Exception as e:
        logging.exception("FTP Operation Error:")
        return jsonify({"error": str(e)}), 500


@app.route('/ftp-metadata', methods=['GET'])
def ftp_metadata():
    try:
        filename = request.args.get('file')
        if not filename:
            return jsonify({"error": "Filename is required."}), 400

        ftp_host = "107.180.25.236"
        ftp_user = "mrworker@alambda.systems"
        ftp_pass = "Christian123!"

        ftp = FTP()
        ftp.connect(ftp_host, 21, timeout=10)
        ftp.login(ftp_user, ftp_pass)

        size = ftp.size(filename)
        modified_time = ftp.sendcmd(f"MDTM {filename}")[4:].strip()
        ftp.quit()

        return jsonify({
            "success": True,
            "filename": filename,
            "size": size,
            "modified_time": modified_time
        })

    except Exception as e:
        logging.exception("FTP Metadata Error:")
        return jsonify({"error": str(e)}), 500

@app.route('/ftp-search', methods=['GET'])
@token_required
def ftp_search():
    try:
        query = request.args.get('query', '').lower()
        if not query:
            return jsonify({"error": "Search query is required."}), 400

        ftp_host = "107.180.25.236"
        ftp_user = "mrworker@alambda.systems"
        ftp_pass = "Christian123!"

        ftp = FTP()
        ftp.connect(ftp_host, 21, timeout=10)
        ftp.login(ftp_user, ftp_pass)

        username = request.user
        user_folder = f"{username}/"

        try:
            ftp.cwd(user_folder)
        except Exception:
            return jsonify({"error": "User folder not found."}), 404

        items = ftp.nlst()
        matching_items = [item for item in items if query in item.lower()]
        ftp.quit()

        return jsonify({"success": True, "query": query, "results": matching_items})

    except Exception as e:
        logging.exception("FTP Search Error:")
        return jsonify({"error": str(e)}), 500

def generate_pre_generated_music():
    """Generate pre-generated music tracks at app startup."""
    SAMPLE_DIR = "static/samples"
    os.makedirs(SAMPLE_DIR, exist_ok=True)
    
    # Generate simple musical notes
    notes = ["C4", "D4", "E4", "F4", "G4", "A4", "B4", "C5"]
    
    # Create a WAV file for each note
    for note in notes:
        output_path = os.path.join(SAMPLE_DIR, f"{note}.wav")
        try:
            create_wave_from_notes([{"note": note, "duration": 1.0}], output_path)
            logging.info(f"Created WAV file for note {note}")
        except Exception as e:
            logging.error(f"Failed to create WAV file for note {note}: {e}")
            continue

    return jsonify({"status": "success", "message": "Pre-generated music created"})

def create_mp3_from_notes(notes, output_path):
    """Convert music notes to an MP3 file."""
    try:
        logging.debug(f"Creating MP3 for notes: {notes}")
        audio = AudioSegment.silent(duration=0)

        for note_info in notes:
            note = note_info.get("note", "C4")
            duration = note_info.get("duration", 1.0)
            pitch = note_to_frequency(note)
            if pitch is not None:
                tone = Sine(pitch).to_audio_segment(duration=duration * 1000)
                audio += tone

        logging.debug(f"Exporting MP3 to {output_path}")
        audio.export(output_path, format="mp3")
        logging.info(f"MP3 file created at {output_path}")
    except Exception as e:
        logging.error(f"Failed to create MP3 file: {e}")

def note_to_midi(note):
    """Convert a musical note (e.g., C4) to a MIDI pitch number."""
    note_map = {"C": 0, "C#": 1, "D": 2, "D#": 3, "E": 4, "F": 5, "F#": 6, "G": 7, "G#": 8, "A": 9, "A#": 10, "B": 11}
    try:
        note_name = note[:-1]
        octave = int(note[-1])
        return 12 * (octave + 1) + note_map[note_name]
    except (KeyError, ValueError):
        logging.error(f"Invalid note: {note}")
        return None

def note_to_frequency(note):
    """Convert a musical note (e.g., C4) to a frequency in Hz."""
    midi_pitch = note_to_midi(note)
    if midi_pitch is not None:
        return 440.0 * (2 ** ((midi_pitch - 69) / 12.0))  # A4 = 440 Hz
    return None

def create_wave_from_notes(notes, output_path):
    """Convert music notes to a WAV file (no ffmpeg required)."""
    try:
        logging.debug(f"Creating WAV for notes: {notes}")
        audio = AudioSegment.silent(duration=0)

        for note_info in notes:
            note = note_info.get("note", "C4")
            duration = note_info.get("duration", 1.0)
            pitch = note_to_frequency(note)
            if pitch is not None:
                tone = Sine(pitch).to_audio_segment(duration=duration * 1000)
                audio += tone

            # Handle chords if present
            chord = note_info.get("chord", [])
            for chord_note in chord:
                chord_pitch = note_to_frequency(chord_note.get("note"))
                if chord_pitch:
                    chord_tone = Sine(chord_pitch).to_audio_segment(duration=chord_note.get("duration", 1.0) * 1000)
                    audio = audio.overlay(chord_tone)

        logging.debug(f"Exporting WAV to {output_path}")
        audio.export(output_path, format="wav")
        logging.info(f"WAV file created at {output_path}")
    except Exception as e:
        logging.error(f"Failed to create WAV file: {e}")

@app.route('/curriculum/generate', methods=['OPTIONS'])
def curriculum_generate_preflight():
    response = jsonify({'status': 'ok'})
    response.headers.add('Access-Control-Allow-Origin', 'https://alambda.systems')
    response.headers.add('Access-Control-Allow-Methods', 'POST, OPTIONS')
    response.headers.add('Access-Control-Allow-Headers', 'Content-Type, Authorization')
    return response, 200

@app.route('/curriculum/generate', methods=['POST'])
@token_required
def generate_curriculum():
    """
    Generate a curriculum module based on the provided track and level.
    """
    try:
        data = request.get_json() or {}
        track = data.get("track", "").strip()
        level = int(data.get("level", 1))
        temperature = float(data.get("temperature", 0.7))
        max_tokens = int(data.get("max_tokens", 400))

        if not track:
            return jsonify({"error": "track is required"}), 400

        # Prompt for the AI model
        prompt = (
            f"Create a beginner learning module for the track '{track}', level {level}.\n"
            "Respond in structured text with:\n"
            "Module Title: <title>\n"
            "Description: <brief overview>\n"
            "Module Link: <URL>\n"
            "Question 1: <text>\n"
            "- A: Option A\n"
            "- B: Option B\n"
            "- C: Option C\n"
            "Answer: <correct option letter>\n"
        )

        payload = {
            "model": DEFAULT_MODEL,
            "messages": [
                {"role": "system", "content": "You are a curriculum writer. Return structured plain text."},
                {"role": "user", "content": prompt}
            ],
            "temperature": temperature,
            "max_tokens": max_tokens,
            "stream": False
        }

        logging.debug(f"Sending curriculum generation request: {json.dumps(payload, indent=2)}")
        llm_resp = session.post(OLLAMA_URL, json=payload, timeout=60)
        if llm_resp.status_code != 200:
            return jsonify({"error": f"Ollama error {llm_resp.status_code}"}), 500


        content = llm_resp.json().get("message", {}).get("content", "").strip()
        if not content:
            return jsonify({"error": "Empty curriculum response"}), 500
        logging.debug(f"LLM curriculum output: {content}")
        # Parse the structured text response
        lines = content.split("\n")
        module = {
            "title": "",
            "description": "",
            "link": "",
            "questions": []
        }
        current_question = None
        for line in lines:
            line = line.strip()
            if line.startswith("Module Title:"):
                module["title"] = line[len("Module Title:"):].strip()
            elif line.startswith("Description:"):
                module["description"] = line[len("Description:"):].strip()
            elif line.startswith("Module Link:"):
                module["link"] = line[len("Module Link:"):].strip()
            elif line.startswith("Question "):
                if current_question:
                    module["questions"].append(current_question)
                current_question = {
                    "text": line[len("Question "):].strip(),
                    "options": [],
                    "answer": ""
                }
            elif current_question is not None:
                if line.startswith("- "):
                    option_text = line[2:].strip()
                    current_question["options"].append(option_text)
                elif line.startswith("Answer:"):
                    current_question["answer"] = line[len("Answer:"):].strip()
        if current_question:
            module["questions"].append(current_question)
        logging.info(f"Generated curriculum module: {module}")
        # Save the module to a file
        module_filename = f"curriculum_{track}_{level}.json"
        with open(module_filename, "w") as f:
            json.dump(module, f, indent=2)
            logging.info(f"Curriculum module saved to {module_filename}")


        return jsonify({"success": True, "module": module}), 200
    except Exception as e:
        logging.exception("Error in generate_curriculum:")
        return jsonify({"error": str(e)}), 500
    
@app.route("/contacts/<int:contact_id>/drafts", methods=["GET"])
@token_required
def get_contact_drafts(contact_id):
    draft_type = request.args.get("type")  # e.g., ?type=rfp or ?type=email
    drafts = load_all_drafts()
    filtered = [d for d in drafts if d["contact_id"] == contact_id]
    if draft_type:
        filtered = [d for d in filtered if d.get("type", "email") == draft_type]
    return jsonify(filtered)

# Create a new draft for a contact
@app.route("/contacts/<int:contact_id>/drafts", methods=["POST"])
@token_required
def create_contact_draft(contact_id):
    payload = request.get_json() or {}
    subject = payload.get("subject", "Anthesis Agent").strip()  # Default subject
    body = payload.get("body", "").strip()
    draft_type = payload.get("type", "email")  # Default to "email"
    workflow_log = payload.get("workflow_log", None)  # Optional workflow log
    
    # Debug logging
    print(f"Creating draft for contact {contact_id}")
    print(f"Subject: {subject}")
    print(f"Type: {draft_type}")
    print(f"Has workflow_log: {workflow_log is not None}")
    if workflow_log:
        print(f"Workflow log keys: {workflow_log.keys() if isinstance(workflow_log, dict) else 'Not a dict'}")
    
    if not subject or not body:
        return jsonify({"error": "Both subject and body are required."}), 400

    drafts = load_all_drafts()
    draft = {
        "id": next_draft_id(drafts),
        "contact_id": contact_id,
        "subject": subject,
        "body": body,
        "type": draft_type,
        "workflow_log": workflow_log,  # Store workflow log
        "created_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat()
    }
    drafts.append(draft)
    save_all_drafts(drafts)
    
    print(f"Draft saved with ID: {draft['id']}")
    return jsonify(draft), 201

@app.route('/cadet/overview', methods=['GET'])
@token_required
def cadet_overview():
    try:
        username = request.user

        # Load user data
        if os.path.exists("users.json"):
            with open("users.json", "r") as f:
                users = json.load(f)
        else:
            users = {}

        user_data = users.get(username, {})
        track = user_data.get("track", "")

        # Return mock data (you can customize this later)
        return jsonify({
            "username": username,
            "track": track,
            "modulesCompleted": 0,
            "quizzesPassed": 0,
            "hoursLogged": 0,
            "assignments": [
                {"title": f"{track} Project 1", "status": "Not started"},
                {"title": f"{track} Quiz 1", "status": "Locked"}
            ]
        })

    except Exception as e:
        logging.exception("Error in /cadet/overview")
        return jsonify({"error": str(e)}), 500


@app.route('/curriculum/quiz', methods=['OPTIONS'])
def quiz_preflight():
    response = jsonify({'status': 'ok'})
    response.headers.add('Access-Control-Allow-Origin', 'https://alambda.systems')
    response.headers.add('Access-Control-Allow-Methods', 'POST, OPTIONS')
    response.headers.add('Access-Control-Allow-Headers', 'Content-Type, Authorization')
    return response
# Update an existing draft
@app.route("/contacts/<int:contact_id>/drafts/<int:draft_id>", methods=["PUT"])
@token_required
def update_contact_draft(contact_id, draft_id):
    payload = request.get_json() or {}
    drafts = load_all_drafts()
    for d in drafts:
        if d["id"] == draft_id and d["contact_id"] == contact_id:
            d["subject"] = payload.get("subject", d["subject"]).strip()
            d["body"]    = payload.get("body", d["body"]).strip()
            d["updated_at"] = datetime.now(timezone.utc).isoformat()   # updated
            save_all_drafts(drafts)
            return jsonify(d), 200
    return jsonify({"error":"Draft not found."}), 404

# Delete a draft
@app.route("/contacts/<int:contact_id>/drafts/<int:draft_id>", methods=["DELETE"])
@token_required
def delete_contact_draft(contact_id, draft_id):
    drafts = load_all_drafts()
    new_list = [d for d in drafts if not (d["id"]==draft_id and d["contact_id"]==contact_id)]
    if len(new_list) == len(drafts):
        return jsonify({"error":"Draft not found."}), 404
    save_all_drafts(new_list)
    return ("", 204)
from flask import Blueprint, request, jsonify, current_app

@app.route('/lectures/generate', methods=['POST'])
@token_required
def generate_lecture():
    try:
        data = request.get_json() or {}
        track = data.get('track')
        week = data.get('week')  # e.g., 3

        if not track or not week:
            return jsonify(error="track and week required"), 400

        # Use DEFAULT_MODEL from app config or fallback to a default value
        model = current_app.config.get('DEFAULT_MODEL', 'llama3.2')

        payload = {
            "model": model,
            "messages": [
                {"role": "system", "content": "You are an instructional designer."},
                {"role": "user", "content": f"Create a detailed PDF-ready lecture for track '{track}', week {week}."}
            ],
            "temperature": 0.7,
            "max_tokens": 1200,
            "stream": False
        }

        llm_resp = session.post(OLLAMA_URL, json=payload, timeout=60)
        if llm_resp.status_code != 200:
            return jsonify(error=f"LLM error: {llm_resp.status_code}"), 500

        lecture_text = llm_resp.json().get("message", {}).get("content", "").strip()
        if not lecture_text:
            return jsonify(error="Empty lecture"), 500

        # Convert to PDF bytes
        pdf_bytes = text_to_pdf_bytes(lecture_text)

        # Prepare FTP connection
        ftp_host = "107.180.25.236"
        ftp_user = "mrworker@alambda.systems"
        ftp_pass = "Christian123!"

        ftp = FTP()
        ftp.connect(ftp_host, 21, timeout=10)
        ftp.login(ftp_user, ftp_pass)

        # Use naming convention with `-` instead of creating directories
        filename = f"Lecture-{track}-Week{week}.pdf"
        ftp.storbinary(f"STOR {filename}", BytesIO(pdf_bytes))
        ftp.quit()

        # Update progress
        username = request.user
        update_progress(username, "lecture", track, week)

        return jsonify({
            "success": True,
            "message": f"Lecture generated and uploaded successfully",
            "filename": filename
        }), 200

    except Exception as e:
        logging.exception("Error in generate_lecture:")
        return jsonify({"error": str(e)}), 500


def text_to_pdf_bytes(text):
    """Convert text to PDF bytes using the reportlab library."""
    from io import BytesIO
    from reportlab.pdfgen import canvas

    buffer = BytesIO()
    pdf = canvas.Canvas(buffer)
    pdf.drawString(100, 750, text)  # Adjust position as needed
    pdf.save()
    buffer.seek(0)
    return buffer.getvalue()

MODULES_FILE = "modules.json"

def load_modules():
    """Load modules data from file."""
    if os.path.exists(MODULES_FILE):
        with open(MODULES_FILE, "r") as f:
            return json.load(f)
    return {}

def save_modules(modules_data):
    """Save modules data to file."""
    with open(MODULES_FILE, "w") as f:
        json.dump(modules_data, f, indent=2)

LECTURES_FILE = "lectures.json"

def load_lectures():
    """Load lectures data from file."""
    if os.path.exists(LECTURES_FILE):
        with open(LECTURES_FILE, "r") as f:
            return json.load(f)
    return {}

def save_lectures(lectures_data):
    """Save lectures data to file."""
    with open(LECTURES_FILE, "w") as f:
        json.dump(lectures_data, f, indent=2)
QUIZZES_FILE = "quizzes.json"

def load_quizzes():
    """Load quizzes data from file."""
    if os.path.exists(QUIZZES_FILE):
        with open(QUIZZES_FILE, "r") as f:
            return json.load(f)
    return {}

def save_quizzes(quizzes_data):
    """Save quizzes data to file."""
    with open(QUIZZES_FILE, "w") as f:
        json.dump(quizzes_data, f, indent=2)
# Progress tracking
PROGRESS_FILE = "cadet_progress.json"

def load_progress():
    """Load progress data from file."""
    if os.path.exists(PROGRESS_FILE):
        with open(PROGRESS_FILE, "r") as f:
            return json.load(f)
    return {}
@app.route('/modules/list', methods=['GET'])
@token_required
def list_modules():
    try:
        modules = load_modules()
        return jsonify(modules), 200
    except Exception as e:
        logging.exception("Error in list_modules:")
        return jsonify({"error": str(e)}), 500
    
@app.route('/modules/add', methods=['POST'])
@token_required
def add_module():
    try:
        data = request.get_json() or {}
        track = data.get('track', '').strip()
        week = data.get('week')
        title = data.get('title', '').strip()
        description = data.get('description', '').strip()

        if not track or not week or not title:
            return jsonify(error="track, week, and title are required"), 400

        module_id = f"{track}_{week}_{title.replace(' ', '_')}"
        modules = load_modules()
        modules[module_id] = {
            "track": track,
            "week": week,
            "title": title,
            "description": description,
            "created_at": datetime.now().isoformat()
        }
        save_modules(modules)

        return jsonify({"success": True, "module_id": module_id, "message": "Module added successfully"}), 201
    except Exception as e:
        logging.exception("Error in add_module:")
        return jsonify({"error": str(e)}), 500


def save_progress(progress_data):
    """Save progress data to file."""
    with open(PROGRESS_FILE, "w") as f:
        json.dump(progress_data, f, indent=2)

def update_progress(username, activity_type, track, week):
    """Update cadet's progress."""
    progress = load_progress()
    
    if username not in progress:
        progress[username] = {
            "lectures": {},
            "quizzes": {},
            "assignments": {}
        }
    
    if activity_type == "lecture":
        progress[username]["lectures"][f"{track}_{week}"] = {
            "completed": True,
            "timestamp": datetime.now().isoformat()
        }
    elif activity_type == "quiz":
        progress[username]["quizzes"][f"{track}_{week}"] = {
            "completed": True,
            "timestamp": datetime.now().isoformat()
        }
    elif activity_type == "assignment":
        progress[username]["assignments"][f"{track}_{week}"] = {
            "completed": True,
            "timestamp": datetime.now().isoformat()
        }
    
    save_progress(progress)
    return progress[username]

@app.route('/progress', methods=['GET'])
@token_required
def get_progress():
    """Get cadet's progress."""
    try:
        username = request.user
        progress = load_progress()
        
        if username in progress:
            return jsonify(progress[username]), 200
        return jsonify({"error": "No progress data found"}), 404
    except Exception as e:
        logging.exception("Error in get_progress:")
        return jsonify({"error": str(e)}), 500

@app.route('/assignments/create', methods=['POST'])
@token_required
def create_assignment():
    try:
        data = request.get_json() or {}
        track = data.get('track', '').strip()
        week = data.get('week')
        title = data.get('title', '').strip()
        description = data.get('description', '').strip()
        due_date = data.get('due_date')

        if not track or not week or not title:
            return jsonify(error="track, week, and title are required"), 400

        # Generate unique assignment ID
        assignment_id = f"{track}_{week}_{title.replace(' ', '_')}"

        # Save assignment details
        assignments = load_assignments()
        assignments[assignment_id] = {
            "track": track,
            "week": week,
            "title": title,
            "description": description,
            "due_date": due_date,
            "created_at": datetime.now().isoformat(),
            "status": "active"
        }
        save_assignments(assignments)

        # Update progress
        username = request.user
        update_progress(username, "assignment", track, week)

        return jsonify({
            "success": True,
            "assignment_id": assignment_id,
            "message": "Assignment created successfully"
        }), 201

    except Exception as e:
        logging.exception("Error in create_assignment:")
        return jsonify({"error": str(e)}), 500

@app.route('/assignments/list', methods=['GET'])
@token_required
def list_assignments():
    try:
        assignments = load_assignments()
        return jsonify(assignments), 200
    except Exception as e:
        logging.exception("Error in list_assignments:")
        return jsonify({"error": str(e)}), 500

@app.route('/assignments/update', methods=['PUT'])
@token_required
def update_assignment():
    try:
        data = request.get_json() or {}
        assignment_id = data.get('assignment_id')
        status = data.get('status')

        if not assignment_id or not status:
            return jsonify(error="assignment_id and status are required"), 400

        assignments = load_assignments()
        if assignment_id not in assignments:
            return jsonify(error="Assignment not found"), 404

        assignments[assignment_id]["status"] = status
        assignments[assignment_id]["updated_at"] = datetime.now().isoformat()
        save_assignments(assignments)

        return jsonify({
            "success": True,
            "message": "Assignment updated successfully"
        }), 200

    except Exception as e:
        logging.exception("Error in update_assignment:")
        return jsonify({"error": str(e)}), 500

# Assignment data management
ASSIGNMENTS_FILE = "assignments.json"

def load_assignments():
    """Load assignments data from file."""
    if os.path.exists(ASSIGNMENTS_FILE):
        with open(ASSIGNMENTS_FILE, "r") as f:
            return json.load(f)
    return {}

def save_assignments(assignments_data):
    """Save assignments data to file."""
    with open(ASSIGNMENTS_FILE, "w") as f:
        json.dump(assignments_data, f, indent=2)

@app.route('/api/upload', methods=['POST'])
@token_required
def upload_file():
    """Handle file uploads."""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400
            
        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "No selected file"}), 400
            
        # Create uploads directory if it doesn't exist
        upload_dir = "static/uploads"
        os.makedirs(upload_dir, exist_ok=True)
        
        # Save the file
        file_path = os.path.join(upload_dir, file.filename)
        file.save(file_path)
        
        return jsonify({
            "status": "success",
            "message": "File uploaded successfully",
            "file_path": file_path
        })
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# Lecture management
LECTURES_FILE = "static/lectures.json"

def load_lectures():
    """Load lectures from JSON file."""
    try:
        if not os.path.exists(LECTURES_FILE):
            return {}
        with open(LECTURES_FILE, 'r') as f:
            return json.load(f)
    except Exception as e:
        logging.error(f"Error loading lectures: {e}")
        return {}

def save_lectures(lectures):
    """Save lectures to JSON file."""
    try:
        with open(LECTURES_FILE, 'w') as f:
            json.dump(lectures, f, indent=2)
    except Exception as e:
        logging.error(f"Error saving lectures: {e}")

# Quiz management
QUIZZES_FILE = "static/quizzes.json"

def load_quizzes():
    """Load quizzes from JSON file."""
    try:
        if not os.path.exists(QUIZZES_FILE):
            return {}
        with open(QUIZZES_FILE, 'r') as f:
            return json.load(f)
    except Exception as e:
        logging.error(f"Error loading quizzes: {e}")
        return {}

def save_quizzes(quizzes):
    """Save quizzes to JSON file."""
    try:
        with open(QUIZZES_FILE, 'w') as f:
            json.dump(quizzes, f, indent=2)
    except Exception as e:
        logging.error(f"Error saving quizzes: {e}")

@app.route('/api/lectures/<track>', methods=['GET'])
@token_required
def get_lectures(track):
    """Get all lectures for a track."""
    try:
        lectures = load_lectures()
        track_lectures = lectures.get(track, [])
        return jsonify(track_lectures)
    except Exception as e:
        logging.error(f"Error getting lectures: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/lectures', methods=['POST'])
@token_required
def create_lecture():
    """Create a new lecture based on track and week."""
    data = request.get_json()
    track = data.get('track')
    week = data.get('week')
    
    if not track or not week:
        return jsonify({"error": "Track and week are required"}), 400

    try:
        # Generate lecture content using Ollama
        prompt = f"""Generate a detailed lecture for week {week} of the {track} track.
        Include key concepts, examples, and practical applications.
        Format the content in markdown with appropriate headings and sections."""
        
        response = requests.post(
            'http://localhost:11434/api/chat',
            json={
                "model": "llama3.2",
                "temperature": 0.7,
                "messages": [{"role": "user", "content": prompt}],
                "stream": False
            },
            timeout=30
        )
        
        if response.status_code != 200:
            logging.error(f"Ollama API error: {response.text}")
            return jsonify({"error": f"Ollama API error: {response.text}"}), 500
            
        try:
            response_data = response.json()
            # Defensive: handle both dict and list (stream chunk) responses
            if isinstance(response_data, list):
                lecture_content = "".join(
                    chunk.get("message", {}).get("content", "")
                    for chunk in response_data
                    if isinstance(chunk, dict)
                ).strip()
            else:
                lecture_content = response_data.get("message", {}).get("content", "").strip()
        except (json.JSONDecodeError, KeyError, TypeError) as e:
            return jsonify({"error": f"Error parsing Ollama response: {str(e)}"}), 500
        
        if not lecture_content:
            return jsonify({"error": "Empty lecture from LLM"}), 502
        
        # Save lecture
        lectures = load_lectures()
        if track not in lectures:
            lectures[track] = []
        lecture_id = str(uuid.uuid4())
        lecture = {
            "id": lecture_id,
            "track": track,
            "week": week,
            "title": f"Week {week} Lecture - {track} Track",
            "content": lecture_content
        }
        lectures[track].append(lecture)
        save_lectures(lectures)
        return jsonify(lecture)
    except requests.exceptions.RequestException as e:
        logging.error(f"Error connecting to Ollama: {e}")
        return jsonify({"error": "Failed to connect to Ollama service"}), 500
    except Exception as e:
        logging.error(f"Error creating lecture: {e}")
        return jsonify({"error": str(e)}), 500

# --- Do the same for /api/quizzes ---
@app.route('/api/quizzes', methods=['POST'])
@token_required
def create_quiz():
    """Create a new quiz."""
    try:
        data = request.get_json()
        track = data.get('track')
        week = data.get('week')
        
        if not track or not week:
            return jsonify({"error": "Track and week are required"}), 400
            
        # Generate quiz content using Ollama
        prompt = f"Generate a quiz for week {week} of the {track} track. Include 5 multiple choice questions with answers."
        
        payload = {
            "model": DEFAULT_MODEL,
            "messages": [
                {"role": "system", "content": "You are a quiz generator. Generate JSON-formatted quizzes with questions and answers."},
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.7,
            "max_tokens": 500,
            "stream": False
        }
        
        response = session.post(OLLAMA_URL, json=payload, timeout=60)
        if response.status_code != 200:
            return jsonify({"error": f"Ollama error {response.status_code}"}), 500
            
        content = response.json()
        # Defensive: handle both dict and list (stream chunk) responses
        if isinstance(content, list):
            quiz_json_str = "".join(
                chunk.get("message", {}).get("content", "")
                for chunk in content
                if isinstance(chunk, dict)
            ).strip()
        else:
            quiz_json_str = content.get("message", {}).get("content", "").strip()
        try:
            quiz_data = json.loads(quiz_json_str)
        except Exception as e:
            return jsonify({"error": f"Quiz JSON parse error: {str(e)}", "raw": quiz_json_str}), 500
        
        # Load existing quizzes
        quizzes = load_quizzes()
        
        # Add new quiz
        if track not in quizzes:
            quizzes[track] = []
            
        quiz_id = f"{track}_week{week}_{len(quizzes[track]) + 1}"
        quiz = {
            "id": quiz_id,
            "track": track,
            "week": week,
            "title": f"Week {week} Quiz",
            "questions": quiz_data.get("questions", [])
        }
        quizzes[track].append(quiz)
        
        # Save updated quizzes
        save_quizzes(quizzes)
        
        return jsonify(quiz)
        
    except Exception as e:
        logging.error(f"Error creating quiz: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/extract-text', methods=['POST'])

def extract_text():
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400
            
        uploaded_file = request.files['file']
        if uploaded_file.filename == '':
            return jsonify({"error": "No selected file"}), 400
            
        filename = uploaded_file.filename
        ext = os.path.splitext(filename)[1].lower()
        
        # Create extraction log
        extraction_log = {
            "timestamp": datetime.now(timezone.utc).isoformat(),
            "filename": filename,
            "extension": ext,
            "status": "processing",
            "size_bytes": 0,
            "extracted_chars": 0,
            "error": None
        }
        
        # Save to a temporary file to measure size and process
        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            uploaded_file.save(temp_file.name)
            temp_path = temp_file.name
            extraction_log["size_bytes"] = os.path.getsize(temp_path)
            
            try:
                if ext == '.txt':
                    with open(temp_path, 'r', encoding='utf-8', errors='ignore') as f:
                        text = f.read()
                elif ext == '.docx':
                    if not docx:
                        extraction_log["status"] = "failed"
                        extraction_log["error"] = "python-docx not installed"
                        return jsonify({"error": "python-docx not installed", "extraction_log": extraction_log}), 500
                    doc = docx.Document(temp_path)
                    text = '\n'.join([p.text for p in doc.paragraphs])
                elif ext == '.pdf':
                    try:
                        # Try PyMuPDF first
                        import fitz
                        doc = fitz.open(temp_path)
                        text = ""
                        for page in doc:
                            text += page.get_text()
                    except ImportError:
                        # Fall back to PyPDF2
                        if not PyPDF2:
                            extraction_log["status"] = "failed"
                            extraction_log["error"] = "No PDF library installed"
                            return jsonify({"error": "PyPDF2 not installed", "extraction_log": extraction_log}), 500
                        reader = PyPDF2.PdfReader(temp_path)
                        text = ''
                        for page in reader.pages:
                            text += page.extract_text() or ''
                else:
                    extraction_log["status"] = "failed"
                    extraction_log["error"] = f"Unsupported file type: {ext}"
                    return jsonify({"error": f"Unsupported file type: {ext}", "extraction_log": extraction_log}), 400
                
                # Update extraction log with success info
                extraction_log["status"] = "success"
                extraction_log["extracted_chars"] = len(text)
                
                # Save the extraction log to file
                log_dir = "extraction_logs"
                os.makedirs(log_dir, exist_ok=True)
                log_filename = f"extract_log_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}.json"
                log_path = os.path.join(log_dir, log_filename)
                with open(log_path, 'w') as f:
                    json.dump(extraction_log, f, indent=2)
                
                return jsonify({"text": text, "extraction_log": extraction_log})
                
            except Exception as e:
                extraction_log["status"] = "error"
                extraction_log["error"] = str(e)
                
                # Still save the failed extraction log
                log_dir = "extraction_logs"
                os.makedirs(log_dir, exist_ok=True)
                log_filename = f"extract_log_error_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}.json"
                log_path = os.path.join(log_dir, log_filename)
                with open(log_path, 'w') as f:
                    json.dump(extraction_log, f, indent=2)
                    
                return jsonify({"error": str(e), "extraction_log": extraction_log}), 500
            
            finally:
                # Clean up the temp file
                try:
                    os.unlink(temp_path)
                except:
                    pass
                    
    except Exception as e:
        logging.exception("Error in extract_text endpoint:")
        return jsonify({"error": str(e)}), 500


# def conversation():
#     """
#     Handles predefined conversation logic with fuzzy matching for user input.
#     """
#     try:
#         # Get the user question from the request
#         data = request.get_json()
#         user_question = data.get("question", "").strip().lower()

#         # Predefined questions and responses
        
#         responses = {
#             "where is terminal 1 shuttle right now?": (
#                 "According to your location, Terminal 1 shuttle (Shuttle #927) is 30min away from your location. "
#                 "Terminal 1 shuttle (Shuttle #923) is 5min away from your location."
#             ),
#             "i’m at lot a. is my shuttle near?": (
#                 "Could you please be more specific? Are you wanting the Terminal 2 shuttle or Terminal 1 shuttle toward Lambert Airport?"
#             ),
#             "i’ve been waiting 30min for a shuttle. when is lot e shuttle coming!": (
#                 "The Lot E shuttle is currently 10 minutes away from your location. Thank you for your patience!"
#             )
#         }

#         # Normalize predefined questions
#         predefined_questions = list(responses.keys())

#         # Use fuzzy matching to find the closest match
#         match, score = process.extractOne(user_question, predefined_questions, scorer=process.fuzz.ratio)

#         # Define a similarity threshold (e.g., 70%)
#         if score >= 60:
#             response = responses[match]
#         else:
#             response = "I'm sorry, I didn't understand your question. Could you rephrase it?"

#         # Return the response
#         return jsonify({"success": True, "response": response}), 200

#     except Exception as e:
#         logging.exception("Error in conversation endpoint:")
#         return jsonify({"success": False, "error": str(e)}), 500


# Add these endpoints to support the alambdaagent.html functionality

@app.route('/contacts/create', methods=['POST'])
@token_required
def create_agent_contact():
    """Create a new AI persona/agent contact."""
    try:
        data = request.get_json()
        
        # Validate required fields
        required_fields = ['name', 'email', 'bio']
        if not all(field in data and data[field] for field in required_fields):
            return jsonify({"error": "Missing required fields"}), 400

        # Generate unique ID
        contact_id = str(uuid.uuid4())
        
        # Create new contact with agent-specific fields
        new_contact = {
            'id': contact_id,
            'name': data['name'],
            'email': data['email'],
            'bio': data['bio'],
            'type': 'agent',  # Mark as agent type
            'created_by': request.user,
            'created_at': datetime.now(timezone.utc).isoformat(),
            'updated_at': datetime.now(timezone.utc).isoformat()
        }

        # Load existing contacts or create new list
        if os.path.exists('agent_contacts.json'):
            with open('agent_contacts.json', 'r') as f:
                contacts = json.load(f)
        else:
            contacts = []

        contacts.append(new_contact)
        
        # Save updated contacts
        with open('agent_contacts.json', 'w') as f:
            json.dump(contacts, f, indent=2)

        return jsonify(new_contact), 201
        
    except Exception as e:
        logging.exception("Error in create_agent_contact:")
        return jsonify({"error": str(e)}), 500

@app.route('/contacts/list', methods=['POST'])
@token_required
def list_agent_contacts():
    """List all AI agent contacts for the authenticated user."""
    try:
        # Load contacts
        if os.path.exists('agent_contacts.json'):
            with open('agent_contacts.json', 'r') as f:
                all_contacts = json.load(f)
        else:
            all_contacts = []
        
        # Filter to show only contacts created by this user
        user_contacts = [
            contact for contact in all_contacts 
            if contact.get('created_by') == request.user or contact.get('type') == 'public'
        ]
        
        return jsonify(user_contacts)
        
    except Exception as e:
        logging.exception("Error in list_agent_contacts:")
        return jsonify({"error": str(e)}), 500

@app.route('/contacts/get', methods=['GET'])
@token_required
def get_agent_contact():
    """Get a specific AI agent contact by ID."""
    try:
        contact_id = request.args.get('id')
        if not contact_id:
            return jsonify({"error": "ID parameter is required"}), 400
            
        # Load contacts
        if os.path.exists('agent_contacts.json'):
            with open('agent_contacts.json', 'r') as f:
                contacts = json.load(f)
        else:
            return jsonify({"error": "No contacts found"}), 404

        contact = next((c for c in contacts if c['id'] == contact_id), None)
        if not contact:
            return jsonify({"error": "Contact not found"}), 404

        return jsonify(contact)
    except Exception as e:
        logging.exception("Error in get_agent_contact:")
        return jsonify({"error": str(e)}), 500

@app.route('/contacts/update', methods=['POST'])
@token_required
def update_agent_contact():
    """Update an existing AI agent contact."""
    try:
        data = request.get_json()
        contact_id = data.get('id')
        
        if not contact_id:
            return jsonify({"error": "ID parameter is required"}), 400

        # Load contacts
        if os.path.exists('agent_contacts.json'):
            with open('agent_contacts.json', 'r') as f:
                contacts = json.load(f)
        else:
            return jsonify({"error": "No contacts found"}), 404

        contact = next((c for c in contacts if c['id'] == contact_id), None)
        if not contact:
            return jsonify({"error": "Contact not found"}), 404

        # Update contact fields
        contact['name'] = data.get('name', contact['name'])
        contact['email'] = data.get('email', contact['email'])
        contact['bio'] = data.get('bio', contact['bio'])
        contact['updated_at'] = datetime.now(timezone.utc).isoformat()

        with open('agent_contacts.json', 'w') as f:
            json.dump(contacts, f, indent=2)

        return jsonify(contact)
    except Exception as e:
        logging.exception("Error in update_agent_contact:")
        return jsonify({"error": str(e)}), 500

@app.route('/contacts/delete', methods=['POST'])
@token_required
def delete_agent_contact():
    """Delete an AI agent contact."""
    try:
        data = request.get_json()
        contact_id = data.get('id')
        
        if not contact_id:
            return jsonify({"error": "ID parameter is required"}), 400

        # Load contacts
        if os.path.exists('agent_contacts.json'):
            with open('agent_contacts.json', 'r') as f:
                contacts = json.load(f)
        else:
            return jsonify({"error": "No contacts found"}), 404

        contacts = [c for c in contacts if c['id'] != contact_id]
        with open('agent_contacts.json', 'w') as f:
            json.dump(contacts, f, indent=2)

        return jsonify({"message": "Contact deleted successfully"})
    except Exception as e:
        logging.exception("Error in delete_agent_contact:")
        return jsonify({"error": str(e)}), 500


@app.route('/predict', methods=['POST'])
def agent_predict():
    """Process a prompt using the AI agent."""
    try:
        data = request.get_json()
        if not data:
            return jsonify({"error": "Missing JSON data"}), 400
            
        prompt = data.get('prompt')
        agent_id = data.get('agent_id')
        
        if not prompt:
            return jsonify({"error": "Prompt is required"}), 400
            
        # Load agent data if agent_id provided
        agent_context = ""
        if agent_id:
            if os.path.exists('agent_contacts.json'):
                with open('agent_contacts.json', 'r') as f:
                    contacts = json.load(f)
                    agent = next((c for c in contacts if c['id'] == agent_id), None)
                    if agent:
                        agent_context = f"You are acting as {agent['name']}.\nBio: {agent['bio']}\n\n"
        
        # Prepare the full prompt with agent context
        full_prompt = f"{agent_context}{prompt}"
        
        # Call the AI model
        payload = {
            "model": DEFAULT_MODEL,
            "messages": [
                {"role": "system", "content": "You are a helpful AI assistant."},
                {"role": "user", "content": full_prompt}
            ],
            "temperature": 0.7,
            "max_tokens": 500,
            "stream": False
        }
        
        response = session.post(OLLAMA_URL, json=payload, timeout=60)
        if response.status_code != 200:
            return jsonify({"error": f"Model error {response.status_code}"}), 500
            
        content = response.json().get("message", {}).get("content", "").strip()
        
        return jsonify({"answer": content})
        
    except Exception as e:
        logging.exception("Error in agent_predict:")
        return jsonify({"error": str(e)}), 500

# Add these routes to handle preflight OPTIONS requests

@app.route('/login', methods=['OPTIONS'])
def login_preflight():
    response = jsonify({'status': 'ok'})
    response.headers.add('Access-Control-Allow-Origin', 'https://alambda.systems')
    response.headers.add('Access-Control-Allow-Methods', 'POST, OPTIONS')
    response.headers.add('Access-Control-Allow-Headers', 'Content-Type')
    return response

@app.route('/register', methods=['OPTIONS'])
def register_preflight():
    response = jsonify({'status': 'ok'})
    response.headers.add('Access-Control-Allow-Origin', 'https://alambda.systems')
    response.headers.add('Access-Control-Allow-Methods', 'POST, OPTIONS')
    response.headers.add('Access-Control-Allow-Headers', 'Content-Type')
    return response

# Handle preflight requests for drafts endpoints
@app.route('/contacts/<int:contact_id>/drafts', methods=['OPTIONS'])
def drafts_preflight(contact_id):
    response = jsonify({'status': 'ok'})
    response.headers.add('Access-Control-Allow-Origin', '*')
    response.headers.add('Access-Control-Allow-Methods', 'GET, POST, PUT, DELETE, OPTIONS')
    response.headers.add('Access-Control-Allow-Headers', 'Content-Type, Authorization')
    return response

@app.route('/contacts/<int:contact_id>/drafts/<int:draft_id>', methods=['OPTIONS'])
def draft_preflight(contact_id, draft_id):
    response = jsonify({'status': 'ok'})
    response.headers.add('Access-Control-Allow-Origin', '*')
    response.headers.add('Access-Control-Allow-Methods', 'GET, POST, PUT, DELETE, OPTIONS')
    response.headers.add('Access-Control-Allow-Headers', 'Content-Type, Authorization')
    return response

if __name__ == '__main__':
    with app.app_context():
        os.makedirs(SAMPLE_DIR, exist_ok=True)
        os.makedirs("static", exist_ok=True)
        os.makedirs("static/uploads", exist_ok=True)
        generate_pre_generated_music()
    app.run(host='0.0.0.0', port=5000, debug=True)
    # app.run(host='127.0.0.1', port=5000, debug=True)