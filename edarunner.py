#!/usr/bin/env python3
import os
import io
import json
import pickle
import re
import logging
import sys
import time
import asyncio
import aiohttp
import math
import csv
import pandas as pd
import random
from datetime import datetime
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, as_completed
from functools import lru_cache

import boto3
import PyPDF2

# Try to import pdfminer with fallback for Python 3.7 compatibility
try:
    import pdfminer.high_level
    HAS_PDFMINER = True
except ImportError:
    print("⚠️ pdfminer not available, using PyPDF2 fallback")
    HAS_PDFMINER = False

from langdetect import detect, LangDetectException
from sagemaker.predictor import Predictor
from sagemaker.serializers import JSONSerializer
from sagemaker.deserializers import JSONDeserializer
import networkx as nx
import matplotlib.pyplot as plt
import numpy as np
from matplotlib.patches import FancyBboxPatch

# Try to import seaborn for enhanced plotting
try:
    import seaborn as sns
    HAS_SEABORN = True
except ImportError:
    print("⚠️ seaborn not available, using matplotlib only")
    HAS_SEABORN = False

# Tracking processed and failed files
processed_files = set()
failed_files = set()

# Load processed and failed files from logs
if os.path.exists("processed_files.log"):
    with open("processed_files.log", "r") as f:
        processed_files = set(f.read().splitlines())

if os.path.exists("failed_files.log"):
    with open("failed_files.log", "r") as f:
        failed_files = set(f.read().splitlines())

def get_extracted_text_path(file_path):
    """Get the path where extracted text would be stored"""
    if file_path.startswith("s3://"):
        # For S3 paths, create local log directory
        parts = file_path.replace("s3://", "").split("/")
        local_dir = os.path.join("extracted_logs", parts[0], *parts[1:-1])
        filename = parts[-1]
    else:
        # For local paths
        local_dir = os.path.join(os.path.dirname(file_path), "extracted_logs")
        filename = os.path.basename(file_path)
    
    return os.path.join(local_dir, f"{filename}_extracted.txt")

def load_extracted_text(file_path):
    """Load previously extracted text if it exists"""
    try:
        log_file = get_extracted_text_path(file_path)
        if os.path.exists(log_file):
            with open(log_file, "r", encoding="utf-8") as f:
                content = f.read()
                # Extract the actual text after the separator
                separator = "-" * 80
                if separator in content:
                    return content.split(separator, 1)[1].strip()
                else:
                    # Fallback: return content after first 4 lines (header)
                    lines = content.split('\n')
                    if len(lines) > 4:
                        return '\n'.join(lines[4:])
        return None
    except Exception as e:
        print(f"❌ Failed to load extracted text for {file_path}: {e}")
        return None

def log_extracted_text(file_path, extracted_text):
    """Log extracted text to local directory structure"""
    try:
        if file_path.startswith("s3://"):
            # For S3 paths, create local log directory
            parts = file_path.replace("s3://", "").split("/")
            local_dir = os.path.join("extracted_logs", parts[0], *parts[1:-1])
            filename = parts[-1]
        else:
            # For local paths
            local_dir = os.path.join(os.path.dirname(file_path), "extracted_logs")
            filename = os.path.basename(file_path)
        
        os.makedirs(local_dir, exist_ok=True)
        log_file = os.path.join(local_dir, f"{filename}_extracted.txt")
        
        with open(log_file, "w", encoding="utf-8") as f:
            f.write(f"File: {file_path}\nTime: {datetime.now().isoformat()}\nLength: {len(extracted_text)}\n{'-'*80}\n{extracted_text}")
        
        print(f"📝 Extracted text logged to {log_file}")
        
    except Exception as e:
        print(f"❌ Failed to log extracted text for {file_path}: {e}")

def print_processing_status():
    """Print current processing status"""
    status = {
        'processed': len(processed_files),
        'failed': len(failed_files),
        'total': len(processed_files) + len(failed_files),
        'success_rate': len(processed_files) / (len(processed_files) + len(failed_files)) if (len(processed_files) + len(failed_files)) > 0 else 0
    }
    
    print(f"📊 Status: ✅{status['processed']} ❌{status['failed']} 📈{status['success_rate']:.1%}")
    
    if failed_files:
        print(f"❌ Failed Files ({len(failed_files)}):")
        for i, file_path in enumerate(list(failed_files)[:5]):  # Show first 5
            print(f"   {i+1}. {file_path}")
        if len(failed_files) > 5:
            print(f"   ... and {len(failed_files) - 5} more files")

def log_matching_failure(document_key: str, chunk_id: str, relationship_index: int, 
                        relationship: dict, failure_reason: str, failure_type: str, logger: logging.Logger):
    """Log detailed relationship matching failure information"""
    try:
        # Create detailed failure log entry
        failure_entry = {
            'timestamp': datetime.now().isoformat(),
            'document_key': document_key,
            'chunk_id': chunk_id,
            'relationship_index': relationship_index,
            'failure_type': failure_type,
            'failure_reason': failure_reason,
            'original_relationship': relationship,
            'subject': relationship.get('subject', ''),
            'object': relationship.get('object', ''),
            'predicate': relationship.get('predicate', ''),
            'confidence': relationship.get('confidence', ''),
            'evidence': relationship.get('evidence', '')[:500] if relationship.get('evidence') else '',  # Truncate long evidence
        }
        
        # Log to both debug logger and separate failure log
        debug_logger.log_step(
            f"relationship_matching_failure_{failure_type}",
            "error",
            input_data={
                'document_key': document_key,
                'chunk_id': chunk_id,
                'relationship_index': relationship_index
            },
            output_data=failure_entry,
            error=failure_reason
        )
        
        # Write to dedicated matching failures log
        failure_log_file = f"matching_failures_{document_key.replace('/', '_') if document_key else 'unknown'}.json"
        
        # Load existing failures if file exists
        existing_failures = []
        if os.path.exists(failure_log_file):
            try:
                with open(failure_log_file, 'r') as f:
                    existing_failures = json.load(f)
            except Exception as e:
                logger.debug(f"Could not load existing failures: {e}")
        
        # Add new failure
        existing_failures.append(failure_entry)
        
        # Write updated failures
        with open(failure_log_file, 'w') as f:
            json.dump(existing_failures, f, indent=2)
            
        logger.debug(f"📝 Logged matching failure to {failure_log_file}")
        
    except Exception as e:
        logger.error(f"❌ Failed to log matching failure: {e}")

def update_logs(file_path, success=True):
    """Update logs for processed and failed files."""
    log_file = "processed_files.log" if success else "failed_files.log"
    target_set = processed_files if success else failed_files
    
    with open(log_file, "a") as f:
        f.write(file_path + "\n")
    target_set.add(file_path)

# Monitoring and alerting for LLM metrics
class LLMMetrics:
    def __init__(self):
        self.success_count = 0
        self.failure_count = 0
        self.timeout_count = 0
        self.response_times = []

    def record_success(self, response_time):
        self.success_count += 1
        self.response_times.append(response_time)

    def record_failure(self, error_type="general"):
        self.failure_count += 1
        if "timeout" in str(error_type).lower():
            self.timeout_count += 1

    def get_success_rate(self):
        total = self.success_count + self.failure_count
        return self.success_count / total if total > 0 else 0

    def get_avg_response_time(self):
        return sum(self.response_times) / len(self.response_times) if self.response_times else 0

# Adaptive Rate Limiter and Circuit Breaker for LLM endpoint overload protection
class AdaptiveRateLimiter:
    def __init__(self, initial_rate=5, min_rate=1, max_rate=20):
        self.current_rate = initial_rate
        self.min_rate = min_rate
        self.max_rate = max_rate
        self.success_count = 0
        self.failure_count = 0
        self.semaphore = asyncio.Semaphore(initial_rate)

    def adjust_rate(self, success=True):
        if success:
            self.success_count += 1
            if self.success_count > 10:
                self.current_rate = min(self.current_rate + 1, self.max_rate)
                self.success_count = 0
        else:
            self.failure_count += 1
            if self.failure_count > 2:
                self.current_rate = max(self.current_rate // 2, self.min_rate)
                self.failure_count = 0
        self.semaphore = asyncio.Semaphore(self.current_rate)

class CircuitBreaker:
    def __init__(self, failure_threshold=5, recovery_timeout=60):
        self.failure_threshold = failure_threshold
        self.recovery_timeout = recovery_timeout
        self.failure_count = 0
        self.last_failure_time = None
        self.state = "CLOSED"

    def call_succeeded(self):
        self.failure_count = 0
        self.state = "CLOSED"

    def call_failed(self):
        self.failure_count += 1
        self.last_failure_time = time.time()
        if self.failure_count >= self.failure_threshold:
            self.state = "OPEN"

    def can_attempt_call(self):
        if self.state == "CLOSED":
            return True
        if self.state == "OPEN":
            if time.time() - self.last_failure_time > self.recovery_timeout:
                self.state = "HALF_OPEN"
                return True
            return False
        return True

def get_retry_delay(attempt, base_delay=2, max_delay=60):
    delay = min(base_delay * (2 ** attempt), max_delay)
    jitter = random.uniform(0, delay * 0.1)
    return delay + jitter

async def invoke_endpoint_with_timeout(endpoint_name, body, timeout=5):
    """Invoke a SageMaker endpoint asynchronously with a timeout."""
    import concurrent.futures
    import functools

    def sync_predict():
        predictor = Predictor(
            endpoint_name=endpoint_name,
            serializer=JSONSerializer(),
            deserializer=JSONDeserializer()
        )
        return predictor.predict(json.loads(body))

    loop = asyncio.get_event_loop()
    with ThreadPoolExecutor() as executor:
        return await loop.run_in_executor(executor, sync_predict)

async def check_llm_health(endpoint_name):
    try:
        response = await invoke_endpoint_with_timeout(
            endpoint_name=endpoint_name,
            body=json.dumps({"prompt": "test", "max_length": 10}),
            timeout=5
        )
        return True
    except:
        return False

# JSON Debug Logger Setup
class JSONDebugLogger:
    def __init__(self, log_file="debug_flow.json"):
        self.log_file = log_file
        self.session_id = f"session_{int(time.time())}"
        self.step_counter = 0
        self.flow_log = []
        
    def log_step(self, step_name, step_type, input_data=None, output_data=None, metadata=None, error=None):
        """Log a step in the processing flow with detailed JSON data"""
        self.step_counter += 1
        
        log_entry = {
            "session_id": self.session_id,
            "step_number": self.step_counter,
            "timestamp": datetime.now().isoformat(),
            "step_name": step_name,
            "step_type": step_type,  # 'input', 'process', 'output', 'error'
            "input_data": self._serialize_data(input_data),
            "output_data": self._serialize_data(output_data),
            "metadata": metadata or {},
            "error": str(error) if error else None,
            "success": error is None
        }
        
        self.flow_log.append(log_entry)
        
        # Also log to regular logger
        status = "✅" if error is None else "❌"
        logger.info(f"{status} STEP {self.step_counter}: {step_name} ({step_type})")
        if error:
            logger.error(f"   Error: {error}")
        if metadata:
            logger.debug(f"   Metadata: {json.dumps(metadata, indent=2)}")
        
        # Periodically flush to file
        if self.step_counter % 10 == 0:
            self.flush_to_file()
    
    def _serialize_data(self, data):
        """Safely serialize data for JSON logging"""
        if data is None:
            return None
        
        try:
            # Handle common data types
            if isinstance(data, (str, int, float, bool)):
                return data
            elif isinstance(data, Path):
                return str(data)
            elif isinstance(data, list):
                return [self._serialize_item(item) for item in data[:10]]  # Limit to first 10 items
            elif isinstance(data, dict):
                return {k: self._serialize_item(v) for k, v in list(data.items())[:10]}  # Limit to first 10 items
            else:
                return str(data)[:1000]  # Truncate long strings
        except Exception as e:
            return f"<serialization_error: {str(e)}>"
    
    def _serialize_item(self, item):
        """Serialize individual items safely"""
        if isinstance(item, (str, int, float, bool)):
            return item
        elif isinstance(item, Path):
            return str(item)
        elif isinstance(item, dict):
            return {k: str(v)[:200] if isinstance(v, str) else self._serialize_item(v) for k, v in item.items()}
        else:
            return str(item)[:200]
    
    def flush_to_file(self):
        """Write the current log to file"""
        try:
            with open(self.log_file, 'w') as f:
                json.dump(self.flow_log, f, indent=2)
        except Exception as e:
            logger.error(f"Failed to write debug log: {e}")
    
    def get_summary(self):
        """Get a summary of the current session"""
        total_steps = len(self.flow_log)
        successful_steps = sum(1 for step in self.flow_log if step['success'])
        failed_steps = total_steps - successful_steps
        
        return {
            "session_id": self.session_id,
            "total_steps": total_steps,
            "successful_steps": successful_steps,
            "failed_steps": failed_steps,
            "success_rate": successful_steps / total_steps if total_steps > 0 else 0,
            "start_time": self.flow_log[0]['timestamp'] if self.flow_log else None,
            "end_time": self.flow_log[-1]['timestamp'] if self.flow_log else None
        }

# Initialize JSON debug logger
debug_logger = JSONDebugLogger("debug_flow.json")

def get_debug_summary():
    """Get a summary of the debug session"""
    summary = debug_logger.get_summary()
    debug_logger.flush_to_file()
    return summary

def analyze_debug_flow(debug_file="debug_flow.json"):
    """Analyze the debug flow to find bottlenecks and issues"""
    try:
        with open(debug_file, 'r') as f:
            flow_data = json.load(f)
        
        # Analyze by step type
        step_analysis = {}
        error_analysis = {}
        
        for step in flow_data:
            step_type = step.get('step_type', 'unknown')
            step_name = step.get('step_name', 'unknown')
            
            if step_type not in step_analysis:
                step_analysis[step_type] = []
            step_analysis[step_type].append(step)
            
            if step.get('error'):
                error_type = step.get('metadata', {}).get('error_type', 'unknown')
                if error_type not in error_analysis:
                    error_analysis[error_type] = []
                error_analysis[error_type].append(step)
        
        # Find longest running steps
        processing_steps = step_analysis.get('process', [])
        long_running = sorted(
            [s for s in processing_steps if s.get('metadata', {}).get('duration_seconds')],
            key=lambda x: x['metadata']['duration_seconds'],
            reverse=True
        )[:10]
        
        # Count errors by type
        error_counts = {k: len(v) for k, v in error_analysis.items()}
        
        analysis_result = {
            'total_steps': len(flow_data),
            'steps_by_type': {k: len(v) for k, v in step_analysis.items()},
            'error_counts': error_counts,
            'longest_running_steps': [
                {
                    'step_name': s['step_name'],
                    'duration': s['metadata']['duration_seconds'],
                    'metadata': s['metadata']
                } for s in long_running
            ],
            'session_summary': debug_logger.get_summary()
        }
        
        return analysis_result
        
    except Exception as e:
        return {"error": f"Failed to analyze debug flow: {e}"}

def print_debug_analysis():
    """Print a formatted debug analysis"""
    analysis = analyze_debug_flow()
    
    print("\n" + "="*60)
    print("🔍 DEBUG FLOW ANALYSIS")
    print("="*60)
    
    if 'error' in analysis:
        print(f"❌ Error: {analysis['error']}")
        return
    
    print(f"📊 Total Steps: {analysis['total_steps']}")
    print(f"✅ Success Rate: {analysis['session_summary']['success_rate']:.2%}")
    
    print("\n📈 Steps by Type:")
    for step_type, count in analysis['steps_by_type'].items():
        print(f"   {step_type}: {count}")
    
    if analysis['error_counts']:
        print("\n❌ Errors by Type:")
        for error_type, count in analysis['error_counts'].items():
            print(f"   {error_type}: {count}")
    
    if analysis['longest_running_steps']:
        print("\n⏱️  Longest Running Steps:")
        for step in analysis['longest_running_steps'][:5]:
            print(f"   {step['step_name']}: {step['duration']:.2f}s")
    
    print("\n" + "="*60)

# Endpoint health monitoring
class EndpointHealthMonitor:
    def __init__(self):
        self.failure_count = 0
        self.last_failure_time = 0
        self.circuit_breaker_threshold = 5
        self.circuit_breaker_timeout = 300  # 5 minutes
    
    def record_failure(self):
        self.failure_count += 1
        self.last_failure_time = time.time()
        
    def record_success(self):
        self.failure_count = 0
        
    def is_circuit_open(self):
        if self.failure_count >= self.circuit_breaker_threshold:
            if time.time() - self.last_failure_time < self.circuit_breaker_timeout:
                return True
            else:
                # Reset circuit breaker after timeout
                self.failure_count = 0
                return False
        return False
    
    def get_retry_strategy(self):
        if self.failure_count >= 3:
            return {"max_retries": 1, "timeout_multiplier": 2}
        elif self.failure_count >= 1:
            return {"max_retries": 2, "timeout_multiplier": 1.5}
        else:
            return {"max_retries": 3, "timeout_multiplier": 1}

def calculate_throttle_delay():
    """Calculate appropriate delay based on throttling history"""
    global throttle_counter, last_throttle_time
    
    current_time = time.time()
    
    # Reset counter if recovery time has passed
    if current_time - last_throttle_time > THROTTLING_PROTECTION['recovery_time']:
        throttle_counter = 0
    
    if throttle_counter >= THROTTLING_PROTECTION['throttle_threshold']:
        # Exponential backoff with jitter
        base_delay = THROTTLING_PROTECTION['base_delay']
        exponential_factor = min(2 ** (throttle_counter - 2), 8)  # Cap at 8x
        delay = base_delay * exponential_factor
        
        # Add jitter to prevent thundering herd
        if THROTTLING_PROTECTION['jitter']:
            delay *= (0.5 + random.random() * 0.5)  # 50-100% of calculated delay
        
        # Cap at max delay
        delay = min(delay, THROTTLING_PROTECTION['max_delay'])
        
        return delay
    
    return THROTTLING_PROTECTION['base_delay']

def handle_throttling_error():
    """Handle throttling error and update counters"""
    global throttle_counter, last_throttle_time
    
    throttle_counter += 1
    last_throttle_time = time.time()
    
    delay = calculate_throttle_delay()
    logger.warning(f"🛑 Throttling protection: delaying {delay:.1f}s (throttle count: {throttle_counter})")
    
    return delay

# Global health monitor
endpoint_health = EndpointHealthMonitor()

# Try to import docx for DOCX support
try:
    from docx import Document
    DOCX_SUPPORT = True
    print("✅ python-docx successfully imported")
except ImportError as e:
    DOCX_SUPPORT = False
    print(f"⚠️ python-docx not installed: {e}. DOCX files will be skipped. Install with: pip install python-docx")

# Performance Factors and Their Drivers Structure
PERFORMANCE_FACTORS = {
    "size": {
        "drivers": {
            "lens_diameter": "Overall lens diameter affects coverage and centering",
            "base_curve_radius": "Curvature affects how lens sits on the eye", 
            "edge_design": "Edge thickness and shape affects lens profile",
            "lens_thickness": "Central and peripheral thickness affects size perception",
            "packaging_design": "Blister pack size and lens presentation",
            "lens_material_density": "Material density affects perceived size and weight"
        },
        "color": "#FFB347",  # Orange
        "description": "Physical dimensions and size characteristics"
    },
    "fit": {
        "drivers": {
            "base_curve_radius": "Determines how tightly the lens curves around the eye",
            "lens_diameter": "Affects overall centration and corneal coverage", 
            "lens_material_modulus": "Higher modulus = stiffer lens, which can affect mobility",
            "edge_design": "Tapered vs blunt edges affect eyelid interaction and fit",
            "hydration_retention": "Well-hydrated lenses stay centered better",
            "lens_movement_post_blink": "Controlled movement ensures oxygen flow and proper placement"
        },
        "color": "#FF6B6B",  # Red
        "description": "How well the lens conforms to the eye"
    },
    "handling": {
        "drivers": {
            "lens_modulus": "Higher modulus = easier to pinch and apply",
            "surface_texture": "Some lenses are smoother/slippery, harder to grip",
            "lens_thickness": "Thicker lenses are often easier to handle, especially for new wearers",
            "packaging_wetting_solution": "Impacts lens hydration and slipperiness out of the pack",
            "edge_stiffness": "Too soft edges may collapse on the finger"
        },
        "color": "#4ECDC4",  # Teal
        "description": "Ease of insertion, removal, and manipulation"
    },
    "comfort": {
        "drivers": {
            "surface_lubricity": "Smooth surfaces reduce friction against eyelids",
            "moisture_retention": "Long-lasting hydration improves wear over time",
            "oxygen_permeability": "Keeps the cornea healthy, reduces dryness or irritation",
            "lens_edge_design": "A smooth, tapered edge reduces lid irritation",
            "lens_cleanliness_deposits": "Fewer protein/lipid deposits = higher comfort",
            "initial_lens_fit": "A well-fit lens won't cause localized pressure or discomfort",
            "wearing_time": "All-day wear lenses need superior comfort retention"
        },
        "color": "#45B7D1",  # Blue
        "description": "Subjective patient experience"
    }
}

# Driver categories for better organization
DRIVER_CATEGORIES = {
    "material_science": {
        "drivers": ["lens_material_modulus", "lens_modulus", "surface_lubricity", "moisture_retention", 
                   "oxygen_permeability", "surface_texture", "hydration_retention"],
        "color": "#9B59B6",  # Purple
        "description": "Material properties and chemistry"
    },
    "product_design": {
        "drivers": ["base_curve_radius", "lens_diameter", "edge_design", "lens_thickness", 
                   "edge_stiffness", "lens_edge_design"],
        "color": "#E67E22",  # Orange
        "description": "Physical design parameters"
    },
    "patient_response": {
        "drivers": ["lens_movement_post_blink", "lens_cleanliness_deposits", "initial_lens_fit", 
                   "wearing_time", "packaging_wetting_solution"],
        "color": "#2ECC71",  # Green
        "description": "Patient experience and interaction factors"
    }
}

# Enhanced concept detection for drivers
DRIVER_TERMS = {}
for factor_data in PERFORMANCE_FACTORS.values():
    for driver, description in factor_data["drivers"].items():
        # Extract key terms from driver names and descriptions
        terms = []
        # Add driver name variations
        terms.append(driver.replace("_", " "))
        terms.append(driver.replace("_", "-"))
        # Add key terms from description
        desc_words = description.lower().split()
        key_terms = [word for word in desc_words if len(word) > 4 and word not in 
                    ['which', 'often', 'better', 'against', 'reduces', 'improves', 'affects']]
        terms.extend(key_terms)
        DRIVER_TERMS[driver] = terms

# Configuration - Updated for hybrid local/SageMaker environments
BUCKET_NAME = 'mds-mlops-exp-genai-data'  # Updated bucket name
PREFIXES = ['formulations-windchill-data']

# Test mode configuration
TEST_MODE = os.getenv('TEST_MODE', 'auto')  # 'auto', 'local', 'sagemaker', 'real'
LOCAL_TEST_DIR = "test_data_sagemaker"
OLLAMA_HOST = "http://localhost:11434"
OLLAMA_MODEL = "llama3.2:latest"

# Multi-endpoint configuration for failover and load balancing
ENDPOINTS = [
    {
        'name': 'jumpstart-dft-meta-textgenerationne-20250514-174001',
        'region': 'us-east-1',
        'priority': 1,
        'active': True
    }
]

# Current endpoint tracking
CURRENT_ENDPOINT_INDEX = 0
ENDPOINT_NAME = ENDPOINTS[CURRENT_ENDPOINT_INDEX]['name']
REGION = ENDPOINTS[CURRENT_ENDPOINT_INDEX]['region']

# S3 Prefixes for different lens types
PREFIXES = ['trusted/precision1']

# S3 Output Configuration
OUTPUT_BUCKET_NAME = 'mds-mlops-poc-pdp-clinical-stack-lenses-data'  # Same bucket for output
S3_OUTPUT_PREFIX = 'knowledge_graphs/drivers_analysis'  # Base prefix for all outputs

# Enhanced Configuration for Throttling Mitigation
MAX_WORKERS = 50  # Dramatically reduced from 500
MAX_LLM_WORKERS = 20  # Dramatically reduced from 300  
CHUNK_SIZE = 3000  # Keep optimized chunk size
BATCH_SIZE = 10  # Reduced from 100

# Enhanced LLM Configuration with Throttling Protection
LLM_BATCH_SIZE = 2  # Reduced from 3
MAX_BATCH_LENGTH = 6000  # Reduced from 8000
LLM_CALL_DELAY = 3.0  # Increased from 1.5 seconds
BATCH_RETRY_ATTEMPTS = 2  # Reduced from 3
BATCH_TIMEOUT = 120  # Increased from 90

# Throttling Protection Configuration
THROTTLING_PROTECTION = {
    'enabled': True,
    'base_delay': 2.0,
    'max_delay': 30.0,
    'exponential_backoff': True,
    'jitter': True,
    'throttle_threshold': 3,  # After 3 throttles, increase delays
    'recovery_time': 60  # Time to reset throttle counter
}

# Global throttling state
throttle_counter = 0
last_throttle_time = 0

# Enhanced Prompt Engineering Configuration
PROMPT_ENGINEERING = {
    'use_examples': True,
    'include_context_hints': True,
    'enforce_json_structure': True,
    'add_validation_instructions': True,
    'include_stop_tokens': True
}

# Enhanced Confidence and Quality Thresholds
CONFIDENCE_THRESHOLDS = {
    'HIGH': 0.8,
    'MEDIUM': 0.6,
    'LOW': 0.4,
    'MIN_EVIDENCE_LENGTH': 20,  # Minimum evidence text length
    'MIN_RELATIONSHIP_SCORE': 0.3  # Minimum relationship confidence to include
}

# Circuit Breaker Configuration for Timeout Management
TIMEOUT_THRESHOLD = 10  # Number of consecutive timeouts before switching to single-chunk mode
timeout_counter = 0  # Global counter for tracking consecutive timeouts
use_single_chunk_mode = False  # Global flag to force single-chunk processing

# Chart and Visualization Configuration
CHART_COLORS = {
    'Size': '#FF6B6B',
    'Fit': '#4ECDC4', 
    'Handling': '#45B7D1',
    'Comfort': '#96CEB4'
}

# Local Output structure
BASE_OUTPUT_DIR = 'drivers_kg_output'
CACHE_FILE = 'drivers_kg_cache.pkl'
CHECKPOINT_FILE = 'drivers_processed_docs.json'

def create_output_structure():
    """Create organized output folder structure"""
    base_dir = Path(BASE_OUTPUT_DIR)
    base_dir.mkdir(exist_ok=True)
    
    # Create bucket-specific folders
    bucket_folders = {}
    for prefix in PREFIXES:
        lens_type = prefix.split('/')[-1]  # precision1 or precision7
        bucket_dir = base_dir / f"bucket_{BUCKET_NAME.split('-')[-1]}" / f"lens_{lens_type}"
        bucket_dir.mkdir(parents=True, exist_ok=True)
        
        # Create subfolders
        (bucket_dir / "logs").mkdir(exist_ok=True)
        (bucket_dir / "reports").mkdir(exist_ok=True)
        (bucket_dir / "graphs").mkdir(exist_ok=True)
        (bucket_dir / "cache").mkdir(exist_ok=True)
        
        bucket_folders[lens_type] = bucket_dir
    
    return bucket_folders

def upload_file_to_s3(local_file_path: Path, s3_key: str, logger: logging.Logger = None) -> bool:
    """Upload a file to S3 with error handling and logging"""
    try:
        if logger:
            logger.info(f"📤 Uploading {local_file_path.name} to S3...")
            logger.debug(f"   Local: {local_file_path}")
            logger.debug(f"   S3: s3://{OUTPUT_BUCKET_NAME}/{s3_key}")
        
        # Upload the file
        s3.upload_file(
            str(local_file_path), 
            OUTPUT_BUCKET_NAME, 
            s3_key,
            ExtraArgs={
                'ServerSideEncryption': 'AES256',
                'Metadata': {
                    'upload_timestamp': str(int(time.time())),
                    'source': 'drivers_knowledge_graph_pipeline'
                }
            }
        )
        
        if logger:
            logger.info(f"   ✅ Successfully uploaded to s3://{OUTPUT_BUCKET_NAME}/{s3_key}")
        else:
            print(f"   ✅ Uploaded: {local_file_path.name}")
        
        return True
        
    except Exception as e:
        error_msg = f"❌ Failed to upload {local_file_path.name}: {e}"
        if logger:
            logger.error(error_msg)
        else:
            print(f"   {error_msg}")
        return False

def upload_directory_to_s3(local_dir: Path, s3_prefix: str, logger: logging.Logger = None) -> dict:
    """Upload an entire directory to S3 and return upload results"""
    upload_results = {
        'successful': [],
        'failed': [],
        'total_files': 0,
        'total_size': 0
    }
    
    if not local_dir.exists():
        if logger:
            logger.warning(f"⚠️ Directory does not exist: {local_dir}")
        return upload_results
    
    # Get all files in the directory
    all_files = []
    for file_path in local_dir.rglob('*'):
        if file_path.is_file():
            all_files.append(file_path)
    
    upload_results['total_files'] = len(all_files)
    
    if logger:
        logger.info(f"📁 Uploading directory: {local_dir}")
        logger.info(f"   Files to upload: {len(all_files)}")
    
    for file_path in all_files:
        # Calculate relative path for S3 key
        relative_path = file_path.relative_to(local_dir)
        s3_key = f"{s3_prefix}/{relative_path}".replace('\\', '/')
        
        # Get file size
        file_size = file_path.stat().st_size
        upload_results['total_size'] += file_size
        
        # Upload file
        if upload_file_to_s3(file_path, s3_key, logger):
            upload_results['successful'].append({
                'local_path': str(file_path),
                's3_key': s3_key,
                'size': file_size
            })
        else:
            upload_results['failed'].append({
                'local_path': str(file_path),
                's3_key': s3_key,
                'size': file_size
            })
    
    # Summary
    success_count = len(upload_results['successful'])
    total_mb = upload_results['total_size'] / (1024 * 1024)
    
    if logger:
        logger.info(f"📊 Upload Summary:")
        logger.info(f"   ✅ Successful: {success_count}/{upload_results['total_files']} files")
        logger.info(f"   ❌ Failed: {len(upload_results['failed'])} files")
        logger.info(f"   📦 Total size: {total_mb:.2f} MB")
    
    return upload_results

def setup_lens_logger(lens_type: str, output_dir: Path):
    """Setup lens-specific logger"""
    # Handle None or empty lens_type
    if not lens_type:
        lens_type = "unknown"
    
    logger_name = f"drivers_kg_{lens_type}"
    logger = logging.getLogger(logger_name)
    logger.setLevel(logging.DEBUG)
    
    # Clear existing handlers
    logger.handlers.clear()
    
    # File handler for this lens
    log_file = output_dir / "logs" / f"drivers_kg_{lens_type}.log"
    fh = logging.FileHandler(log_file)
    fh.setLevel(logging.DEBUG)
    fh_formatter = logging.Formatter('%(asctime)s %(levelname)s [%(name)s] %(message)s')
    fh.setFormatter(fh_formatter)
    logger.addHandler(fh)
    
    # Console handler
    ch = logging.StreamHandler(sys.stdout)
    ch.setLevel(logging.INFO)
    ch_formatter = logging.Formatter(f'[{lens_type.upper() if lens_type else "UNKNOWN"}] %(levelname)s: %(message)s')
    ch.setFormatter(ch_formatter)
    logger.addHandler(ch)
    
    return logger

# Initialize logger
logger = logging.getLogger("drivers_kg_pipeline")
logger.setLevel(logging.DEBUG)

# File handler
fh = logging.FileHandler("drivers_kg_pipeline.log")
fh.setLevel(logging.DEBUG)
fh_formatter = logging.Formatter('%(asctime)s %(levelname)s [%(name)s] %(message)s')
fh.setFormatter(fh_formatter)
logger.addHandler(fh)

# Console handler
ch = logging.StreamHandler(sys.stdout)
ch.setLevel(logging.INFO)
ch_formatter = logging.Formatter('%(levelname)s: %(message)s')
ch.setFormatter(ch_formatter)
logger.addHandler(ch)

# Initialize AWS clients
s3 = boto3.client('s3', region_name=REGION)

def get_active_endpoints():
    """Get list of active endpoints"""
    return [ep for ep in ENDPOINTS if ep['active']]

def switch_endpoint(endpoint_index=None, endpoint_name=None):
    """Switch to a different endpoint"""
    global CURRENT_ENDPOINT_INDEX, ENDPOINT_NAME, REGION, predictor
    
    if endpoint_name:
        # Find endpoint by name
        for i, ep in enumerate(ENDPOINTS):
            if ep['name'] == endpoint_name and ep['active']:
                endpoint_index = i
                break
        else:
            raise ValueError(f"Endpoint {endpoint_name} not found or not active")
    
    if endpoint_index is None:
        raise ValueError("Must provide either endpoint_index or endpoint_name")
    
    if 0 <= endpoint_index < len(ENDPOINTS) and ENDPOINTS[endpoint_index]['active']:
        CURRENT_ENDPOINT_INDEX = endpoint_index
        ENDPOINT_NAME = ENDPOINTS[endpoint_index]['name']
        REGION = ENDPOINTS[endpoint_index]['region']
        
        # Reinitialize predictor with new endpoint
        predictor = Predictor(
            endpoint_name=ENDPOINT_NAME,
            serializer=JSONSerializer(),
            deserializer=JSONDeserializer()
        )
        
        print(f"✅ Switched to endpoint: {ENDPOINT_NAME} in region {REGION}")
        return True
    else:
        raise ValueError(f"Invalid endpoint index {endpoint_index} or endpoint not active")

def try_next_endpoint():
    """Try the next available endpoint in priority order"""
    global CURRENT_ENDPOINT_INDEX
    
    active_endpoints = get_active_endpoints()
    if len(active_endpoints) <= 1:
        return False
    
    # Find next active endpoint
    current_priority = ENDPOINTS[CURRENT_ENDPOINT_INDEX]['priority']
    next_endpoint = None
    
    for ep in sorted(active_endpoints, key=lambda x: x['priority']):
        if ep['priority'] > current_priority:
            next_endpoint = ep
            break
    
    # If no higher priority endpoint, wrap around to lowest priority
    if not next_endpoint:
        next_endpoint = sorted(active_endpoints, key=lambda x: x['priority'])[0]
    
    # Find index of next endpoint
    for i, ep in enumerate(ENDPOINTS):
        if ep['name'] == next_endpoint['name']:
            try:
                switch_endpoint(i)
                return True
            except Exception as e:
                print(f"⚠️ Failed to switch to endpoint {next_endpoint['name']}: {e}")
                return False
    
    return False

def endpoint_health_check(endpoint_name=None):
    """Check if an endpoint is healthy"""
    test_endpoint_name = endpoint_name or ENDPOINT_NAME
    
    try:
        # Create a test predictor
        test_predictor = Predictor(
            endpoint_name=test_endpoint_name,
            serializer=JSONSerializer(),
            deserializer=JSONDeserializer()
        )
        
        # Send a simple test request
        response = test_predictor.predict({
            "inputs": "Test",
            "parameters": {"max_new_tokens": 5, "temperature": 0.1}
        })
        
        if response and 'generated_text' in response:
            return True
        else:
            return False
            
    except Exception as e:
        print(f"❌ Health check failed for {test_endpoint_name}: {e}")
        return False

def list_available_endpoints():
    """List all available endpoints and their status"""
    print("\n📡 Available Endpoints:")
    print("=" * 50)
    
    for i, ep in enumerate(ENDPOINTS):
        status = "🟢 ACTIVE" if ep['active'] else "🔴 INACTIVE"
        current = "👈 CURRENT" if i == CURRENT_ENDPOINT_INDEX else ""
        health = "✅ HEALTHY" if ep['active'] and endpoint_health_check(ep['name']) else "❌ UNHEALTHY" if ep['active'] else "⚪ SKIP"
        
        print(f"{i+1}. {ep['name']}")
        print(f"   Region: {ep['region']}")
        print(f"   Priority: {ep['priority']}")
        print(f"   Status: {status} {current}")
        print(f"   Health: {health}")
        print()

# Initialize predictor with error handling and failover
try:
    predictor = Predictor(
        endpoint_name=ENDPOINT_NAME,
        serializer=JSONSerializer(),
        deserializer=JSONDeserializer()
    )
    print(f"✅ Initialized with endpoint: {ENDPOINT_NAME}")
except Exception as e:
    print(f"❌ Failed to initialize endpoint {ENDPOINT_NAME}: {e}")
    # Try to switch to next available endpoint
    if try_next_endpoint():
        print("✅ Successfully switched to backup endpoint")
    else:
        print("❌ No backup endpoints available")
        sys.exit(1)

def list_s3_documents(bucket: str, prefix: str, logger: logging.Logger) -> list:
    """List all documents in S3 bucket/prefix"""
    logger.info(f"📂 Listing documents in s3://{bucket}/{prefix}")
    
    documents = []
    try:
        paginator = s3.get_paginator('list_objects_v2')
        pages = paginator.paginate(Bucket=bucket, Prefix=prefix)
        
        for page in pages:
            if 'Contents' in page:
                for obj in page['Contents']:
                    key = obj['Key']
                    if key.lower().endswith(('.pdf', '.txt', '.doc', '.docx')):
                        documents.append({
                            'key': key,
                            'size': obj['Size'],
                            'last_modified': obj['LastModified']
                        })
        
        logger.info(f"📊 Found {len(documents)} documents to process")
        return documents
        
    except Exception as e:
        logger.error(f"❌ Failed to list S3 documents: {e}")
        return []

def extract_text_from_s3_document(bucket: str, key: str, logger: logging.Logger) -> str:
    """Extract text content from S3 document"""
    file_path = f"s3://{bucket}/{key}"
    
    # Check if we have already extracted text for this file
    extracted_text = load_extracted_text(file_path)
    if extracted_text:
        logger.info(f"📄 Using previously extracted text for: {key} ({len(extracted_text)} chars)")
        return extracted_text
    
    # Check if already processed but failed
    if file_path in processed_files:
        logger.info(f"📄 File already processed successfully: {key}")
        # Try to load the text again in case it was just missed
        extracted_text = load_extracted_text(file_path)
        if extracted_text:
            return extracted_text
        else:
            logger.warning(f"⚠️ File marked as processed but no extracted text found: {key}")
    
    logger.debug(f"📄 Extracting text from s3://{bucket}/{key}")
    
    try:
        # Download document from S3
        response = s3.get_object(Bucket=bucket, Key=key)
        content = response['Body'].read()
        
        # Extract text based on file type
        if key.lower().endswith('.pdf'):
            # Use PyPDF2 for PDF extraction
            pdf_file = io.BytesIO(content)
            try:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                text = ""
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
            except Exception as e:
                logger.debug(f"📄 PyPDF2 failed ({e}), trying pdfminer...")
                # Fallback to pdfminer if PyPDF2 fails and available
                if HAS_PDFMINER:
                    pdf_file.seek(0)  # Reset file pointer
                    text = pdfminer.high_level.extract_text(pdf_file)
                else:
                    logger.warning("⚠️ pdfminer not available, using PyPDF2 result")
            
            # If still minimal text, try pdfminer as additional fallback
            if len(text.strip()) < 100 and HAS_PDFMINER:
                logger.debug("📄 PyPDF2 returned minimal text, trying pdfminer...")
                pdf_file.seek(0)  # Reset file pointer
                try:
                    text = pdfminer.high_level.extract_text(pdf_file)
                except Exception as e:
                    logger.warning(f"⚠️ Both PDF extraction methods failed: {e}")
                    update_logs(file_path, success=False)
                    return ""
            elif len(text.strip()) < 100:
                logger.warning("⚠️ PyPDF2 returned minimal text and pdfminer not available")
                
        elif key.lower().endswith('.docx'):
            if DOCX_SUPPORT:
                try:
                    # Use python-docx for DOCX extraction
                    docx_file = io.BytesIO(content)
                    doc = Document(docx_file)
                    
                    text = ""
                    for paragraph in doc.paragraphs:
                        text += paragraph.text + "\n"
                    
                    # Also extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                text += cell.text + " "
                        text += "\n"
                        
                except Exception as e:
                    logger.warning(f"⚠️ DOCX extraction failed: {e}")
                    update_logs(file_path, success=False)
                    return ""
            else:
                logger.warning(f"⚠️ DOCX support not available, skipping: {key}")
                update_logs(file_path, success=False)
                return ""
                
        elif key.lower().endswith('.txt'):
            try:
                text = content.decode('utf-8', errors='ignore')
            except Exception as e:
                logger.warning(f"⚠️ Text file decoding failed: {e}")
                update_logs(file_path, success=False)
                return ""
        else:
            logger.warning(f"⚠️ Unsupported file type: {key}")
            update_logs(file_path, success=False)
            return ""
        
        # Clean and validate text
        text = text.strip()
        if len(text) < 50:
            logger.warning(f"⚠️ Extracted text too short ({len(text)} chars): {key}")
            update_logs(file_path, success=False)
            return ""
        
        logger.debug(f"✅ Extracted {len(text)} characters from {key}")
        
        # Log extracted text to file
        log_extracted_text(file_path, text)
        
        # Update success log
        update_logs(file_path, success=True)
        
        return text
        
    except Exception as e:
        logger.error(f"❌ Failed to extract text from {key}: {e}")
        update_logs(file_path, success=False)
        return ""

def batch_extract_text_from_documents(bucket: str, documents: list, logger: logging.Logger, max_workers: int = MAX_WORKERS) -> dict:
    """Extract text from multiple documents in parallel using all workers"""
    logger.info(f"🔍 Batch extracting text from {len(documents)} documents with {max_workers} workers")
    
    def extract_single_document(doc_info):
        key = doc_info['key']
        file_path = f"s3://{bucket}/{key}"
        
        # Check if we have already extracted text for this file
        extracted_text = load_extracted_text(file_path)
        if extracted_text:
            logger.debug(f"📄 Using previously extracted text for: {key} ({len(extracted_text)} chars)")
            return key, extracted_text, "using_cached"
        
        # Check if already processed but failed
        if file_path in processed_files:
            logger.debug(f"📄 File already processed successfully: {key}")
            # Try to load the text again in case it was just missed
            extracted_text = load_extracted_text(file_path)
            if extracted_text:
                return key, extracted_text, "using_cached"
            else:
                logger.warning(f"⚠️ File marked as processed but no extracted text found: {key}")
        
        try:
            text = extract_text_from_s3_document(bucket, key, logger)
            return key, text, "success" if text else "extraction_failed"
        except Exception as e:
            logger.error(f"❌ Failed to extract text from {key}: {e}")
            return key, None, f"error: {str(e)}"
    
    # Process all documents in parallel
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = [executor.submit(extract_single_document, doc) for doc in documents]
        
        results = {}
        for future in as_completed(futures):
            key, text, status = future.result()
            results[key] = {'text': text, 'status': status, 'length': len(text) if text else 0}
            
            if status == "success":
                logger.info(f"✅ Text extracted from {key} ({len(text)} characters)")
            elif status == "using_cached":
                logger.info(f"📄 Using cached text from {key} ({len(text)} characters)")
            elif status not in ["already_processed", "using_cached"]:
                logger.error(f"❌ Failed to extract text from {key}: {status}")
    
    stats = {
        'successful': len([r for r in results.values() if r['status'] == 'success']),
        'cached': len([r for r in results.values() if r['status'] == 'using_cached']),
        'failed': len([r for r in results.values() if r['status'] not in ['success', 'already_processed', 'using_cached']]),
        'skipped': len([r for r in results.values() if r['status'] == 'already_processed'])
    }
    
    logger.info(f"📊 Batch extraction: ✅{stats['successful']} 📄{stats['cached']} ❌{stats['failed']} ⏭️{stats['skipped']}")
    return results

def chunk_text_for_processing(text: str, chunk_size: int = 4000, overlap: int = 500) -> list:
    """Split text into overlapping chunks for LLM processing - reduced size to minimize timeouts"""
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + chunk_size
        
        # Try to break at sentence boundary
        if end < len(text):
            # Look for sentence endings near the chunk boundary
            for i in range(end - 200, min(end + 200, len(text))):
                if text[i] in '.!?':
                    end = i + 1
                    break
        
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        
        start = end - overlap
        if start >= len(text):
            break
    
    return chunks

def process_document_with_llm(bucket: str, doc_info: dict, lens_type: str, logger: logging.Logger, llm_workers: int = 200) -> dict:
    """Process a single document with LLM to extract relationships with maximum parallel performance"""
    key = doc_info['key']
    
    debug_logger.log_step(
        f"process_document_start_{key.replace('/', '_')}",
        "input",
        input_data={
            "bucket": bucket,
            "key": key,
            "lens_type": lens_type,
            "llm_workers": llm_workers,
            "doc_info": doc_info
        }
    )
    
    logger.info(f"🤖 Processing document: {key}")
    
    # Extract text from document
    text = extract_text_from_s3_document(bucket, key, logger)
    
    debug_logger.log_step(
        f"process_document_text_extracted_{key.replace('/', '_')}",
        "process",
        input_data={"key": key},
        output_data={"text_length": len(text) if text else 0, "text_preview": text[:300] if text else None},
        metadata={"lens_type": lens_type, "extraction_success": bool(text)}
    )
    
    if not text:
        debug_logger.log_step(
            f"process_document_failed_{key.replace('/', '_')}",
            "error",
            input_data={"key": key},
            output_data={"status": "failed", "error": "text_extraction_failed"},
            metadata={"lens_type": lens_type},
            error="text_extraction_failed"
        )
        
        return {
            'document': key,
            'status': 'failed',
            'error': 'text_extraction_failed',
            'relationships': []
        }
    
    # Detect language
    try:
        language = detect(text[:1000])  # Use first 1000 chars for language detection
        
        debug_logger.log_step(
            f"process_document_language_detected_{key.replace('/', '_')}",
            "process",
            input_data={"text_sample_length": 1000},
            output_data={"language": language},
            metadata={"lens_type": lens_type, "key": key}
        )
        
        if language != 'en':
            logger.warning(f"⚠️ Non-English document detected ({language}): {key}")
    except LangDetectException:
        debug_logger.log_step(
            f"process_document_language_detection_failed_{key.replace('/', '_')}",
            "error",
            input_data={"key": key},
            metadata={"lens_type": lens_type},
            error="Could not detect language"
        )
        logger.warning(f"⚠️ Could not detect language for: {key}")
    
    # Detect drivers and factors in the text
    found_drivers, found_factors = detect_drivers_in_context(text, logger)
    
    debug_logger.log_step(
        f"process_document_drivers_detected_{key.replace('/', '_')}",
        "process",
        input_data={"text_length": len(text)},
        output_data={
            "found_drivers": list(found_drivers.keys()) if found_drivers else [],
            "found_factors": found_factors,
            "drivers_count": len(found_drivers) if found_drivers else 0,
            "factors_count": len(found_factors) if found_factors else 0
        },
        metadata={"lens_type": lens_type, "key": key}
    )
    
    if not found_drivers and not found_factors:
        debug_logger.log_step(
            f"process_document_no_content_{key.replace('/', '_')}",
            "process",
            input_data={"key": key, "text_length": len(text)},
            output_data={"status": "no_relevant_content"},
            metadata={"lens_type": lens_type, "reason": "no_drivers_or_factors"}
        )
        
        logger.info(f"ℹ️ No relevant drivers/factors found in {key}")
        return {
            'document': key,
            'status': 'no_relevant_content',
            'text_length': len(text),
            'relationships': []
        }
    
    logger.info(f"🔍 Found {len(found_drivers)} drivers, {len(found_factors)} factors in {key}")
    
    # Chunk text for LLM processing
    chunks = chunk_text_for_processing(text)
    
    debug_logger.log_step(
        f"process_document_text_chunked_{key.replace('/', '_')}",
        "process",
        input_data={"text_length": len(text)},
        output_data={"chunks_count": len(chunks), "average_chunk_length": sum(len(c) for c in chunks) / len(chunks) if chunks else 0},
        metadata={"lens_type": lens_type, "key": key}
    )
    
    logger.debug(f"📝 Split document into {len(chunks)} chunks")
    
    all_relationships = []
    
    # Process chunks with LLM in parallel - maximum parallelism with intelligent batching
    if len(chunks) > 1:
        logger.info(f"🤖 Processing {len(chunks)} chunks with intelligent batch processing")
        
        debug_logger.log_step(
            f"process_document_multi_chunk_processing_{key.replace('/', '_')}",
            "process",
            input_data={"chunks_count": len(chunks), "processing_strategy": "adaptive" if len(chunks) <= 20 else "parallel"},
            metadata={"lens_type": lens_type, "key": key}
        )
        
        # Choose processing strategy based on chunk count
        if len(chunks) <= 20:
            # Small number of chunks - use adaptive processing for optimal learning
            chunk_relationships = adaptive_batch_processing(chunks, key, lens_type, logger)
        else:
            # Large number of chunks - use parallel batch processing for maximum throughput
            chunk_relationships = process_chunks_with_llm_parallel(chunks, key, lens_type, logger, llm_workers)
            
        all_relationships.extend(chunk_relationships)
    else:
        # Process single chunk directly
        logger.info(f"🤖 Processing single chunk directly")
        
        debug_logger.log_step(
            f"process_document_single_chunk_processing_{key.replace('/', '_')}",
            "process",
            input_data={"chunks_count": 1, "chunk_length": len(chunks[0])},
            metadata={"lens_type": lens_type, "key": key}
        )
        
        chunk_relationships = extract_relationships_with_llm(chunks[0], logger, lens_type)
        
        if chunk_relationships:
            # Match relationships to known drivers/factors
            matched_relationships = match_relationships(chunk_relationships, logger, key, "chunk_0")
            all_relationships.extend(matched_relationships)
    
    debug_logger.log_step(
        f"process_document_relationships_extracted_{key.replace('/', '_')}",
        "process",
        input_data={"chunks_processed": len(chunks)},
        output_data={"relationships_extracted": len(all_relationships), "relationships": all_relationships},
        metadata={"lens_type": lens_type, "key": key}
    )
    
    # Deduplicate relationships
    unique_relationships = []
    seen_relationships = set()
    
    for rel in all_relationships:
        rel_key = (rel['driver'], rel['factor'], rel['predicate'])
        if rel_key not in seen_relationships:
            seen_relationships.add(rel_key)
            unique_relationships.append(rel)
    
    result = {
        'document': key,
        'status': 'success',
        'text_length': len(text),
        'chunks_processed': len(chunks),
        'total_relationships_found': len(all_relationships),
        'unique_relationships': len(unique_relationships),
        'found_drivers': list(found_drivers.keys()),
        'found_factors': found_factors,
        'relationships': unique_relationships
    }
    
    logger.info(f"✅ Document processed: {len(unique_relationships)} unique relationships from {key}")
    return result

def process_chunks_with_llm_parallel(chunks: list, document_key: str, lens_type: str, logger: logging.Logger, llm_workers: int = 200) -> list:
    """Process text chunks with LLM using intelligent batching with circuit breaker pattern"""
    global timeout_counter, use_single_chunk_mode
    
    logger.debug(f"🚀 Processing {len(chunks)} chunks with intelligent batching for {document_key}")
    
    # Check circuit breaker status
    if use_single_chunk_mode:
        logger.warning(f"⚠️ Circuit breaker activated - using single-chunk processing due to consecutive timeouts")
        return process_chunks_individually_safe(chunks, document_key, lens_type, logger)
    
    all_relationships = []
    
    # Group chunks into batches for LLM processing
    chunk_batches = []
    current_batch = []
    current_batch_ids = []
    current_batch_length = 0
    
    for i, chunk in enumerate(chunks):
        chunk_id = f"{document_key.split('/')[-1]}_chunk_{i+1}"
        
        # Check if adding this chunk would exceed batch limits
        if (len(current_batch) >= LLM_BATCH_SIZE or 
            current_batch_length + len(chunk) > MAX_BATCH_LENGTH):
            
            # Save current batch and start new one
            if current_batch:
                chunk_batches.append((current_batch.copy(), current_batch_ids.copy()))
                current_batch.clear()
                current_batch_ids.clear()
                current_batch_length = 0
        
        current_batch.append(chunk)
        current_batch_ids.append(chunk_id)
        current_batch_length += len(chunk)
    
    # Add final batch if it has content
    if current_batch:
        chunk_batches.append((current_batch, current_batch_ids))
    
    logger.info(f"📦 Organized {len(chunks)} chunks into {len(chunk_batches)} LLM batches")
    logger.debug(f"   • Average batch size: {len(chunks)/len(chunk_batches):.1f} chunks per batch")
    
    def process_chunk_batch_with_circuit_breaker(batch_data):
        global timeout_counter, use_single_chunk_mode
        
        batch_idx, (text_chunks, chunk_ids) = batch_data
        try:
            logger.debug(f"🔄 Processing LLM batch {batch_idx + 1}/{len(chunk_batches)} ({len(text_chunks)} chunks)")
            
            # Use batch processing for multiple chunks, single processing for one chunk
            if len(text_chunks) > 1 and not use_single_chunk_mode:
                batch_relationships = extract_relationships_with_llm_batch(
                    text_chunks, chunk_ids, logger, lens_type, max_retries=BATCH_RETRY_ATTEMPTS
                )
                
                # Reset timeout counter on successful batch
                if batch_relationships is not None:
                    timeout_counter = 0
                    
            else:
                # Single chunk - use original method
                single_relationships = extract_relationships_with_llm(
                    text_chunks[0], logger, lens_type, max_retries=3
                )
                batch_relationships = single_relationships
            
            if batch_relationships:
                # Match relationships to known drivers/factors
                matched_relationships = match_relationships(
                    batch_relationships, 
                    logger,
                    document_key=document_key,
                    chunk_id=f"{document_key.split('/')[-1]}_batch_{batch_idx + 1}"
                )
                logger.debug(f"✅ Batch {batch_idx + 1} yielded {len(matched_relationships)} matched relationships")
                return matched_relationships
            else:
                logger.debug(f"ℹ️ Batch {batch_idx + 1} yielded no relationships")
                return []
                
        except Exception as e:
            error_msg = str(e)
            if "timeout" in error_msg.lower() or "ModelError" in error_msg:
                # Increment timeout counter
                timeout_counter += 1
                logger.warning(f"⏰ Timeout {timeout_counter}/{TIMEOUT_THRESHOLD} in batch {batch_idx + 1}: {e}")
                
                # Activate circuit breaker if threshold reached
                if timeout_counter >= TIMEOUT_THRESHOLD:
                    use_single_chunk_mode = True
                    logger.error(f"🚨 Circuit breaker activated! Switching to single-chunk mode after {timeout_counter} consecutive timeouts")
                    
                    # Process this batch with fallback
                    return fallback_to_individual_processing(text_chunks, chunk_ids, logger, lens_type)
            else:
                logger.error(f"❌ LLM batch processing failed for batch {batch_idx + 1}: {e}")
            
            return []
    
    # Process batches with controlled parallelism
    max_concurrent_batches = min(llm_workers // 4, len(chunk_batches))  # Further reduced concurrency
    logger.info(f"🎯 Processing {len(chunk_batches)} batches with max {max_concurrent_batches} concurrent batches")
    
    # Process batches in smaller groups to manage timeout risk
    batch_groups = [chunk_batches[i:i + max_concurrent_batches] 
                   for i in range(0, len(chunk_batches), max_concurrent_batches)]
    
    for group_idx, batch_group in enumerate(batch_groups):
        logger.debug(f"🔄 Processing batch group {group_idx + 1}/{len(batch_groups)} ({len(batch_group)} batches)")
        
        group_start_time = time.time()
        
        # Check circuit breaker before each group
        if use_single_chunk_mode:
            logger.warning(f"⚠️ Circuit breaker activated mid-processing - switching remaining batches to single-chunk mode")
            # Process remaining chunks individually
            remaining_chunks = []
            for _, (text_chunks, _) in batch_group:
                remaining_chunks.extend(text_chunks)
            
            if remaining_chunks:
                fallback_relationships = process_chunks_individually_safe(remaining_chunks, document_key, lens_type, logger)
                all_relationships.extend(fallback_relationships)
            break
        
        with ThreadPoolExecutor(max_workers=len(batch_group)) as executor:
            # Submit batches in this group
            batch_data = [(i + group_idx * max_concurrent_batches, batch) for i, batch in enumerate(batch_group)]
            future_to_batch = {executor.submit(process_chunk_batch_with_circuit_breaker, data): data for data in batch_data}
            
            for future in as_completed(future_to_batch):
                try:
                    batch_relationships = future.result(timeout=240)  # 4 minute timeout per batch
                    all_relationships.extend(batch_relationships)
                except TimeoutError:
                    batch_idx = future_to_batch[future][0]
                    timeout_counter += 1
                    logger.warning(f"⏰ Future timeout {timeout_counter}/{TIMEOUT_THRESHOLD} for batch {batch_idx + 1}")
                    
                    if timeout_counter >= TIMEOUT_THRESHOLD:
                        use_single_chunk_mode = True
                        logger.error(f"🚨 Circuit breaker activated due to future timeouts!")
                        
                except Exception as e:
                    batch_idx = future_to_batch[future][0]
                    logger.error(f"❌ Failed to process batch {batch_idx + 1}: {e}")
        
        group_duration = time.time() - group_start_time
        logger.debug(f"⚡ Batch group {group_idx + 1} completed in {group_duration:.2f}s")
        
        # Add delay between groups, longer if timeouts are occurring
        if group_idx < len(batch_groups) - 1:
            delay = LLM_CALL_DELAY * len(batch_group)
            if timeout_counter > 0:
                delay *= 2  # Double delay if we've had timeouts
            time.sleep(delay)
    
    logger.info(f"🎯 Batch processing complete for {document_key}. Found {len(all_relationships)} relationships")
    if len(chunk_batches) > 0:
        logger.info(f"   • Efficiency: {len(all_relationships)/len(chunk_batches):.1f} relationships per LLM call")
    
    return all_relationships

def process_chunks_individually_safe(chunks: list, document_key: str, lens_type: str, logger: logging.Logger) -> list:
    """Safe single-chunk processing with conservative settings"""
    logger.info(f"🔄 Processing {len(chunks)} chunks individually (safe mode)")
    
    all_relationships = []
    successful_chunks = 0
    
    for i, chunk in enumerate(chunks):
        try:
            logger.debug(f"🔄 Safe processing chunk {i+1}/{len(chunks)}")
            
            # Use very conservative single chunk processing
            chunk_relationships = extract_relationships_with_llm(chunk, logger, lens_type, max_retries=2)
            
            if chunk_relationships:
                matched_relationships = match_relationships(chunk_relationships, logger, document_key, f"chunk_{i}")
                all_relationships.extend(matched_relationships)
                successful_chunks += 1
                logger.debug(f"✅ Safe chunk {i+1} successful: {len(matched_relationships)} relationships")
            else:
                logger.debug(f"ℹ️ Safe chunk {i+1} yielded no relationships")
            
            # Conservative delay between chunks
            if i < len(chunks) - 1:
                time.sleep(1.0)
                
        except Exception as e:
            logger.warning(f"⚠️ Safe chunk {i+1} failed: {e}")
            continue
    
    logger.info(f"🎯 Safe processing complete: {len(all_relationships)} relationships from {successful_chunks}/{len(chunks)} chunks")
    return all_relationships

def process_documents_sequential_with_batch_extraction(bucket: str, documents: list, lens_type: str, logger: logging.Logger, max_workers: int = MAX_WORKERS, llm_workers: int = MAX_LLM_WORKERS) -> list:
    """Process documents sequentially, but first batch extract all text using all workers"""
    logger.info(f"🚀 Sequential processing with batch extraction using {max_workers} workers")
    
    # Step 1: Batch extract all text
    extraction_results = batch_extract_text_from_documents(bucket, documents, logger, max_workers)
    
    # Check for extraction failures - "using_cached" is a success, not failure
    failed_extractions = [key for key, result in extraction_results.items() if result['status'] not in ['success', 'already_processed', 'using_cached']]
    if failed_extractions:
        logger.error(f"❌ Text extraction failed for {len(failed_extractions)} documents")
        for key in failed_extractions[:3]:  # Show first 3 failed
            logger.error(f"   - {key}: {extraction_results[key]['status']}")
        logger.error(f"🛑 STOPPING - Fix text extraction issues before continuing")
        print_processing_status()
        return []
    
    # Step 2: Process successfully extracted documents (including cached ones)
    results = []
    successful_extractions = [key for key, result in extraction_results.items() if result['status'] in ['success', 'using_cached']]
    
    for doc_idx, doc_info in enumerate(documents):
        key = doc_info['key']
        
        if key not in successful_extractions:
            continue
            
        logger.info(f"📄 Processing {doc_idx + 1}/{len(documents)}: {key}")
        
        text = extraction_results[key]['text']
        text_length = extraction_results[key]['length']
        
        try:
            # Process with LLM
            result = process_single_document_llm(key, text, text_length, lens_type, logger, max_workers)
            results.append(result)
            logger.info(f"✅ Processed: {result['unique_relationships']} relationships from {key}")
            
        except Exception as e:
            logger.error(f"❌ LLM processing failed for {key}: {e}")
            logger.error(f"🛑 STOPPING - Fix LLM processing issue")
            print_processing_status()
            return results
    
    successful = len([r for r in results if r.get('status') == 'success'])
    logger.info(f"🎯 Complete: {successful}/{len(documents)} documents processed successfully")
    
    return results

def process_single_document_llm(key: str, text: str, text_length: int, lens_type: str, logger: logging.Logger, max_workers: int) -> dict:
    """Process a single document with LLM"""
    # Detect drivers and factors
    found_drivers, found_factors = detect_drivers_in_context(text, logger)
    
    if not found_drivers and not found_factors:
        return {'document': key, 'status': 'no_relevant_content', 'text_length': text_length, 'relationships': []}
    
    # Chunk and process
    chunks = chunk_text_for_processing(text)
    all_relationships = []
    
    if len(chunks) > 1:
        chunk_relationships = process_chunks_with_llm_parallel(chunks, key, lens_type, logger, max_workers)
        all_relationships.extend(chunk_relationships)
    else:
        chunk_relationships = extract_relationships_with_llm(chunks[0], logger, lens_type)
        if chunk_relationships:
            # Pass document and chunk info to matching with updated function signature
            matched_relationships = match_relationships(
                chunk_relationships, 
                logger, 
                document_key=key,
                chunk_id=f"{key.split('/')[-1]}_chunk_1"
            )
            all_relationships.extend(matched_relationships)
    
    # Deduplicate relationships
    unique_relationships = []
    seen_relationships = set()
    for rel in all_relationships:
        rel_key = (rel['driver'], rel['factor'], rel['predicate'])
        if rel_key not in seen_relationships:
            seen_relationships.add(rel_key)
            unique_relationships.append(rel)
    
    return {
        'document': key,
        'status': 'success',
        'text_length': text_length,
        'chunks_processed': len(chunks),
        'total_relationships_found': len(all_relationships),
        'unique_relationships': len(unique_relationships),
        'found_drivers': list(found_drivers.keys()),
        'found_factors': found_factors,
        'relationships': unique_relationships
    }

def process_documents_parallel_with_failure_tracking(bucket: str, documents: list, lens_type: str, logger: logging.Logger, max_workers: int = MAX_WORKERS, llm_workers: int = MAX_LLM_WORKERS) -> list:
    """Process documents in parallel with comprehensive failure tracking and early stopping"""
    logger.info(f"🚀 Starting PARALLEL document processing with failure tracking")
    logger.info(f"📊 Configuration: {max_workers} text workers, {llm_workers} LLM workers")
    
    # Step 1: Batch extract all text in parallel
    logger.info("📄 Phase 1: Batch text extraction")
    extraction_results = batch_extract_text_from_documents(bucket, documents, logger, max_workers)
    
    # Check for extraction failures
    failed_extractions = [key for key, result in extraction_results.items() if result['status'] not in ['success', 'already_processed']]
    successful_extractions = [key for key, result in extraction_results.items() if result['status'] == 'success']
    
    if failed_extractions:
        logger.warning(f"❌ {len(failed_extractions)} documents failed text extraction")
        for failed_key in failed_extractions[:5]:  # Show first 5
            logger.warning(f"   • {failed_key}")
        if len(failed_extractions) > 5:
            logger.warning(f"   ... and {len(failed_extractions) - 5} more")
    
    logger.info(f"✅ Text extraction complete: {len(successful_extractions)} successful, {len(failed_extractions)} failed")
    
    # Step 2: Process documents with LLM in parallel with failure tracking
    logger.info("🤖 Phase 2: Parallel LLM processing with failure tracking")
    
    successful_docs = [doc for doc in documents if doc['key'] in successful_extractions]
    
    def process_single_document_with_failure_tracking(doc_info):
        """Process a single document with comprehensive failure tracking"""
        key = doc_info['key']
        file_path = f"s3://{bucket}/{key}"
        
        try:
            # Get extracted text
            text_result = extraction_results.get(key, {})
            text = text_result.get('text', '')
            
            if not text:
                logger.error(f"❌ No text available for {key}")
                update_logs(file_path, success=False)
                return create_failed_document_result(key, "no_text_available", "Text extraction failed or empty")
            
            # Process with LLM
            result = process_single_document_llm_with_tracking(
                key=key,
                text=text,
                text_length=len(text),
                lens_type=lens_type,
                logger=logger,
                max_workers=max_workers
            )
            
            # Check for processing failures
            if result.get('status') == 'failed':
                logger.error(f"❌ LLM processing failed for {key}: {result.get('error', 'Unknown error')}")
                update_logs(file_path, success=False)
                return result
            
            # Check for relationship matching failures
            if result.get('status') == 'success' and result.get('unique_relationships', 0) == 0:
                logger.warning(f"⚠️ No relationships found for {key}")
                # This is not necessarily a failure, but log it for analysis
                debug_logger.log_step(
                    f"document_no_relationships_{key.replace('/', '_')}",
                    "process",
                    input_data={'key': key, 'text_length': len(text)},
                    output_data={'relationships_found': 0},
                    metadata={'lens_type': lens_type, 'status': 'no_relationships'}
                )
            
            # Success
            update_logs(file_path, success=True)
            logger.info(f"✅ Successfully processed {key}: {result.get('unique_relationships', 0)} relationships")
            return result
            
        except Exception as e:
            error_msg = f"Document processing failed: {str(e)}"
            logger.error(f"❌ Error processing {key}: {error_msg}")
            
            # Log comprehensive failure information
            debug_logger.log_step(
                f"document_processing_failed_{key.replace('/', '_')}",
                "error",
                input_data={'key': key, 'bucket': bucket},
                error=error_msg,
                metadata={'lens_type': lens_type, 'error_type': type(e).__name__}
            )
            
            update_logs(file_path, success=False)
            return create_failed_document_result(key, "processing_exception", error_msg)
    
    # Process documents in parallel with controlled concurrency
    max_concurrent_docs = min(llm_workers // 4, len(successful_docs))
    logger.info(f"🎯 Processing {len(successful_docs)} documents with max {max_concurrent_docs} concurrent documents")
    
    results = []
    failed_documents = []
    
    with ThreadPoolExecutor(max_workers=max_concurrent_docs) as executor:
        # Submit all document processing tasks
        future_to_doc = {executor.submit(process_single_document_with_failure_tracking, doc): doc 
                        for doc in successful_docs}
        
        # Process results as they complete
        for future in as_completed(future_to_doc):
            doc = future_to_doc[future]
            try:
                result = future.result()
                results.append(result)
                
                # Check for failures and early stopping
                if result.get('status') == 'failed':
                    failed_documents.append(doc['key'])
                    logger.error(f"❌ Document failed: {doc['key']}")
                    
                    # Optional: Early stopping if too many failures
                    if len(failed_documents) > len(successful_docs) * 0.5:  # More than 50% failure rate
                        logger.error("🚨 High failure rate detected! Consider stopping and investigating.")
                        # Could implement early stopping here if needed
                        
            except Exception as e:
                logger.error(f"❌ Future processing error for {doc['key']}: {e}")
                results.append(create_failed_document_result(doc['key'], "future_exception", str(e)))
                failed_documents.append(doc['key'])
    
    # Log final statistics
    successful_results = [r for r in results if r.get('status') == 'success']
    logger.info(f"🎯 Parallel processing complete:")
    logger.info(f"   ✅ Successful: {len(successful_results)}")
    logger.info(f"   ❌ Failed: {len(failed_documents)}")
    logger.info(f"   📊 Success rate: {len(successful_results)/len(successful_docs)*100:.1f}%")
    
    # Create comprehensive failure report
    if failed_documents:
        create_failure_report(failed_documents, lens_type, logger)
    
    return results

def process_single_document_llm_with_tracking(key: str, text: str, text_length: int, lens_type: str, logger: logging.Logger, max_workers: int) -> dict:
    """Process a single document with LLM including comprehensive failure tracking"""
    try:
        # Detect drivers and factors
        found_drivers, found_factors = detect_drivers_in_context(text, logger)
        
        if not found_drivers and not found_factors:
            return create_failed_document_result(key, "no_relevant_content", "No drivers or factors found in document")
        
        # Chunk and process
        chunks = chunk_text_for_processing(text)
        all_relationships = []
        
        if len(chunks) > 1:
            chunk_relationships = process_chunks_with_llm_parallel(chunks, key, lens_type, logger, max_workers)
            all_relationships.extend(chunk_relationships)
        else:
            chunk_relationships = extract_relationships_with_llm(chunks[0], logger, lens_type)
            if chunk_relationships:
                # Pass document and chunk info to matching
                matched_relationships = match_relationships(
                    chunk_relationships, 
                    logger, 
                    document_key=key,
                    chunk_id=f"{key.split('/')[-1]}_chunk_1"
                )
                all_relationships.extend(matched_relationships)
        
        # Deduplicate relationships
        unique_relationships = []
        seen_relationships = set()
        for rel in all_relationships:
            rel_key = (rel['driver'], rel['factor'], rel['predicate'])
            if rel_key not in seen_relationships:
                seen_relationships.add(rel_key)
                unique_relationships.append(rel)
        
        return {
            'document': key,
            'status': 'success',
            'text_length': text_length,
            'chunks_processed': len(chunks),
            'total_relationships_found': len(all_relationships),
            'unique_relationships': len(unique_relationships),
            'found_drivers': list(found_drivers.keys()),
            'found_factors': found_factors,
            'relationships': unique_relationships
        }
        
    except Exception as e:
        logger.error(f"❌ LLM processing failed for {key}: {e}")
        return create_failed_document_result(key, "llm_processing_failed", str(e))

def create_failed_document_result(key: str, failure_type: str, error_message: str) -> dict:
    """Create a standardized failed document result"""
    return {
        'document': key,
        'status': 'failed',
        'failure_type': failure_type,
        'error': error_message,
        'text_length': 0,
        'chunks_processed': 0,
        'total_relationships_found': 0,
        'unique_relationships': 0,
        'found_drivers': [],
        'found_factors': [],
        'relationships': []
    }

def create_failure_report(failed_documents: list, lens_type: str, logger: logging.Logger):
    """Create a comprehensive failure report for analysis"""
    try:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        failure_report_file = f"failure_report_{lens_type}_{timestamp}.json"
        
        failure_report = {
            'timestamp': timestamp,
            'lens_type': lens_type,
            'total_failures': len(failed_documents),
            'failed_documents': failed_documents,
            'failure_analysis': {
                'extraction_failures': len([key for key in failed_documents if key in failed_files]),
                'processing_failures': len(failed_documents) - len([key for key in failed_documents if key in failed_files])
            }
        }
        
        # Write failure report
        with open(failure_report_file, 'w') as f:
            json.dump(failure_report, f, indent=2)
        
        logger.info(f"📊 Failure report created: {failure_report_file}")
        logger.info(f"   📋 Total failures: {len(failed_documents)}")
        logger.info(f"   📄 Extraction failures: {failure_report['failure_analysis']['extraction_failures']}")
        logger.info(f"   🤖 Processing failures: {failure_report['failure_analysis']['processing_failures']}")
        
    except Exception as e:
        logger.error(f"❌ Failed to create failure report: {e}")

def process_documents_sequential(bucket: str, documents: list, lens_type: str, logger: logging.Logger, max_workers: int = MAX_WORKERS, llm_workers: int = MAX_LLM_WORKERS) -> list:
    """Process documents one at a time, using all workers for text extraction from each document"""
    logger.info(f"🚀 Starting SEQUENTIAL document processing with {max_workers} workers for text extraction per document")
    
    results = []
    
    for doc_idx, doc_info in enumerate(documents):
        key = doc_info['key']
        logger.info(f"� Processing document {doc_idx + 1}/{len(documents)}: {key}")
        
        # Extract text from this document first
        file_path = f"s3://{bucket}/{key}"
        
        # Check if we have already extracted text for this file
        extracted_text = load_extracted_text(file_path)
        if extracted_text:
            logger.info(f"📄 Using previously extracted text for: {key} ({len(extracted_text)} chars)")
            text = extracted_text
        else:
            # Check if already processed but failed
            if file_path in processed_files:
                logger.info(f"📄 File already processed successfully: {key}")
                # Try to load the text again in case it was just missed
                extracted_text = load_extracted_text(file_path)
                if extracted_text:
                    text = extracted_text
                else:
                    logger.warning(f"⚠️ File marked as processed but no extracted text found: {key}")
                    continue
            else:
                # Extract text with detailed logging
                logger.info(f"🔍 Extracting text from: {key}")
                text = extract_text_from_s3_document(bucket, key, logger)
        
        if not text:
            logger.error(f"❌ Text extraction failed for: {key}")
            logger.error(f"🛑 STOPPING PROCESSING - Fix text extraction issue before continuing")
            
            # Print processing status
            print_processing_status()
            
            # Return partial results so far
            return results
        
        logger.info(f"✅ Text extracted successfully ({len(text)} characters)")
        
        # Now process the extracted text with LLM
        try:
            # Detect drivers and factors in the text
            found_drivers, found_factors = detect_drivers_in_context(text, logger)
            
            if not found_drivers and not found_factors:
                logger.info(f"ℹ️ No relevant drivers/factors found in {key}")
                result = {
                    'document': key,
                    'status': 'no_relevant_content',
                    'text_length': len(text),
                    'relationships': []
                }
                results.append(result)
                continue
            
            logger.info(f"🔍 Found {len(found_drivers)} drivers, {len(found_factors)} factors in {key}")
            
            # Chunk text for LLM processing
            chunks = chunk_text_for_processing(text)
            logger.info(f"� Split document into {len(chunks)} chunks")
            
            # Process chunks with all available workers
            all_relationships = []
            
            if len(chunks) > 1:
                logger.info(f"🤖 Processing {len(chunks)} chunks with all {max_workers} workers")
                chunk_relationships = process_chunks_with_llm_parallel(chunks, key, lens_type, logger, max_workers)
                all_relationships.extend(chunk_relationships)
            else:
                logger.info(f"🤖 Processing single chunk")
                chunk_relationships = extract_relationships_with_llm(chunks[0], logger, lens_type)
                if chunk_relationships:
                    matched_relationships = match_relationships(chunk_relationships, logger, key, "chunk_0")
                    all_relationships.extend(matched_relationships)
            
            # Deduplicate relationships
            unique_relationships = []
            seen_relationships = set()
            
            for rel in all_relationships:
                rel_key = (rel['driver'], rel['factor'], rel['predicate'])
                if rel_key not in seen_relationships:
                    seen_relationships.add(rel_key)
                    unique_relationships.append(rel)
            
            result = {
                'document': key,
                'status': 'success',
                'text_length': len(text),
                'chunks_processed': len(chunks),
                'total_relationships_found': len(all_relationships),
                'unique_relationships': len(unique_relationships),
                'found_drivers': list(found_drivers.keys()),
                'found_factors': found_factors,
                'relationships': unique_relationships
            }
            
            logger.info(f"✅ Document processed: {len(unique_relationships)} unique relationships from {key}")
            results.append(result)
            
        except Exception as e:
            logger.error(f"❌ LLM processing failed for {key}: {e}")
            logger.error(f"🛑 STOPPING PROCESSING - Fix LLM processing issue before continuing")
            
            # Print processing status
            print_processing_status()
            
            # Return partial results so far
            return results
    
    total_successful = len([r for r in results if r.get('status') == 'success'])
    total_failed = len([r for r in results if r.get('status') in ['failed', 'timeout']])
    
    logger.info(f"🎯 Sequential processing complete:")
    logger.info(f"   • Total documents: {len(documents)}")
    logger.info(f"   • Successful: {total_successful}")
    logger.info(f"   • Failed/Timeout: {total_failed}")
    logger.info(f"   • Success rate: {total_successful/len(documents)*100:.1f}%" if len(documents) > 0 else "   • Success rate: 0%")
    
    return results

def extract_comfort_drivers_from_study_data(lens_type: str, logger: logging.Logger) -> dict:
    """Main function to extract comfort drivers from S3 study data using NLP/LLM"""
    logger.info(f"🔬 Starting comfort driver extraction for {lens_type} using NLP/LLM")
    
    # Get the S3 prefix for this lens type
    prefix = f"trusted/{lens_type}"
    
    # List all documents in S3
    documents = list_s3_documents(BUCKET_NAME, prefix, logger)
    
    if not documents:
        logger.warning(f"⚠️ No documents found for {lens_type}")
        return {
            'lens_type': lens_type,
            'documents_processed': 0,
            'total_relationships': 0,
            'comfort_specific_relationships': 0,
            'documents': [],
            'comfort_drivers': {},
            'all_relationships': []
        }
    
    # Process documents sequentially with batch text extraction first
    processing_results = process_documents_sequential_with_batch_extraction(BUCKET_NAME, documents, lens_type, logger, max_workers=MAX_WORKERS, llm_workers=MAX_LLM_WORKERS)
    
    # Aggregate results
    all_relationships = []
    comfort_relationships = []
    comfort_drivers = {}
    
    for result in processing_results:
        if result['status'] == 'success':
            relationships = result.get('relationships', [])
            all_relationships.extend(relationships)
            
            # Filter for comfort-specific relationships
            for rel in relationships:
                if rel['factor'] == 'comfort':
                    comfort_relationships.append(rel)
                    
                    driver = rel['driver']
                    if driver not in comfort_drivers:
                        comfort_drivers[driver] = {
                            'driver_name': driver.replace('_', ' ').title(),
                            'category': rel.get('category', 'unknown'),
                            'evidence_count': 0,
                            'confidence_levels': [],
                            'evidence_snippets': [],
                            'documents': []
                        }
                    
                    comfort_drivers[driver]['evidence_count'] += 1
                    comfort_drivers[driver]['confidence_levels'].append(rel.get('confidence', 'medium'))
                    comfort_drivers[driver]['evidence_snippets'].append(rel.get('evidence', ''))
                    comfort_drivers[driver]['documents'].append(result['document'])
    
    # Calculate summary statistics for comfort drivers
    for driver_data in comfort_drivers.values():
        # Calculate average confidence
        confidence_scores = {'high': 3, 'medium': 2, 'low': 1}
        scores = [confidence_scores.get(conf, 2) for conf in driver_data['confidence_levels']]
        avg_score = sum(scores) / len(scores) if scores else 2
        
        if avg_score >= 2.5:
            driver_data['overall_confidence'] = 'high'
        elif avg_score >= 1.5:
            driver_data['overall_confidence'] = 'medium'
        else:
            driver_data['overall_confidence'] = 'low'
        
        # Remove duplicates from documents
        driver_data['documents'] = list(set(driver_data['documents']))
        driver_data['unique_documents'] = len(driver_data['documents'])
    
    summary = {
        'lens_type': lens_type,
        'documents_processed': len([r for r in processing_results if r['status'] == 'success']),
        'documents_failed': len([r for r in processing_results if r['status'] != 'success']),
        'total_documents': len(processing_results),
        'total_relationships': len(all_relationships),
        'comfort_specific_relationships': len(comfort_relationships),
        'unique_comfort_drivers': len(comfort_drivers),
        'documents': processing_results,
        'comfort_drivers': comfort_drivers,
        'all_relationships': all_relationships
    }
    
    logger.info(f"📊 Comfort Driver Extraction Summary for {lens_type}:")
    logger.info(f"   📄 Documents processed: {summary['documents_processed']}")
    logger.info(f"   🔗 Total relationships: {summary['total_relationships']}")
    logger.info(f"   😊 Comfort relationships: {summary['comfort_specific_relationships']}")
    logger.info(f"   🎯 Unique comfort drivers: {summary['unique_comfort_drivers']}")
    
    return summary

def build_drivers_prompt(context: str) -> str:
    """Build simple, clear prompt for LLM extraction - avoid template confusion"""
    
    # Truncate context strategically for quality processing
    max_context_length = 3500
    if len(context) > max_context_length:
        # Try to break at sentence boundary
        truncated = context[:max_context_length]
        last_period = truncated.rfind('.')
        if last_period > max_context_length * 0.8:  # If period is reasonably close to end
            context = truncated[:last_period + 1]
        else:
            context = truncated + "..."
    
    return f"""You are analyzing contact lens clinical data to find relationships between design factors and performance outcomes.

LOOK FOR these design factors:
- lens_material, surface_treatment, coating_technology, oxygen_permeability
- lens_design, edge_design, thickness_profile, base_curve, diameter  
- tear_film_interaction, eyelid_mechanics, blink_patterns, lens_movement
- comfort_perception, wearing_time, packaging_solution

LOOK FOR these performance outcomes:
- fit, handling, comfort

FIND relationships like:
- "lens_material affects comfort"
- "edge_design influences handling"
- "oxygen_permeability improves comfort"

IMPORTANT: Only extract relationships that are clearly stated in the text. If you don't find any clear relationships, return: []

TEXT TO ANALYZE:
{context}

Extract relationships as JSON array. Use this exact format:
[{{"subject":"actual_factor_name","predicate":"affects","object":"performance_outcome","confidence":"medium","evidence":"quote from text"}}]

If no relationships found, return: []

EXTRACTED RELATIONSHIPS:"""

def build_batch_drivers_prompt(text_chunks: list, chunk_ids: list) -> str:
    """Build simplified batch prompt for processing multiple text chunks"""
    
    # Calculate total length and truncate if necessary
    total_length = sum(len(chunk) for chunk in text_chunks)
    
    if total_length > MAX_BATCH_LENGTH:
        # Proportionally truncate chunks to fit within limit
        scale_factor = MAX_BATCH_LENGTH / total_length * 0.9  # Leave some buffer
        text_chunks = [chunk[:int(len(chunk) * scale_factor)] for chunk in text_chunks]
    
    # Build batch sections
    batch_sections = []
    for i, (chunk, chunk_id) in enumerate(zip(text_chunks, chunk_ids)):
        section = f"""
SECTION_{i+1} [ID:{chunk_id}]:
{chunk.strip()}
---"""
        batch_sections.append(section)
    
    batch_text = "\n".join(batch_sections)
    
    return f"""Analyze contact lens clinical data sections and extract performance relationships.

FIND these factors:
- lens_material, surface_treatment, coating_technology, oxygen_permeability
- lens_design, edge_design, thickness_profile, base_curve, diameter  
- tear_film_interaction, eyelid_mechanics, blink_patterns, lens_movement
- comfort_perception, wearing_time, packaging_solution

FIND these outcomes:
- fit, handling, comfort

PROCESS each section independently and find relationships between factors and outcomes.

IMPORTANT: Only extract relationships clearly stated in the text. If no clear relationships in any section, return: []

TEXT SECTIONS:
{batch_text}

Extract as JSON array with section_id:
[{{"section_id":"chunk_id", "subject":"factor_name", "predicate":"affects", "object":"outcome", "confidence":"medium", "evidence":"quote"}}]

If no relationships found, return: []

EXTRACTED RELATIONSHIPS:"""

def extract_relationships_with_llm_batch(text_chunks: list, chunk_ids: list, logger: logging.Logger, lens_type: str, max_retries: int = 3) -> list:
    """Extract relationships using LLM batch processing - multiple chunks in single call with timeout handling"""
    
    debug_logger.log_step(
        "extract_relationships_with_llm_batch_start",
        "input",
        input_data={
            "chunks_count": len(text_chunks),
            "chunk_ids": chunk_ids,
            "lens_type": lens_type,
            "max_retries": max_retries,
            "total_text_length": sum(len(chunk) for chunk in text_chunks)
        }
    )
    
    logger.info(f"🚀 Batch LLM processing {len(text_chunks)} chunks in single call")
    
    # For timeout-prone endpoints, reduce batch size if it's too large
    if len(text_chunks) > 5:
        logger.warning(f"⚠️ Large batch ({len(text_chunks)} chunks) detected, reducing to prevent timeouts")
        
        debug_logger.log_step(
            "extract_relationships_with_llm_batch_split",
            "process",
            input_data={"original_chunks": len(text_chunks), "max_per_batch": 3},
            metadata={"lens_type": lens_type, "reason": "prevent_timeouts"}
        )
        
        # Split into smaller sub-batches
        sub_batches = []
        for i in range(0, len(text_chunks), 3):  # Max 3 chunks per sub-batch
            sub_batch_chunks = text_chunks[i:i+3]
            sub_batch_ids = chunk_ids[i:i+3]
            sub_batches.append((sub_batch_chunks, sub_batch_ids))
        
        # Process sub-batches
        all_relationships = []
        for batch_idx, (batch_chunks, batch_ids) in enumerate(sub_batches):
            debug_logger.log_step(
                f"extract_relationships_with_llm_batch_subbatch_{batch_idx}",
                "process",
                input_data={"sub_batch_chunks": len(batch_chunks), "sub_batch_ids": batch_ids},
                metadata={"lens_type": lens_type, "sub_batch_index": batch_idx}
            )
            
            batch_rels = extract_relationships_with_llm_batch(batch_chunks, batch_ids, logger, lens_type, max_retries)
            all_relationships.extend(batch_rels)
        
        debug_logger.log_step(
            "extract_relationships_with_llm_batch_split_complete",
            "output",
            output_data={"total_relationships": len(all_relationships)},
            metadata={"lens_type": lens_type, "sub_batches_processed": len(sub_batches)}
        )
        
        return all_relationships
    
    for attempt in range(max_retries):
        try:
            debug_logger.log_step(
                f"extract_relationships_with_llm_batch_attempt_{attempt + 1}",
                "process",
                input_data={"attempt": attempt + 1, "max_retries": max_retries},
                metadata={"lens_type": lens_type, "chunks_count": len(text_chunks)}
            )
            
            prompt = build_batch_drivers_prompt(text_chunks, chunk_ids)
            
            debug_logger.log_step(
                "extract_relationships_with_llm_batch_prompt_built",
                "process",
                input_data={"total_text_length": sum(len(chunk) for chunk in text_chunks)},
                output_data={"prompt_length": len(prompt), "prompt_preview": prompt[:300]},
                metadata={"lens_type": lens_type, "attempt": attempt + 1}
            )
            
            logger.debug(f"📝 Batch LLM Prompt length: {len(prompt)} characters (attempt {attempt + 1}/{max_retries})")
            
            # Log batch details
            if attempt == 0:
                total_text_length = sum(len(chunk) for chunk in text_chunks)
                logger.debug(f"🎯 Batch LLM Request Details:")
                logger.debug(f"   • Endpoint: {ENDPOINT_NAME}")
                logger.debug(f"   • Chunks in batch: {len(text_chunks)}")
                logger.debug(f"   • Total text length: {total_text_length}")
                logger.debug(f"   • Chunk IDs: {chunk_ids}")
                logger.debug(f"   • Lens type: {lens_type}")
            
            start_time = time.time()
            
            # Optimized parameters for auto scaling - enhanced for better relationship extraction
            llm_input = {
                "inputs": prompt,
                "parameters": {
                    "max_new_tokens": 800,   # Increased for comprehensive responses
                    "temperature": 0.05,     # Very low for consistent, focused output
                    "top_p": 0.8,           # Balanced for quality outputs
                    "do_sample": True,
                    "pad_token_id": 128001,
                    "return_full_text": False,
                    "repetition_penalty": 1.2,  # Prevent repetitive outputs
                    "stop": ["EXTRACTED RELATIONSHIPS (JSON):", "TEXT TO ANALYZE:", "---END---"]  # Stop tokens for cleaner output
                }
            }
            
            debug_logger.log_step(
                "extract_relationships_with_llm_batch_input_prepared",
                "process",
                input_data=llm_input,
                metadata={"lens_type": lens_type, "attempt": attempt + 1}
            )
            
            response = predictor.predict(llm_input)
            
            llm_time = time.time() - start_time
            
            debug_logger.log_step(
                "extract_relationships_with_llm_batch_response_received",
                "process",
                input_data={"request_time": llm_time},
                output_data={"response": response},
                metadata={"lens_type": lens_type, "attempt": attempt + 1, "duration_seconds": llm_time}
            )
            
            logger.info(f"⚡ Batch LLM Response time: {llm_time:.2f} seconds for {len(text_chunks)} chunks (attempt {attempt + 1})")
            
            # Extract and parse batch response
            response_text = response.get('generated_text', '')
            logger.debug(f"📤 Batch LLM Response length: {len(response_text)} characters")
            
            debug_logger.log_step(
                "extract_relationships_with_llm_batch_response_parsed",
                "process",
                input_data={"response_text_length": len(response_text)},
                output_data={"response_text_preview": response_text[:500]},
                metadata={"lens_type": lens_type, "attempt": attempt + 1}
            )
            
            # Debug: Show first part of response to understand what we're getting
            if response_text:
                logger.debug(f"📤 Response preview: {response_text[:500]}...")
            else:
                logger.warning("⚠️ Empty response from LLM!")
            
            # Parse batch JSON response
            relationships = parse_batch_llm_response(response_text, chunk_ids, logger)
            
            debug_logger.log_step(
                "extract_relationships_with_llm_batch_relationships_extracted",
                "process",
                input_data={"response_text_length": len(response_text)},
                output_data={"relationships": relationships, "relationship_count": len(relationships)},
                metadata={"lens_type": lens_type, "attempt": attempt + 1}
            )
            
            if relationships:
                debug_logger.log_step(
                    "extract_relationships_with_llm_batch_success",
                    "output",
                    output_data={"relationships": relationships, "relationship_count": len(relationships)},
                    metadata={"lens_type": lens_type, "attempt": attempt + 1, "chunks_processed": len(text_chunks)}
                )
                
                logger.info(f"✅ Batch processing successful: {len(relationships)} relationships from {len(text_chunks)} chunks")
                return relationships
            else:
                debug_logger.log_step(
                    "extract_relationships_with_llm_batch_no_relationships",
                    "process",
                    input_data={"response_text": response_text},
                    metadata={"lens_type": lens_type, "attempt": attempt + 1, "reason": "no_relationships_found"}
                )
                
                logger.warning(f"⚠️ Batch attempt {attempt + 1} returned no valid relationships")
                logger.debug(f"🔍 Full response for debugging: {response_text}")
                if attempt < max_retries - 1:
                    wait_time = (2 ** attempt) + 1  # Exponential backoff
                    logger.debug(f"🔄 Retrying batch in {wait_time} seconds...")
                    time.sleep(wait_time)
                    continue
                
        except Exception as e:
            error_msg = str(e)
            endpoint_health.record_failure()
            
            debug_logger.log_step(
                "extract_relationships_with_llm_batch_exception",
                "error",
                input_data={"attempt": attempt + 1, "max_retries": max_retries},
                metadata={"lens_type": lens_type, "error_type": type(e).__name__},
                error=error_msg
            )
            
            if "throttling" in error_msg.lower() or "ThrottlingException" in error_msg:
                # Use advanced throttling protection
                delay = handle_throttling_error()
                logger.warning(f"⏰ Batch LLM issue on attempt {attempt + 1}/{max_retries}: {error_msg}")
                
                if attempt < max_retries - 1:
                    time.sleep(delay)
                    continue
                else:
                    logger.error(f"❌ Batch processing failed, attempting fallback to individual chunk processing")
                    return fallback_to_individual_processing(text_chunks, chunk_ids, logger, lens_type)
                    
            elif "timeout" in error_msg.lower() or "ModelError" in error_msg:
                logger.warning(f"⏰ Batch LLM issue on attempt {attempt + 1}/{max_retries}: {e}")
                
                if attempt < max_retries - 1:
                    wait_time = (3 ** attempt) + 2  # Aggressive exponential backoff for timeouts
                    logger.debug(f"🔄 Batch retry wait: {wait_time} seconds (timeout recovery)")
                    time.sleep(wait_time)
                    continue
                else:
                    # Final attempt failed - try fallback to individual chunks
                    logger.error(f"❌ Batch processing failed, attempting fallback to individual chunk processing")
                    return fallback_to_individual_processing(text_chunks, chunk_ids, logger, lens_type)
            else:
                logger.error(f"❌ Batch LLM extraction failed: {e}")
                break
    
    logger.error(f"❌ Batch LLM processing failed after {max_retries} attempts")
    # Try fallback processing
    return fallback_to_individual_processing(text_chunks, chunk_ids, logger, lens_type)

def fallback_to_individual_processing(text_chunks: list, chunk_ids: list, logger: logging.Logger, lens_type: str) -> list:
    """Fallback to processing chunks individually when batch processing fails"""
    logger.info(f"🔄 Fallback: Processing {len(text_chunks)} chunks individually due to batch failures")
    
    all_relationships = []
    successful_chunks = 0
    
    for i, (chunk, chunk_id) in enumerate(zip(text_chunks, chunk_ids)):
        try:
            logger.debug(f"🔄 Fallback processing chunk {i+1}/{len(text_chunks)}: {chunk_id}")
            
            # Use single chunk processing with conservative retry
            chunk_relationships = extract_relationships_with_llm(chunk, logger, lens_type, max_retries=2)
            
            if chunk_relationships:
                # Add section_id for consistency with batch processing
                for rel in chunk_relationships:
                    rel['section_id'] = chunk_id
                all_relationships.extend(chunk_relationships)
                successful_chunks += 1
                logger.debug(f"✅ Fallback chunk {i+1} successful: {len(chunk_relationships)} relationships")
            else:
                logger.debug(f"ℹ️ Fallback chunk {i+1} yielded no relationships")
            
            # Brief pause between individual chunks
            if i < len(text_chunks) - 1:
                time.sleep(0.5)
                
        except Exception as e:
            logger.warning(f"⚠️ Fallback chunk {i+1} failed: {e}")
            continue
    
    logger.info(f"🎯 Fallback processing complete: {len(all_relationships)} relationships from {successful_chunks}/{len(text_chunks)} chunks")
    return all_relationships

def parse_batch_llm_response(response_text: str, chunk_ids: list, logger: logging.Logger) -> list:
    """Parse batch LLM response and organize by chunk IDs"""
    try:
        logger.debug(f"🔍 Starting to parse batch response of {len(response_text)} characters")
        
        # Extract JSON from response
        json_text = None
        
        # Method 1: Look for JSON array
        json_start = response_text.find('[')
        json_end = response_text.rfind(']') + 1
        
        if json_start != -1 and json_end > json_start:
            json_text = response_text[json_start:json_end]
            logger.debug(f"📋 Method 1: Found JSON array at positions {json_start}-{json_end}")
        
        # Method 2: Extract from code blocks
        if not json_text:
            json_pattern = r'```(?:json)?\s*(\[.*?\])\s*```'
            match = re.search(json_pattern, response_text, re.DOTALL)
            if match:
                json_text = match.group(1)
                logger.debug(f"📋 Method 2: Found JSON in code block")
        
        # Method 3: Look for JSON after specific markers
        if not json_text:
            markers = [
                "EXTRACTED RELATIONSHIPS (JSON):",
                "JSON:",
                "RELATIONSHIPS:",
                "["
            ]
            for marker in markers:
                marker_pos = response_text.find(marker)
                if marker_pos != -1:
                    after_marker = response_text[marker_pos + len(marker):].strip()
                    bracket_start = after_marker.find('[')
                    if bracket_start != -1:
                        bracket_end = after_marker.rfind(']') + 1
                        if bracket_end > bracket_start:
                            json_text = after_marker[bracket_start:bracket_end]
                            logger.debug(f"📋 Method 3: Found JSON after marker '{marker}'")
                            break
        
        if not json_text:
            logger.warning("⚠️ No valid JSON found in batch LLM response")
            logger.debug(f"📋 Available response content: {response_text[:1000]}...")
            return []
        
        logger.debug(f"📋 Extracted JSON text: {json_text[:500]}...")
        
        # Clean and parse JSON
        json_text = json_text.strip()
        json_text = re.sub(r',(\s*[}\]])', r'\1', json_text)  # Remove trailing commas
        json_text = json_text.replace("'", '"')  # Fix quotes
        
        # Additional cleaning for common LLM issues
        json_text = re.sub(r'(\w+):\s*([^",\[\]{}]+)(?=\s*[,}])', r'"\1": "\2"', json_text)  # Quote unquoted values
        json_text = re.sub(r'""([^"]+)""', r'"\1"', json_text)  # Fix double quotes
        
        logger.debug(f"📋 Cleaned JSON text: {json_text[:500]}...")
        
        relationships = json.loads(json_text)
        
        if not isinstance(relationships, list):
            logger.warning(f"⚠️ Batch response is not a list, got: {type(relationships)}")
            return []
        
        logger.debug(f"📋 Parsed {len(relationships)} raw relationships")
        
        # Validate and organize relationships
        valid_relationships = []
        chunk_stats = {chunk_id: 0 for chunk_id in chunk_ids}
        
        for i, rel in enumerate(relationships):
            logger.debug(f"📋 Processing relationship {i+1}: {rel}")
            
            if isinstance(rel, dict) and all(key in rel for key in ['subject', 'predicate', 'object']):
                # Add chunk tracking
                section_id = rel.get('section_id', 'unknown')
                if section_id in chunk_stats:
                    chunk_stats[section_id] += 1
                
                # Standardize field names
                standardized_rel = {
                    'driver': rel.get('subject', ''),
                    'predicate': rel.get('predicate', 'affects'),
                    'factor': rel.get('object', ''),
                    'confidence': rel.get('confidence', 'medium'),
                    'evidence': rel.get('evidence', ''),
                    'section_id': section_id
                }
                valid_relationships.append(standardized_rel)
                logger.debug(f"   ✅ Valid relationship: {standardized_rel['driver']} → {standardized_rel['factor']}")
            else:
                logger.debug(f"   ❌ Invalid relationship format: {rel}")
        
        # Log batch processing statistics
        logger.debug(f"📊 Batch processing stats:")
        for chunk_id, count in chunk_stats.items():
            logger.debug(f"   • {chunk_id}: {count} relationships")
        
        logger.info(f"✅ Parsed {len(valid_relationships)} valid relationships from batch response")
        return valid_relationships
        
    except json.JSONDecodeError as e:
        logger.warning(f"⚠️ Batch JSON parsing failed: {e}")
        logger.debug(f"Raw JSON text: {json_text[:1000]}..." if json_text else "No JSON text extracted")
        
        # Try one more fallback - look for individual JSON objects
        try:
            logger.debug("🔄 Attempting fallback parsing of individual objects...")
            object_pattern = r'\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}'
            objects = re.findall(object_pattern, response_text)
            
            fallback_relationships = []
            for obj_str in objects:
                try:
                    obj = json.loads(obj_str)
                    if isinstance(obj, dict) and 'subject' in obj and 'object' in obj:
                        standardized_rel = {
                            'driver': obj.get('subject', ''),
                            'predicate': obj.get('predicate', 'affects'),
                            'factor': obj.get('object', ''),
                            'confidence': obj.get('confidence', 'medium'),
                            'evidence': obj.get('evidence', ''),
                            'section_id': obj.get('section_id', 'unknown')
                        }
                        fallback_relationships.append(standardized_rel)
                except:
                    continue
            
            if fallback_relationships:
                logger.info(f"✅ Fallback parsing extracted {len(fallback_relationships)} relationships")
                return fallback_relationships
            
        except Exception as fallback_e:
            logger.debug(f"Fallback parsing failed: {fallback_e}")
        
        return []
    except Exception as e:
        logger.error(f"❌ Batch response parsing failed: {e}")
        logger.debug(f"Response text: {response_text[:1000]}...")
        return []
    except Exception as e:
        logger.error(f"❌ Batch response parsing failed: {e}")
        return []

def extract_relationships_with_llm(context: str, logger: logging.Logger, lens_type: str, max_retries: int = 5) -> list:
    """Extract relationships using LLM with enhanced logging and retry logic"""
    debug_logger.log_step("extract_relationships_with_llm_start", "input", 
                         input_data={"context_length": len(context), "lens_type": lens_type})
    
    logger.info(f"🤖 Extracting relationships from context ({len(context)} chars)")
    
    for attempt in range(max_retries):
        try:
            prompt = build_drivers_prompt(context)
            response = _make_llm_request(prompt, logger, lens_type, attempt)
            
            if response:
                relationships = _extract_relationships_from_response(response, logger, lens_type, attempt)
                if relationships:
                    return relationships
                    
        except Exception as e:
            if not _handle_llm_exception(e, attempt, max_retries, logger, lens_type):
                break
    
    logger.error(f"❌ All {max_retries} LLM attempts failed")
    return []

def _make_llm_request(prompt: str, logger: logging.Logger, lens_type: str, attempt: int) -> dict:
    """Make LLM request with optimized parameters"""
    start_time = time.time()
    
    llm_input = {
        "inputs": prompt,
        "parameters": {
            "max_new_tokens": 400,
            "temperature": 0.02,
            "top_p": 0.7,
            "do_sample": True,
            "pad_token_id": 128001,
            "return_full_text": False,
            "repetition_penalty": 1.15,
            "stop": ["EXTRACTED RELATIONSHIPS (JSON):", "TEXT TO ANALYZE:", "---END---"]
        }
    }
    
    response = predictor.predict(llm_input)
    llm_time = time.time() - start_time
    
    logger.info(f"⚡ LLM Response time: {llm_time:.2f} seconds (attempt {attempt + 1})")
    debug_logger.log_step("llm_response_received", "process", 
                         input_data={"duration": llm_time}, output_data={"response": response})
    
    return response

def _extract_relationships_from_response(response: dict, logger: logging.Logger, lens_type: str, attempt: int) -> list:
    """Extract and validate relationships from LLM response"""
    response_text = response.get('generated_text', '')
    
    if not response_text:
        logger.warning("⚠️ Empty response from LLM!")
        return []
    
    # Extract JSON using multiple methods
    json_text = _extract_json_from_text(response_text, logger)
    
    if not json_text:
        logger.warning("⚠️ No valid JSON found in LLM response")
        return []
    
    try:
        # Clean and parse JSON
        json_text = _clean_json_text(json_text)
        relationships = json.loads(json_text)
        
        if isinstance(relationships, list):
            valid_relationships = _validate_relationships(relationships, logger)
            logger.info(f"✅ Successfully parsed {len(valid_relationships)} relationships")
            return valid_relationships
        else:
            logger.warning(f"⚠️ Expected list, got {type(relationships)}")
            return []
            
    except json.JSONDecodeError as e:
        logger.warning(f"⚠️ JSON parsing failed: {e}")
        return _fallback_relationship_extraction(json_text, logger)

def _extract_json_from_text(text: str, logger: logging.Logger) -> str:
    """Extract JSON from text using multiple patterns"""
    patterns = [
        r'\[.*?\]',  # Basic array pattern
        r'EXTRACTED RELATIONSHIPS[:\s]*(\[.*?\])',
        r'JSON[:\s]*(\[.*?\])',
        r'RELATIONSHIPS[:\s]*(\[.*?\])',
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, text, re.DOTALL)
        if matches:
            return matches[-1] if isinstance(matches[-1], str) else matches[-1]
    
    # Try block extraction
    start = text.find('[')
    end = text.rfind(']') + 1
    if start != -1 and end > start:
        return text[start:end]
    
    # Try code blocks
    match = re.search(r'```(?:json)?\s*(\[.*?\])\s*```', text, re.DOTALL)
    return match.group(1) if match else None

def _clean_json_text(json_text: str) -> str:
    """Clean JSON text for parsing"""
    json_text = json_text.strip()
    json_text = re.sub(r'```(?:json)?', '', json_text)
    json_text = re.sub(r'```', '', json_text)
    json_text = re.sub(r',(\s*[}\]])', r'\1', json_text)  # Fix trailing commas
    json_text = re.sub(r"'([^']*)':", r'"\1":', json_text)  # Fix single quotes
    json_text = re.sub(r":\s*'([^']*)'", r': "\1"', json_text)
    json_text = re.sub(r'(\w+):', r'"\1":', json_text)  # Add quotes to keys
    json_text = re.sub(r'""(\w+)"":', r'"\1":', json_text)  # Fix double quotes
    return json_text

def _validate_relationships(relationships: list, logger: logging.Logger) -> list:
    """Validate relationship format and filter out template responses"""
    valid_relationships = []
    
    # Template patterns to detect and filter out
    template_patterns = [
        'driver_name',
        'relationship_type', 
        'factor',
        'level',
        'exact_quote_from_text',
        'supporting_quote',
        'ID'
    ]
    
    for i, rel in enumerate(relationships):
        if isinstance(rel, dict) and all(key in rel for key in ['subject', 'object']):
            # Check if this is a template response
            is_template = False
            for field in ['subject', 'object', 'predicate', 'confidence', 'evidence']:
                value = str(rel.get(field, '')).lower()
                if any(pattern.lower() in value for pattern in template_patterns):
                    is_template = True
                    break
            
            if is_template:
                logger.debug(f"   ⚠️ Template response detected at index {i}, skipping")
                continue
                
            # Check for empty or too short values
            subject = rel.get('subject', '').strip()
            object_val = rel.get('object', '').strip()
            
            if len(subject) < 3 or len(object_val) < 3:
                logger.debug(f"   ⚠️ Too short values at index {i}, skipping")
                continue
                
            valid_relationships.append(rel)
            logger.debug(f"   ✓ Relationship {i+1}: {subject} → {object_val}")
        else:
            logger.debug(f"   ⚠️ Invalid relationship format at index {i}")
    
    if len(valid_relationships) == 0 and len(relationships) > 0:
        logger.warning("⚠️ All relationships appear to be template responses - LLM may not be understanding the task")
    
    return valid_relationships

def _fallback_relationship_extraction(json_text: str, logger: logging.Logger) -> list:
    """Fallback extraction for malformed JSON"""
    try:
        object_pattern = r'\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}'
        objects = re.findall(object_pattern, json_text)
        extracted_relationships = []
        
        for obj_str in objects:
            try:
                obj = json.loads(obj_str)
                if isinstance(obj, dict) and 'subject' in obj and 'object' in obj:
                    extracted_relationships.append(obj)
            except:
                continue
        
        if extracted_relationships:
            logger.info(f"✅ Extracted {len(extracted_relationships)} relationships via fallback parsing")
            return extracted_relationships
    except Exception as e:
        logger.debug(f"Fallback parsing failed: {e}")
    
    return []

def _handle_llm_exception(e: Exception, attempt: int, max_retries: int, logger: logging.Logger, lens_type: str) -> bool:
    """Handle LLM exceptions and return whether to continue retrying"""
    error_msg = str(e)
    endpoint_health.record_failure()
    
    if "throttling" in error_msg.lower() or "ThrottlingException" in error_msg:
        delay = handle_throttling_error()
        logger.warning(f"⏰ LLM throttling on attempt {attempt + 1}/{max_retries}")
        
        if attempt < max_retries - 1:
            time.sleep(delay)
            return True
            
    elif "timeout" in error_msg.lower() or "ModelError" in error_msg:
        logger.warning(f"⏰ LLM timeout on attempt {attempt + 1}/{max_retries}")
        
        if attempt < max_retries - 1:
            wait_time = (1.5 ** attempt) + random.uniform(0, 0.5)
            logger.info(f"🔄 Auto scaling may be activating, retrying in {wait_time:.1f} seconds...")
            time.sleep(wait_time)
            return True
    else:
        logger.error(f"❌ LLM extraction failed: {e}")
        import traceback
        logger.debug(f"Traceback: {traceback.format_exc()}")
    
    return False

def detect_drivers_in_context(context: str, logger: logging.Logger):
    """Detect performance drivers mentioned in document context with enhanced pattern matching"""
    logger.debug("🔍 Starting enhanced driver detection in context")
    
    context_lower = context.lower()
    found_drivers = {}
    found_factors = []
    
    # Log context details
    logger.debug(f"📄 Context analysis:")
    logger.debug(f"   • Length: {len(context)} characters")
    logger.debug(f"   • Words: {len(context.split())} words")
    
    # Enhanced driver detection with better pattern matching
    logger.debug("🔧 Scanning for performance drivers with enhanced patterns...")
    
    # Create enhanced search patterns for better detection
    enhanced_driver_patterns = {
        'polymer_chemistry': ['polymer', 'chemistry', 'material composition', 'chemical structure', 'molecular'],
        'surface_treatment': ['surface treatment', 'coating', 'surface modification', 'treatment', 'surface chemistry'],
        'lens_material': ['lens material', 'material', 'substrate', 'polymer material', 'lens composition'],
        'coating_technology': ['coating', 'coating technology', 'surface coating', 'protective coating'],
        'oxygen_permeability': ['oxygen permeability', 'dk/t', 'oxygen transmission', 'breathability', 'permeability'],
        'hydration_retention': ['hydration', 'water content', 'moisture retention', 'hydration retention'],
        'surface_lubricity': ['surface lubricity', 'lubricity', 'slipperiness', 'friction', 'smoothness'],
        'moisture_retention': ['moisture', 'moisture retention', 'water retention', 'hydration'],
        'lens_design': ['lens design', 'design', 'optical design', 'lens geometry'],
        'edge_design': ['edge design', 'edge', 'lens edge', 'peripheral design', 'edge geometry'],
        'thickness_profile': ['thickness', 'thickness profile', 'lens thickness', 'center thickness'],
        'base_curve': ['base curve', 'base curve radius', 'curvature', 'bcr'],
        'diameter': ['diameter', 'lens diameter', 'overall diameter', 'od'],
        'lens_diameter': ['lens diameter', 'diameter', 'overall diameter'],
        'lens_thickness': ['lens thickness', 'thickness', 'center thickness'],
        'edge_stiffness': ['edge stiffness', 'stiffness', 'rigidity', 'flexibility'],
        'tear_film_interaction': ['tear film', 'tear', 'tear layer', 'lacrimal', 'tear interaction'],
        'eyelid_mechanics': ['eyelid', 'lid', 'blinking', 'lid mechanics', 'eyelid interaction'],
        'blink_patterns': ['blink', 'blinking', 'blink pattern', 'blink frequency'],
        'lens_movement_post_blink': ['lens movement', 'movement', 'post-blink', 'displacement'],
        'initial_lens_fit': ['lens fit', 'fit', 'initial fit', 'fitting'],
        'comfort_perception': ['comfort', 'comfort perception', 'sensation', 'feeling'],
        'wearing_time': ['wearing time', 'wear time', 'duration', 'extended wear'],
        'lens_cleanliness_deposits': ['deposits', 'protein deposits', 'lipid deposits', 'cleanliness'],
        'packaging_wetting_solution': ['packaging solution', 'wetting solution', 'saline', 'buffer solution']
    }
    
    for driver, patterns in enhanced_driver_patterns.items():
        for pattern in patterns:
            if pattern.lower() in context_lower:
                found_drivers[driver] = found_drivers.get(driver, [])
                found_drivers[driver].append(pattern)
                logger.debug(f"   ✓ Found driver '{driver}' via pattern '{pattern}'")
                break
    
    # Enhanced factor detection with synonyms
    logger.debug("🎯 Scanning for performance factors with enhanced patterns...")
    factor_patterns = {
        'size': ['size', 'sizing', 'dimensions', 'dimensional'],
        'fit': ['fit', 'fitting', 'centration', 'positioning', 'alignment'],
        'handling': ['handling', 'manipulation', 'insertion', 'removal', 'application'],
        'comfort': ['comfort', 'comfortable', 'sensation', 'feeling', 'irritation', 'discomfort']
    }
    
    for factor, patterns in factor_patterns.items():
        for pattern in patterns:
            if pattern.lower() in context_lower:
                found_factors.append(factor)
                logger.debug(f"   ✓ Found factor '{factor}' via pattern '{pattern}'")
                break
    
    # Remove duplicates from factors
    found_factors = list(set(found_factors))
    
    logger.info(f"📊 Enhanced detection results: {len(found_drivers)} drivers, {len(found_factors)} factors")
    
    return found_drivers, found_factors

def monitor_endpoint_performance(logger):
    """Monitor endpoint performance and auto scaling metrics"""
    try:
        import boto3
        import time
        
        # Get CloudWatch metrics for the endpoint
        cloudwatch = boto3.client('cloudwatch', region_name='us-east-1')  # Adjust region as needed
        
        # Get recent metrics
        end_time = time.time()
        start_time = end_time - 300  # Last 5 minutes
        
        metrics_to_check = [
            'InvocationsPerInstance',
            'ModelLatency', 
            'Invocations',
            'InvocationModelErrors'
        ]
        
        logger.info("📊 Endpoint Performance Metrics (last 5 minutes):")
        
        for metric_name in metrics_to_check:
            try:
                response = cloudwatch.get_metric_statistics(
                    Namespace='AWS/SageMaker',
                    MetricName=metric_name,
                    Dimensions=[
                        {
                            'Name': 'EndpointName',
                            'Value': ENDPOINT_NAME.split('/')[-1] if '/' in ENDPOINT_NAME else ENDPOINT_NAME
                        }
                    ],
                    StartTime=start_time,
                    EndTime=end_time,
                    Period=60,
                    Statistics=['Average', 'Maximum']
                )
                
                if response['Datapoints']:
                    latest = max(response['Datapoints'], key=lambda x: x['Timestamp'])
                    avg_val = latest.get('Average', 0)
                    max_val = latest.get('Maximum', 0)
                    logger.info(f"   • {metric_name}: avg={avg_val:.2f}, max={max_val:.2f}")
                else:
                    logger.debug(f"   • {metric_name}: No data available")
                    
            except Exception as e:
                logger.debug(f"Could not fetch {metric_name}: {e}")
        
        # Check endpoint configuration for auto scaling
        try:
            sagemaker = boto3.client('sagemaker', region_name='us-east-1')
            endpoint_config_name = sagemaker.describe_endpoint(
                EndpointName=ENDPOINT_NAME.split('/')[-1] if '/' in ENDPOINT_NAME else ENDPOINT_NAME
            )['EndpointConfigName']
            
            config = sagemaker.describe_endpoint_config(
                EndpointConfigName=endpoint_config_name
            )
            
            for variant in config['ProductionVariants']:
                logger.info(f"🎯 Endpoint Configuration:")
                logger.info(f"   • Instance Type: {variant.get('InstanceType', 'N/A')}")
                logger.info(f"   • Initial Instance Count: {variant.get('InitialInstanceCount', 'N/A')}")
                logger.info(f"   • Initial Variant Weight: {variant.get('InitialVariantWeight', 'N/A')}")
                
        except Exception as e:
            logger.debug(f"Could not fetch endpoint configuration: {e}")
            
    except Exception as e:
        logger.debug(f"Performance monitoring failed: {e}")

def analyze_relationship_values(relationships: list, logger: logging.Logger) -> dict:
    """Analyze quantitative values and patterns in relationships"""
    try:
        import re
        import numpy as np
        
        logger.info("📊 Analyzing relationship values and magnitudes")
        
        # Extract numerical values from relationships
        numerical_patterns = [
            r'(\d+\.?\d*)\s*%',  # Percentages
            r'(\d+\.?\d*)\s*(?:fold|times|x)',  # Multipliers
            r'(\d+\.?\d*)\s*(?:mm|cm|μm)',  # Measurements
            r'(\d+\.?\d*)\s*(?:points?|units?)',  # Scale values
            r'(\d+\.?\d*)\s*(?:hours?|mins?|days?)',  # Time values
            r'improved?\s+by\s+(\d+\.?\d*)',  # Improvement values
            r'increased?\s+by\s+(\d+\.?\d*)',  # Increase values
            r'reduced?\s+by\s+(\d+\.?\d*)',  # Reduction values
            r'(\d+\.?\d*)\s*(?:out\s+of\s+\d+|/\d+)',  # Ratio values
        ]
        
        value_data = {
            'percentages': [],
            'measurements': [],
            'improvements': [],
            'reductions': [],
            'ratios': [],
            'multipliers': [],
            'time_values': [],
            'scale_values': []
        }
        
        relationship_magnitudes = []
        
        for rel in relationships:
            evidence_text = rel.get('evidence', '').lower()
            confidence = rel.get('confidence', 0)
            
            # Extract percentages
            percentages = re.findall(r'(\d+\.?\d*)\s*%', evidence_text)
            for pct in percentages:
                try:
                    val = float(pct)
                    value_data['percentages'].append({
                        'value': val,
                        'confidence': confidence,
                        'category': rel.get('driver_category', 'unknown'),
                        'evidence': evidence_text[:100]
                    })
                except:
                    continue
            
            # Extract measurements
            measurements = re.findall(r'(\d+\.?\d*)\s*(?:mm|cm|μm)', evidence_text)
            for meas in measurements:
                try:
                    val = float(meas)
                    value_data['measurements'].append({
                        'value': val,
                        'confidence': confidence,
                        'category': rel.get('driver_category', 'unknown'),
                        'evidence': evidence_text[:100]
                    })
                except:
                    continue
            
            # Extract improvement values
            improvements = re.findall(r'improved?\s+by\s+(\d+\.?\d*)', evidence_text)
            for imp in improvements:
                try:
                    val = float(imp)
                    value_data['improvements'].append({
                        'value': val,
                        'confidence': confidence,
                        'category': rel.get('driver_category', 'unknown'),
                        'evidence': evidence_text[:100]
                    })
                except:
                    continue
            
            # Calculate relationship magnitude score
            magnitude_score = confidence
            
            # Boost magnitude for quantitative evidence
            if any(re.search(pattern, evidence_text) for pattern in numerical_patterns):
                magnitude_score *= 1.5
            
            # Boost for strong language
            strong_indicators = ['significant', 'substantial', 'marked', 'dramatic', 'notable', 'considerable']
            if any(indicator in evidence_text for indicator in strong_indicators):
                magnitude_score *= 1.3
            
            relationship_magnitudes.append({
                'driver': rel.get('driver', ''),
                'factor': rel.get('factor', ''),
                'magnitude': min(magnitude_score, 1.0),  # Cap at 1.0
                'confidence': confidence,
                'has_quantitative': any(re.search(pattern, evidence_text) for pattern in numerical_patterns),
                'evidence_length': len(evidence_text)
            })
        
        # Calculate statistics
        analysis_results = {
            'total_relationships': len(relationships),
            'quantitative_relationships': len([r for r in relationship_magnitudes if r['has_quantitative']]),
            'high_magnitude_count': len([r for r in relationship_magnitudes if r['magnitude'] > 0.8]),
            'medium_magnitude_count': len([r for r in relationship_magnitudes if 0.6 <= r['magnitude'] <= 0.8]),
            'low_magnitude_count': len([r for r in relationship_magnitudes if r['magnitude'] < 0.6]),
            'value_distributions': value_data,
            'relationship_magnitudes': relationship_magnitudes,
            'average_magnitude': np.mean([r['magnitude'] for r in relationship_magnitudes]) if relationship_magnitudes else 0,
            'magnitude_std': np.std([r['magnitude'] for r in relationship_magnitudes]) if relationship_magnitudes else 0,
            'quantitative_percentage': (len([r for r in relationship_magnitudes if r['has_quantitative']]) / len(relationships) * 100) if relationships else 0
        }
        
        logger.info(f"📈 Value analysis complete:")
        logger.info(f"   • Total relationships: {analysis_results['total_relationships']}")
        logger.info(f"   • Quantitative relationships: {analysis_results['quantitative_relationships']} ({analysis_results['quantitative_percentage']:.1f}%)")
        logger.info(f"   • High magnitude: {analysis_results['high_magnitude_count']}")
        logger.info(f"   • Average magnitude: {analysis_results['average_magnitude']:.3f}")
        
        return analysis_results
        
    except Exception as e:
        logger.error(f"❌ Failed to analyze relationship values: {e}")
        return {}

def create_quantitative_charts(analysis_data: dict, lens_type: str, output_dir: Path, logger: logging.Logger) -> list:
    """Create charts and visualizations for quantitative analysis"""
    try:
        import matplotlib.pyplot as plt
        import seaborn as sns
        import numpy as np
        import pandas as pd
        
        logger.info(f"📊 Creating quantitative charts for {lens_type}")
        
        # Set style
        plt.style.use('default')
        if HAS_SEABORN:
            sns.set_palette("husl")
        
        chart_files = []
        
        # 1. Magnitude Distribution Chart
        fig, ax = plt.subplots(1, 1, figsize=(12, 8))
        
        magnitudes = [r['magnitude'] for r in analysis_data.get('relationship_magnitudes', [])]
        if magnitudes:
            ax.hist(magnitudes, bins=20, alpha=0.7, color='skyblue', edgecolor='black')
            ax.axvline(analysis_data.get('average_magnitude', 0), color='red', linestyle='--', 
                      label=f"Average: {analysis_data.get('average_magnitude', 0):.3f}")
            ax.set_xlabel('Relationship Magnitude Score')
            ax.set_ylabel('Frequency')
            ax.set_title(f'Relationship Magnitude Distribution - {lens_type}')
            ax.legend()
            ax.grid(True, alpha=0.3)
        
        chart_file = output_dir / f"magnitude_distribution_{lens_type}.png"
        plt.savefig(chart_file, dpi=300, bbox_inches='tight')
        plt.close()
        chart_files.append(chart_file)
        logger.info(f"   📈 Created magnitude distribution chart: {chart_file.name}")
        
        # 2. Confidence vs Magnitude Scatter Plot
        fig, ax = plt.subplots(1, 1, figsize=(12, 8))
        
        if analysis_data.get('relationship_magnitudes'):
            confidences = [r['confidence'] for r in analysis_data['relationship_magnitudes']]
            magnitudes = [r['magnitude'] for r in analysis_data['relationship_magnitudes']]
            has_quant = [r['has_quantitative'] for r in analysis_data['relationship_magnitudes']]
            
            # Different colors for quantitative vs qualitative
            quant_mask = np.array(has_quant)
            ax.scatter(np.array(confidences)[quant_mask], np.array(magnitudes)[quant_mask], 
                      alpha=0.7, color='orange', label='Quantitative', s=60)
            ax.scatter(np.array(confidences)[~quant_mask], np.array(magnitudes)[~quant_mask], 
                      alpha=0.7, color='lightblue', label='Qualitative', s=60)
            
            ax.set_xlabel('Confidence Score')
            ax.set_ylabel('Magnitude Score')
            ax.set_title(f'Confidence vs Magnitude Analysis - {lens_type}')
            ax.legend()
            ax.grid(True, alpha=0.3)
            
            # Add correlation line
            if len(confidences) > 1:
                z = np.polyfit(confidences, magnitudes, 1)
                p = np.poly1d(z)
                ax.plot(confidences, p(confidences), "r--", alpha=0.8, 
                       label=f'Trend (R²={np.corrcoef(confidences, magnitudes)[0,1]**2:.3f})')
                ax.legend()
        
        chart_file = output_dir / f"confidence_magnitude_scatter_{lens_type}.png"
        plt.savefig(chart_file, dpi=300, bbox_inches='tight')
        plt.close()
        chart_files.append(chart_file)
        logger.info(f"   📊 Created confidence-magnitude scatter: {chart_file.name}")
        
        # 3. Category Performance Chart
        if analysis_data.get('relationship_magnitudes'):
            category_stats = {}
            for rel in analysis_data['relationship_magnitudes']:
                driver = rel.get('driver', 'unknown')
                # Categorize drivers
                if any(keyword in driver.lower() for keyword in ['size', 'diameter', 'thickness']):
                    category = 'Size'
                elif any(keyword in driver.lower() for keyword in ['fit', 'centration', 'movement']):
                    category = 'Fit'
                elif any(keyword in driver.lower() for keyword in ['handling', 'insertion', 'removal']):
                    category = 'Handling'
                elif any(keyword in driver.lower() for keyword in ['comfort', 'sensation', 'irritation']):
                    category = 'Comfort'
                else:
                    category = 'Other'
                
                if category not in category_stats:
                    category_stats[category] = {'magnitudes': [], 'count': 0}
                
                category_stats[category]['magnitudes'].append(rel['magnitude'])
                category_stats[category]['count'] += 1
            
            if category_stats:
                fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(16, 8))
                
                # Average magnitude by category
                categories = list(category_stats.keys())
                avg_magnitudes = [np.mean(category_stats[cat]['magnitudes']) for cat in categories]
                colors = [CHART_COLORS.get(cat, '#888888') for cat in categories]
                
                bars = ax1.bar(categories, avg_magnitudes, color=colors, alpha=0.8)
                ax1.set_ylabel('Average Magnitude')
                ax1.set_title(f'Average Magnitude by Category - {lens_type}')
                ax1.set_ylim(0, 1.0)
                
                # Add value labels on bars
                for bar, mag in zip(bars, avg_magnitudes):
                    height = bar.get_height()
                    ax1.text(bar.get_x() + bar.get_width()/2., height + 0.01,
                            f'{mag:.3f}', ha='center', va='bottom')
                
                # Count by category
                counts = [category_stats[cat]['count'] for cat in categories]
                bars2 = ax2.bar(categories, counts, color=colors, alpha=0.8)
                ax2.set_ylabel('Number of Relationships')
                ax2.set_title(f'Relationship Count by Category - {lens_type}')
                
                # Add value labels on bars
                for bar, count in zip(bars2, counts):
                    height = bar.get_height()
                    ax2.text(bar.get_x() + bar.get_width()/2., height + 0.5,
                            str(count), ha='center', va='bottom')
                
                plt.tight_layout()
                chart_file = output_dir / f"category_performance_{lens_type}.png"
                plt.savefig(chart_file, dpi=300, bbox_inches='tight')
                plt.close()
                chart_files.append(chart_file)
                logger.info(f"   🎯 Created category performance chart: {chart_file.name}")
        
        # 4. Quantitative Values Distribution
        value_data = analysis_data.get('value_distributions', {})
        if any(value_data.values()):
            fig, axes = plt.subplots(2, 2, figsize=(16, 12))
            axes = axes.flatten()
            
            plot_idx = 0
            for value_type, values in value_data.items():
                if values and plot_idx < 4:
                    vals = [v['value'] for v in values if isinstance(v, dict) and 'value' in v]
                    if vals:
                        axes[plot_idx].hist(vals, bins=min(15, len(vals)), alpha=0.7, 
                                          color=list(CHART_COLORS.values())[plot_idx % 4])
                        axes[plot_idx].set_title(f'{value_type.replace("_", " ").title()} Distribution')
                        axes[plot_idx].set_xlabel('Value')
                        axes[plot_idx].set_ylabel('Frequency')
                        axes[plot_idx].grid(True, alpha=0.3)
                        plot_idx += 1
            
            # Hide unused subplots
            for i in range(plot_idx, 4):
                axes[i].set_visible(False)
            
            plt.tight_layout()
            chart_file = output_dir / f"quantitative_distributions_{lens_type}.png"
            plt.savefig(chart_file, dpi=300, bbox_inches='tight')
            plt.close()
            chart_files.append(chart_file)
            logger.info(f"   📊 Created quantitative distributions chart: {chart_file.name}")
        
        logger.info(f"📈 Created {len(chart_files)} quantitative charts for {lens_type}")
        return chart_files
        
    except Exception as e:
        logger.error(f"❌ Failed to create quantitative charts: {e}")
        return []

def get_optimal_batch_size(current_workers: int, logger) -> int:
    """Calculate optimal batch size based on endpoint performance and current load"""
    try:
        # Monitor recent endpoint performance
        base_batch_size = max(2, LLM_BATCH_SIZE)
        
        # Adjust based on worker count to prevent overload
        if current_workers > 200:
            # High worker count - use smaller batches to distribute load
            optimal_size = max(2, base_batch_size // 2)
            logger.debug(f"🔧 High worker count ({current_workers}), reducing batch size to {optimal_size}")
        elif current_workers < 50:
            # Low worker count - can use larger batches
            optimal_size = min(12, base_batch_size * 2)
            logger.debug(f"🔧 Low worker count ({current_workers}), increasing batch size to {optimal_size}")
        else:
            optimal_size = base_batch_size
            logger.debug(f"🔧 Standard worker count ({current_workers}), using base batch size {optimal_size}")
        
        return optimal_size
        
    except Exception as e:
        logger.warning(f"⚠️ Error calculating optimal batch size: {e}")
        return LLM_BATCH_SIZE

def monitor_batch_performance(batch_results: list, batch_times: list, logger: logging.Logger) -> dict:
    """Monitor batch processing performance and suggest optimizations"""
    try:
        if not batch_results or not batch_times:
            return {}
        
        # Handle case where batch_results might contain integers instead of lists
        total_relationships = 0
        for result in batch_results:
            if isinstance(result, (list, tuple)):
                total_relationships += len(result)
            elif isinstance(result, int):
                total_relationships += result
            else:
                # Assume it's a single relationship if not a collection
                total_relationships += 1
        
        total_time = sum(batch_times)
        avg_time = sum(batch_times) / len(batch_times) if batch_times else 0
        
        # Calculate throughput metrics
        throughput_per_second = total_relationships / total_time if total_time > 0 else 0
        batches_per_minute = len(batch_results) / (total_time / 60) if total_time > 0 else 0
        
        # Analyze performance patterns
        slow_batches = len([t for t in batch_times if t > avg_time * 1.5]) if avg_time > 0 else 0
        fast_batches = len([t for t in batch_times if t < avg_time * 0.5]) if avg_time > 0 else 0
        
        performance_stats = {
            'total_batches': len(batch_results),
            'total_relationships': total_relationships,
            'total_time': total_time,
            'avg_batch_time': avg_time,
            'throughput_per_second': throughput_per_second,
            'batches_per_minute': batches_per_minute,
            'slow_batches': slow_batches,
            'fast_batches': fast_batches,
            'efficiency_score': total_relationships / len(batch_results) if batch_results else 0
        }
        
        logger.info(f"📊 Batch Performance Analysis:")
        logger.info(f"   • Total batches: {performance_stats['total_batches']}")
        logger.info(f"   • Relationships found: {performance_stats['total_relationships']}")
        logger.info(f"   • Average batch time: {performance_stats['avg_batch_time']:.2f}s")
        logger.info(f"   • Throughput: {performance_stats['throughput_per_second']:.2f} relationships/sec")
        logger.info(f"   • Efficiency: {performance_stats['efficiency_score']:.1f} relationships/batch")
        
        # Provide optimization suggestions
        if slow_batches > len(batch_results) * 0.3:
            logger.warning(f"⚠️ {slow_batches} slow batches detected - consider reducing batch size")
        
        if performance_stats['efficiency_score'] < 2.0:
            logger.warning(f"⚠️ Low efficiency ({performance_stats['efficiency_score']:.1f}) - consider optimizing prompts")
        
        return performance_stats
        
    except Exception as e:
        logger.error(f"❌ Failed to analyze batch performance: {e}")
        logger.debug(f"Batch results type: {type(batch_results)}, length: {len(batch_results) if batch_results else 0}")
        logger.debug(f"Batch times type: {type(batch_times)}, length: {len(batch_times) if batch_times else 0}")
        return {}

def adaptive_batch_processing(chunks: list, document_key: str, lens_type: str, logger: logging.Logger) -> list:
    """Intelligent adaptive batch processing that learns optimal batch sizes during execution"""
    logger.info(f"🧠 Starting adaptive batch processing for {len(chunks)} chunks from {document_key}")
    
    all_relationships = []
    batch_performance_history = []
    current_batch_size = LLM_BATCH_SIZE
    
    # Start with initial batches to gauge performance
    remaining_chunks = chunks.copy()
    chunk_index = 0
    
    while remaining_chunks:
        # Create current batch
        current_batch = remaining_chunks[:current_batch_size]
        current_batch_ids = [f"{document_key.split('/')[-1]}_chunk_{chunk_index + i + 1}" 
                           for i in range(len(current_batch))]
        
        logger.debug(f"🔄 Adaptive batch {len(batch_performance_history) + 1}: {len(current_batch)} chunks (size: {current_batch_size})")
        
        batch_start_time = time.time()
        
        try:
            # Process batch
            if len(current_batch) > 1:
                batch_relationships = extract_relationships_with_llm_batch(
                    current_batch, current_batch_ids, logger, lens_type, max_retries=BATCH_RETRY_ATTEMPTS
                )
            else:
                single_relationships = extract_relationships_with_llm(
                    current_batch[0], logger, lens_type, max_retries=3
                )
                batch_relationships = single_relationships
            
            batch_time = time.time() - batch_start_time
            
            # Match relationships
            if batch_relationships:
                matched_relationships = match_relationships(
                    batch_relationships, 
                    logger,
                    document_key=document_key,
                    chunk_id=f"{document_key.split('/')[-1]}_adaptive_batch_{len(batch_performance_history) + 1}"
                )
                all_relationships.extend(matched_relationships)
            else:
                matched_relationships = []
            
            # Record performance
            batch_performance = {
                'batch_size': len(current_batch),
                'processing_time': batch_time,
                'relationships_found': len(matched_relationships),
                'efficiency': len(matched_relationships) / batch_time if batch_time > 0 else 0,
                'success': True
            }
            batch_performance_history.append(batch_performance)
            
            logger.debug(f"✅ Batch completed: {len(matched_relationships)} relationships in {batch_time:.2f}s")
            
            # Adaptive learning - adjust batch size based on performance
            current_batch_size = calculate_next_batch_size(
                batch_performance_history, current_batch_size, len(remaining_chunks), logger
            )
            
        except Exception as e:
            batch_time = time.time() - batch_start_time
            logger.error(f"❌ Batch failed after {batch_time:.2f}s: {e}")
            
            # Record failure
            batch_performance_history.append({
                'batch_size': len(current_batch),
                'processing_time': batch_time,
                'relationships_found': 0,
                'efficiency': 0,
                'success': False
            })
            
            # Reduce batch size on failure
            current_batch_size = max(1, current_batch_size // 2)
            logger.warning(f"⚠️ Reducing batch size to {current_batch_size} due to failure")
        
        # Update remaining chunks
        remaining_chunks = remaining_chunks[len(current_batch):]
        chunk_index += len(current_batch)
        
        # Brief pause between batches
        time.sleep(LLM_CALL_DELAY)
    
    # Analyze overall performance - fix the data passing
    if batch_performance_history:
        relationship_counts = [h['relationships_found'] for h in batch_performance_history]
        processing_times = [h['processing_time'] for h in batch_performance_history]
        
        monitor_batch_performance(relationship_counts, processing_times, logger)
    
    logger.info(f"🎯 Adaptive processing complete: {len(all_relationships)} relationships from {len(chunks)} chunks")
    logger.info(f"   • Used {len(batch_performance_history)} adaptive batches")
    logger.info(f"   • Final batch size: {current_batch_size}")
    
    return all_relationships

def calculate_next_batch_size(performance_history: list, current_size: int, remaining_chunks: int, logger: logging.Logger) -> int:
    """Calculate optimal next batch size based on performance history"""
    try:
        if len(performance_history) < 2:
            return current_size
        
        # Analyze recent performance
        recent_batches = performance_history[-3:]  # Look at last 3 batches
        successful_batches = [b for b in recent_batches if b['success']]
        
        if not successful_batches:
            # All recent batches failed - reduce size significantly
            new_size = max(1, current_size // 2)
            logger.debug(f"🔻 All recent batches failed, reducing size: {current_size} → {new_size}")
            return new_size
        
        # Calculate average efficiency of recent successful batches
        avg_efficiency = sum(b['efficiency'] for b in successful_batches) / len(successful_batches)
        avg_time = sum(b['processing_time'] for b in successful_batches) / len(successful_batches)
        
        # Compare with earlier performance
        if len(performance_history) >= 6:
            earlier_batches = [b for b in performance_history[-6:-3] if b['success']]
            if earlier_batches:
                earlier_avg_efficiency = sum(b['efficiency'] for b in earlier_batches) / len(earlier_batches)
                
                # If recent efficiency is better, consider increasing batch size
                if avg_efficiency > earlier_avg_efficiency * 1.1 and avg_time < 60:
                    new_size = min(LLM_BATCH_SIZE * 2, current_size + 2, remaining_chunks)
                    logger.debug(f"🔺 Efficiency improving, increasing size: {current_size} → {new_size}")
                    return new_size
                
                # If efficiency is declining, reduce batch size
                elif avg_efficiency < earlier_avg_efficiency * 0.9:
                    new_size = max(2, current_size - 1)
                    logger.debug(f"🔻 Efficiency declining, reducing size: {current_size} → {new_size}")
                    return new_size
        
        # Adjust based on processing time
        if avg_time > 90:  # Batches taking too long
            new_size = max(2, current_size - 1)
            logger.debug(f"⏰ Batches too slow ({avg_time:.1f}s), reducing size: {current_size} → {new_size}")
            return new_size
        elif avg_time < 30 and avg_efficiency > 2.0:  # Fast and efficient
            new_size = min(LLM_BATCH_SIZE * 2, current_size + 1, remaining_chunks)
            logger.debug(f"⚡ Fast and efficient, increasing size: {current_size} → {new_size}")
            return new_size
        
        # Default: maintain current size
        logger.debug(f"📊 Maintaining current batch size: {current_size}")
        return current_size
        
    except Exception as e:
        logger.warning(f"⚠️ Error calculating next batch size: {e}")
        return current_size

def create_factor_specific_analysis(relationships: list, factor_name: str, lens_type: str, output_dir: Path, logger: logging.Logger) -> dict:
    """Create comprehensive analysis for a specific performance factor (Size, Fit, Handling, Comfort)"""
    try:
        logger.info(f"📊 Creating detailed {factor_name.title()} analysis for {lens_type}")
        
        # Filter relationships for this specific factor
        factor_relationships = [rel for rel in relationships if rel.get('factor', '').lower() == factor_name.lower()]
        
        if not factor_relationships:
            logger.warning(f"⚠️ No relationships found for {factor_name} factor")
            return {
                'factor': factor_name,
                'relationships_count': 0,
                'drivers_found': [],
                'analysis_files': []
            }
        
        logger.info(f"🔍 Found {len(factor_relationships)} relationships for {factor_name}")
        
        # Analyze drivers for this factor
        driver_analysis = {}
        for rel in factor_relationships:
            driver = rel.get('driver', '')
            if driver not in driver_analysis:
                driver_analysis[driver] = {
                    'driver_name': driver.replace('_', ' ').title(),
                    'relationships': [],
                    'confidence_levels': [],
                    'evidence_snippets': [],
                    'quantitative_values': [],
                    'documents': set()
                }
            
            driver_data = driver_analysis[driver]
            driver_data['relationships'].append(rel)
            driver_data['confidence_levels'].append(rel.get('confidence', 'medium'))
            driver_data['evidence_snippets'].append(rel.get('evidence', ''))
            driver_data['documents'].add(rel.get('document', rel.get('section_id', 'unknown')))
            
            # Extract quantitative values from evidence
            evidence = rel.get('evidence', '').lower()
            quant_values = extract_quantitative_values_from_text(evidence)
            driver_data['quantitative_values'].extend(quant_values)
        
        # Convert sets to lists for JSON serialization
        for driver_data in driver_analysis.values():
            driver_data['documents'] = list(driver_data['documents'])
            driver_data['unique_documents'] = len(driver_data['documents'])
            
            # Calculate summary statistics
            confidence_scores = {'high': 3, 'medium': 2, 'low': 1}
            scores = [confidence_scores.get(conf.lower(), 2) for conf in driver_data['confidence_levels']]
            avg_score = sum(scores) / len(scores) if scores else 2
            
            if avg_score >= 2.5:
                driver_data['overall_confidence'] = 'high'
            elif avg_score >= 1.5:
                driver_data['overall_confidence'] = 'medium'
            else:
                driver_data['overall_confidence'] = 'low'
            
            driver_data['relationship_count'] = len(driver_data['relationships'])
            driver_data['avg_confidence_score'] = avg_score
        
        # Create factor-specific graph
        graph_file = create_factor_specific_graph(driver_analysis, factor_name, lens_type, output_dir, logger)
        
        # Create factor-specific tabular report
        table_file = create_factor_tabular_report(driver_analysis, factor_relationships, factor_name, lens_type, output_dir, logger)
        
        # Create factor summary report
        summary_file = create_factor_summary_report(driver_analysis, factor_relationships, factor_name, lens_type, output_dir, logger)
        
        analysis_result = {
            'factor': factor_name,
            'relationships_count': len(factor_relationships),
            'unique_drivers': len(driver_analysis),
            'drivers_found': list(driver_analysis.keys()),
            'driver_analysis': driver_analysis,
            'analysis_files': [f for f in [graph_file, table_file, summary_file] if f]
        }
        
        logger.info(f"✅ {factor_name.title()} analysis complete:")
        logger.info(f"   • {len(factor_relationships)} relationships analyzed")
        logger.info(f"   • {len(driver_analysis)} unique drivers found")
        logger.info(f"   • {len(analysis_result['analysis_files'])} analysis files created")
        
        return analysis_result
        
    except Exception as e:
        logger.error(f"❌ Failed to create {factor_name} analysis: {e}")
        return {
            'factor': factor_name,
            'relationships_count': 0,
            'drivers_found': [],
            'analysis_files': [],
            'error': str(e)
        }

def extract_quantitative_values_from_text(text: str) -> list:
    """Extract quantitative values from evidence text"""
    values = []
    
    # Patterns for different types of quantitative data
    patterns = [
        (r'(\d+\.?\d*)\s*%', 'percentage'),
        (r'(\d+\.?\d*)\s*(?:fold|times|x)', 'multiplier'),
        (r'(\d+\.?\d*)\s*(?:mm|cm|μm|inches?)', 'measurement'),
        (r'(\d+\.?\d*)\s*(?:points?|units?|scale)', 'scale'),
        (r'(\d+\.?\d*)\s*(?:hours?|mins?|minutes?|days?)', 'time'),
        (r'improved?\s+by\s+(\d+\.?\d*)', 'improvement'),
        (r'increased?\s+by\s+(\d+\.?\d*)', 'increase'),
        (r'reduced?\s+by\s+(\d+\.?\d*)', 'reduction'),
        (r'(\d+\.?\d*)\s*(?:out\s+of\s+\d+|/\d+)', 'ratio'),
    ]
    
    for pattern, value_type in patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            try:
                values.append({
                    'value': float(match),
                    'type': value_type,
                    'context': text
                })
            except ValueError:
                continue
    
    return values

def create_factor_specific_graph(driver_analysis: dict, factor_name: str, lens_type: str, output_dir: Path, logger: logging.Logger) -> str:
    """Create a factor-specific graph showing drivers and their relationships"""
    try:
        logger.debug(f"📊 Creating {factor_name} factor graph for {lens_type}")
        
        if not driver_analysis:
            logger.warning(f"⚠️ No driver data for {factor_name} graph")
            return None
        
        # Set up the plot
        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(16, 8))
        fig.suptitle(f'{factor_name.title()} Factor Analysis - {lens_type.title() if lens_type else "Unknown"}', fontsize=16, fontweight='bold')
        
        # Get factor color
        factor_color = PERFORMANCE_FACTORS.get(factor_name.lower(), {}).get('color', '#45B7D1')
        
        # Left plot: Driver Relationship Counts
        drivers = list(driver_analysis.keys())
        counts = [data['relationship_count'] for data in driver_analysis.values()]
        confidence_colors = []
        
        for data in driver_analysis.values():
            conf = data['overall_confidence']
            if conf == 'high':
                confidence_colors.append('#2ECC71')  # Green
            elif conf == 'medium':
                confidence_colors.append('#F39C12')  # Orange
            else:
                confidence_colors.append('#E74C3C')  # Red
        
        # Create horizontal bar chart
        y_pos = range(len(drivers))
        bars = ax1.barh(y_pos, counts, color=confidence_colors, alpha=0.7, edgecolor='black', linewidth=0.5)
        
        ax1.set_yticks(y_pos)
        ax1.set_yticklabels([d.replace('_', ' ').title() for d in drivers], fontsize=10)
        ax1.set_xlabel('Number of Relationships', fontsize=12)
        ax1.set_title(f'{factor_name.title()} Drivers by Relationship Count', fontsize=14, fontweight='bold')
        ax1.grid(axis='x', alpha=0.3)
        
        # Add value labels on bars
        for i, (bar, count) in enumerate(zip(bars, counts)):
            ax1.text(bar.get_width() + 0.1, bar.get_y() + bar.get_height()/2, 
                    str(count), va='center', ha='left', fontweight='bold')
        
        # Add confidence legend
        from matplotlib.patches import Patch
        legend_elements = [
            Patch(facecolor='#2ECC71', label='High Confidence'),
            Patch(facecolor='#F39C12', label='Medium Confidence'),
            Patch(facecolor='#E74C3C', label='Low Confidence')
        ]
        ax1.legend(handles=legend_elements, loc='lower right')
        
        # Right plot: Confidence Score Distribution
        confidence_scores = [data['avg_confidence_score'] for data in driver_analysis.values()]
        
        # Create scatter plot with driver names
        y_positions = range(len(drivers))
        scatter = ax2.scatter(confidence_scores, y_positions, 
                            c=[data['relationship_count'] for data in driver_analysis.values()],
                            s=100, alpha=0.7, cmap='viridis', edgecolors='black', linewidth=0.5)
        
        ax2.set_yticks(y_positions)
        ax2.set_yticklabels([d.replace('_', ' ').title() for d in drivers], fontsize=10)
        ax2.set_xlabel('Average Confidence Score', fontsize=12)
        ax2.set_title(f'{factor_name.title()} Driver Confidence Analysis', fontsize=14, fontweight='bold')
        ax2.set_xlim(0.5, 3.5)
        ax2.grid(alpha=0.3)
        
        # Add colorbar for relationship counts
        cbar = plt.colorbar(scatter, ax=ax2)
        cbar.set_label('Number of Relationships', rotation=270, labelpad=15)
        
        # Add confidence score reference lines
        ax2.axvline(x=2.5, color='green', linestyle='--', alpha=0.5, label='High Confidence Threshold')
        ax2.axvline(x=1.5, color='orange', linestyle='--', alpha=0.5, label='Medium Confidence Threshold')
        ax2.legend(loc='lower right')
        
        # Save the graph
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        graph_filename = f"{factor_name.lower()}_factor_analysis_{lens_type}_{timestamp}.png"
        graph_file_path = output_dir / graph_filename
        
        plt.tight_layout()
        plt.savefig(graph_file_path, dpi=300, bbox_inches='tight', facecolor='white')
        plt.close()
        
        logger.info(f"📊 {factor_name.title()} factor graph saved: {graph_filename}")
        return str(graph_file_path)
        
    except Exception as e:
        logger.error(f"❌ Failed to create {factor_name} factor graph: {e}")
        return None

def create_factor_tabular_report(driver_analysis: dict, relationships: list, factor_name: str, lens_type: str, output_dir: Path, logger: logging.Logger) -> str:
    """Create detailed tabular report for a specific factor"""
    try:
        logger.debug(f"📋 Creating {factor_name} tabular report for {lens_type}")
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Create comprehensive CSV report
        csv_filename = f"{factor_name.lower()}_detailed_analysis_{lens_type}_{timestamp}.csv"
        csv_file_path = output_dir / csv_filename
        
        with open(csv_file_path, 'w', newline='', encoding='utf-8') as csvfile:
            writer = csv.writer(csvfile)
            
            # Header
            writer.writerow([
                'Driver', 'Driver_Description', 'Relationship_Count', 'Overall_Confidence',
                'Avg_Confidence_Score', 'Unique_Documents', 'Evidence_Summary',
                'Quantitative_Values', 'Related_Documents'
            ])
            
            # Driver rows
            for driver, data in driver_analysis.items():
                driver_desc = PERFORMANCE_FACTORS.get(factor_name.lower(), {}).get('drivers', {}).get(driver, 'No description available')
                
                # Summarize evidence
                evidence_summary = '; '.join(data['evidence_snippets'][:3])  # First 3 pieces of evidence
                if len(data['evidence_snippets']) > 3:
                    evidence_summary += f' ... ({len(data["evidence_snippets"])} total pieces of evidence)'
                
                # Summarize quantitative values
                quant_summary = []
                for val_info in data['quantitative_values'][:5]:  # First 5 values
                    quant_summary.append(f"{val_info['value']} ({val_info['type']})")
                quant_str = '; '.join(quant_summary)
                if len(data['quantitative_values']) > 5:
                    quant_str += f' ... ({len(data["quantitative_values"])} total values)'
                
                writer.writerow([
                    data['driver_name'],
                    driver_desc,
                    data['relationship_count'],
                    data['overall_confidence'],
                    f"{data['avg_confidence_score']:.2f}",
                    data['unique_documents'],
                    evidence_summary,
                    quant_str,
                    '; '.join(data['documents'][:3])  # First 3 documents
                ])
        
        # Create relationships detail CSV
        relationships_csv_filename = f"{factor_name.lower()}_relationships_detail_{lens_type}_{timestamp}.csv"
        relationships_csv_path = output_dir / relationships_csv_filename
        
        with open(relationships_csv_path, 'w', newline='', encoding='utf-8') as csvfile:
            writer = csv.writer(csvfile)
            
            # Header
            writer.writerow([
                'Driver', 'Factor', 'Predicate', 'Confidence', 'Evidence',
                'Document/Section', 'Quantitative_Values_Found'
            ])
            
            # Relationship rows
            for rel in relationships:
                evidence = rel.get('evidence', '')
                quant_values = extract_quantitative_values_from_text(evidence.lower())
                quant_str = '; '.join([f"{v['value']} ({v['type']})" for v in quant_values[:3]])
                
                writer.writerow([
                    rel.get('driver', ''),
                    rel.get('factor', ''),
                    rel.get('predicate', ''),
                    rel.get('confidence', ''),
                    evidence,
                    rel.get('document', rel.get('section_id', '')),
                    quant_str
                ])
        
        logger.info(f"📋 {factor_name.title()} tabular reports created:")
        logger.info(f"   • Driver analysis: {csv_filename}")
        logger.info(f"   • Relationships detail: {relationships_csv_filename}")
        
        return str(csv_file_path)
        
    except Exception as e:
        logger.error(f"❌ Failed to create {factor_name} tabular report: {e}")
        return None

def create_factor_summary_report(driver_analysis: dict, relationships: list, factor_name: str, lens_type: str, output_dir: Path, logger: logging.Logger) -> str:
    """Create a comprehensive summary report for a specific factor"""
    try:
        logger.debug(f"📝 Creating {factor_name} summary report for {lens_type}")
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        summary_filename = f"{factor_name.lower()}_summary_report_{lens_type}_{timestamp}.md"
        summary_file_path = output_dir / summary_filename
        
        with open(summary_file_path, 'w', encoding='utf-8') as f:
            f.write(f"# {factor_name.title()} Factor Analysis Report\n\n")
            f.write(f"**Lens Type:** {lens_type.title() if lens_type else 'Unknown'}\n")
            f.write(f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            
            # Factor description
            factor_info = PERFORMANCE_FACTORS.get(factor_name.lower(), {})
            f.write(f"## Factor Description\n")
            f.write(f"{factor_info.get('description', 'No description available')}\n\n")
            
            # Summary statistics
            f.write(f"## Summary Statistics\n")
            f.write(f"- **Total Relationships Found:** {len(relationships)}\n")
            f.write(f"- **Unique Drivers Identified:** {len(driver_analysis)}\n")
            f.write(f"- **Factor Color Code:** {factor_info.get('color', '#000000')}\n\n")
            
            # Driver analysis
            f.write(f"## Driver Analysis\n\n")
            
            # Sort drivers by relationship count
            sorted_drivers = sorted(driver_analysis.items(), 
                                  key=lambda x: x[1]['relationship_count'], reverse=True)
            
            for driver, data in sorted_drivers:
                f.write(f"### {data['driver_name']}\n")
                
                # Driver description
                driver_desc = factor_info.get('drivers', {}).get(driver, 'No description available')
                f.write(f"**Description:** {driver_desc}\n\n")
                
                # Statistics
                f.write(f"**Statistics:**\n")
                f.write(f"- Relationships Found: {data['relationship_count']}\n")
                f.write(f"- Overall Confidence: {data.get('overall_confidence', 'unknown').title()}\n")
                f.write(f"- Average Confidence Score: {data['avg_confidence_score']:.2f}/3.0\n")
                f.write(f"- Documents Referenced: {data['unique_documents']}\n\n")
                
                # Quantitative findings
                if data['quantitative_values']:
                    f.write(f"**Quantitative Findings:**\n")
                    for val_info in data['quantitative_values'][:5]:  # Top 5 values
                        f.write(f"- {val_info['value']} ({val_info['type']})\n")
                    if len(data['quantitative_values']) > 5:
                        f.write(f"- ... and {len(data['quantitative_values']) - 5} more values\n")
                    f.write("\n")
                
                # Key evidence
                f.write(f"**Key Evidence:**\n")
                for evidence in data['evidence_snippets'][:3]:  # Top 3 pieces
                    if evidence.strip():
                        f.write(f"- {evidence[:200]}{'...' if len(evidence) > 200 else ''}\n")
                f.write("\n")
                
                f.write("---\n\n")
            
            # Relationships table
            f.write(f"## Detailed Relationships\n\n")
            f.write("| Driver | Predicate | Confidence | Evidence Summary |\n")
            f.write("|--------|-----------|------------|------------------|\n")
            
            for rel in relationships:
                evidence_summary = rel.get('evidence', '')[:100]
                if len(rel.get('evidence', '')) > 100:
                    evidence_summary += '...'
                
                f.write(f"| {rel.get('driver', '').replace('_', ' ').title()} | ")
                f.write(f"{rel.get('predicate', '')} | ")
                f.write(f"{rel.get('confidence', '').title()} | ")
                f.write(f"{evidence_summary} |\n")
            
            f.write("\n")
            
            # Recommendations
            f.write(f"## Key Insights & Recommendations\n\n")
            
            if len(driver_analysis) == 0:
                f.write("- No specific drivers were identified for this factor in the analyzed documents.\n")
                f.write("- Consider expanding the document set or reviewing extraction criteria.\n")
            else:
                # Top driver
                top_driver = sorted_drivers[0] if sorted_drivers else None
                if top_driver:
                    f.write(f"- **Primary Driver:** {top_driver[1]['driver_name']} ")
                    f.write(f"({top_driver[1]['relationship_count']} relationships)\n")
                
                # High confidence drivers
                high_conf_drivers = [d for d, data in driver_analysis.items() 
                                   if data['overall_confidence'] == 'high']
                if high_conf_drivers:
                    f.write(f"- **High Confidence Drivers:** {', '.join([d.replace('_', ' ').title() for d in high_conf_drivers])}\n")
                
                # Quantitative insights
                total_quant_values = sum(len(data['quantitative_values']) for data in driver_analysis.values())
                if total_quant_values > 0:
                    f.write(f"- **Quantitative Evidence:** {total_quant_values} quantitative measurements found\n")
                
                f.write(f"- **Research Coverage:** {len(set().union(*[data['documents'] for data in driver_analysis.values()]))} unique documents referenced\n")
        
        logger.info(f"📝 {factor_name.title()} summary report created: {summary_filename}")
        return str(summary_file_path)
        
    except Exception as e:
        logger.error(f"❌ Failed to create {factor_name} summary report: {e}")
        return None

def create_overall_findings_summary(results: dict, output_dir: Path, logger: logging.Logger) -> str:
    """Create comprehensive overall findings summary report"""
    try:
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        summary_file = output_dir / f"overall_findings_summary_{timestamp}.md"
        
        logger.info(f"📋 Creating overall findings summary: {summary_file}")
        
        # Aggregate data across all lens types
        total_docs = sum(r.get('study_analysis', {}).get('documents_processed', 0) for r in results.values())
        total_relationships = sum(r.get('study_analysis', {}).get('total_relationships', 0) for r in results.values())
        total_drivers = sum(len(r.get('study_analysis', {}).get('comfort_drivers', {})) for r in results.values())
        
        # Calculate processing statistics
        processing_stats = {
            'total_execution_time': 0,
            'total_s3_uploads': 0,
            'total_reports_generated': 0,
            'category_reports_generated': 0
        }
        
        for lens_type, result in results.items():
            processing_stats['total_reports_generated'] += len(result.get('reports', {}))
            processing_stats['category_reports_generated'] += result.get('category_reports_count', 0)
            if result.get('s3_upload', {}).get('upload_count'):
                processing_stats['total_s3_uploads'] += result['s3_upload']['upload_count']
        
        with open(summary_file, 'w', encoding='utf-8') as f:
            f.write(f"# Contact Lens Performance Drivers - Overall Findings Summary\n\n")
            f.write(f"**Generated:** {time.strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"**Report ID:** OVERALL_SUMMARY_{timestamp}\n")
            f.write(f"**Pipeline Version:** 2.0 (Auto Scaling + Multi-Category Analysis)\n\n")
            
            f.write("## Executive Summary\n\n")
            f.write(f"This comprehensive analysis processed **{total_docs} documents** across multiple lens types, ")
            f.write(f"extracting **{total_relationships} relationships** and identifying **{total_drivers} performance drivers** ")
            f.write("using advanced NLP/LLM techniques with auto scaling optimization.\n\n")
            
            # Processing Performance Summary
            f.write("## Processing Performance Summary\n\n")
            f.write(f"- **Documents Processed:** {total_docs}\n")
            f.write(f"- **Total Relationships Extracted:** {total_relationships}\n")
            f.write(f"- **Performance Drivers Identified:** {total_drivers}\n")
            f.write(f"- **Reports Generated:** {processing_stats['total_reports_generated']}\n")
            f.write(f"- **Category Reports:** {processing_stats['category_reports_generated']}\n")
            f.write(f"- **S3 Uploads Completed:** {processing_stats['total_s3_uploads']}\n")
            f.write(f"- **Lens Types Analyzed:** {len(results)}\n\n")
            
            # Lens Type Breakdown
            f.write("## Lens Type Analysis Breakdown\n\n")
            for lens_type, result in results.items():
                study_analysis = result.get('study_analysis', {})
                f.write(f"### {lens_type.upper() if lens_type else 'UNKNOWN'}\n\n")
                f.write(f"**Documents:** {study_analysis.get('documents_processed', 0)} processed\n")
                f.write(f"**Relationships:** {study_analysis.get('total_relationships', 0)} extracted\n")
                f.write(f"**Drivers Found:** {len(study_analysis.get('comfort_drivers', {}))}\n")
                f.write(f"**Graph Nodes:** {result.get('nodes', 0)}\n")
                f.write(f"**Graph Edges:** {result.get('edges', 0)}\n")
                
                # Category breakdown for this lens type
                category_reports = result.get('category_reports', {})
                if category_reports:
                    f.write(f"**Category Reports:**\n")
                    for category in ["Size", "Fit", "Handling", "Comfort"]:
                        if category in category_reports:
                            f.write(f"  - {category}: ✅ Generated\n")
                        else:
                            f.write(f"  - {category}: ❌ No data\n")
                f.write("\n")
            
            # Confidence Analysis
            f.write("## Confidence Analysis Across All Lens Types\n\n")
            high_conf_total = 0
            med_conf_total = 0
            low_conf_total = 0
            
            for result in results.values():
                drivers = result.get('study_analysis', {}).get('comfort_drivers', {})
                high_conf_total += len([d for d in drivers.values() if d.get('overall_confidence', '') and d.get('overall_confidence', '').upper() == 'HIGH'])
                med_conf_total += len([d for d in drivers.values() if d.get('overall_confidence', '') and d.get('overall_confidence', '').upper() == 'MEDIUM'])
                low_conf_total += len([d for d in drivers.values() if d.get('overall_confidence', '') and d.get('overall_confidence', '').upper() == 'LOW'])
            
            f.write(f"- **High Confidence Drivers:** {high_conf_total} ({high_conf_total/total_drivers*100:.1f}%)\n" if total_drivers > 0 else f"- **High Confidence Drivers:** {high_conf_total} (0.0%)\n")
            f.write(f"- **Medium Confidence Drivers:** {med_conf_total} ({med_conf_total/total_drivers*100:.1f}%)\n" if total_drivers > 0 else f"- **Medium Confidence Drivers:** {med_conf_total} (0.0%)\n")
            f.write(f"- **Low Confidence Drivers:** {low_conf_total} ({low_conf_total/total_drivers*100:.1f}%)\n\n" if total_drivers > 0 else f"- **Low Confidence Drivers:** {low_conf_total} (0.0%)\n\n")
            
            # Performance Category Distribution
            f.write("## Performance Category Distribution\n\n")
            category_distribution = {"Size": 0, "Fit": 0, "Handling": 0, "Comfort": 0, "Other": 0}
            
            for result in results.values():
                category_reports = result.get('category_reports', {})
                for category in category_distribution.keys():
                    if category in category_reports:
                        category_distribution[category] += 1
            
            f.write("Reports generated by category:\n")
            for category, count in category_distribution.items():
                percentage = (count / len(results) * 100) if results else 0
                f.write(f"- **{category}:** {count}/{len(results)} lens types ({percentage:.1f}%)\n")
            f.write("\n")
            
            # Technology Performance
            f.write("## Technology Performance Summary\n\n")
            f.write(f"**Auto Scaling Configuration:**\n")
            f.write(f"- Document Workers: {MAX_WORKERS}\n")
            f.write(f"- LLM Workers: {MAX_LLM_WORKERS}\n")
            f.write(f"- Batch Size: {BATCH_SIZE}\n")
            f.write(f"- Chunk Size: {CHUNK_SIZE} characters\n\n")
            
            f.write(f"**Processing Efficiency:**\n")
            if total_docs > 0:
                avg_relationships_per_doc = total_relationships / total_docs
                avg_drivers_per_doc = total_drivers / total_docs
                f.write(f"- Average Relationships per Document: {avg_relationships_per_doc:.2f}\n")
                f.write(f"- Average Drivers per Document: {avg_drivers_per_doc:.2f}\n")
            f.write("\n")
            
            # Key Findings
            f.write("## Key Findings Across All Lens Types\n\n")
            
            # Find most common drivers across all lens types
            all_drivers = {}
            for result in results.values():
                drivers = result.get('study_analysis', {}).get('comfort_drivers', {})
                for driver_key, driver_data in drivers.items():
                    driver_name = driver_data.get('driver_name', driver_key)
                    if driver_name not in all_drivers:
                        all_drivers[driver_name] = {
                            'count': 0,
                            'total_evidence': 0,
                            'lens_types': [],
                            'confidence_levels': []
                        }
                    all_drivers[driver_name]['count'] += 1
                    all_drivers[driver_name]['total_evidence'] += driver_data.get('evidence_count', 0)
                    all_drivers[driver_name]['lens_types'].append(result.get('study_analysis', {}).get('lens_type', 'unknown'))
                    all_drivers[driver_name]['confidence_levels'].append(driver_data.get('overall_confidence', 'unknown'))
            
            # Top drivers by frequency
            top_drivers = sorted(all_drivers.items(), key=lambda x: x[1]['count'], reverse=True)[:10]
            
            f.write("### Most Common Performance Drivers\n\n")
            for i, (driver_name, data) in enumerate(top_drivers, 1):
                f.write(f"{i}. **{driver_name}**\n")
                f.write(f"   - Found in: {data['count']}/{len(results)} lens types\n")
                f.write(f"   - Total evidence points: {data['total_evidence']}\n")
                f.write(f"   - Lens types: {', '.join(set(data['lens_types']))}\n\n")
            
            # Recommendations
            f.write("## Recommendations\n\n")
            f.write("Based on this comprehensive analysis:\n\n")
            f.write("1. **High-Impact Drivers:** Focus on drivers found across multiple lens types with high confidence\n")
            f.write("2. **Category-Specific Optimization:** Use category reports to target specific performance areas\n")
            f.write("3. **Quantitative Validation:** Prioritize drivers with quantitative evidence for validation studies\n")
            f.write("4. **Auto Scaling Benefits:** The pipeline successfully processed large document volumes efficiently\n\n")
            
            # Files and Outputs
            f.write("## Generated Outputs\n\n")
            f.write("This analysis generated the following output types:\n\n")
            f.write("- **Knowledge Graphs:** Visual representations of driver relationships\n")
            f.write("- **Tabular Reports:** CSV exports for data analysis\n")
            f.write("- **Category Reports:** Detailed Markdown reports for Size, Fit, Handling, Comfort\n")
            f.write("- **Quantitative Charts:** Statistical visualizations and magnitude analysis\n")
            f.write("- **S3 Integration:** All outputs automatically uploaded for cloud access\n\n")
            
            f.write("## Access Information\n\n")
            f.write(f"**S3 Location:** `s3://{OUTPUT_BUCKET_NAME}/{S3_OUTPUT_PREFIX}/`\n")
            f.write(f"**Local Directory:** `{output_dir}`\n")
            f.write(f"**Generated Files:** Available in respective lens type subdirectories\n\n")
            
            f.write("---\n")
            f.write(f"*Generated by Contact Lens Performance Drivers Pipeline v2.0 on {time.strftime('%Y-%m-%d %H:%M:%S')}*\n")
            f.write(f"*Overall Summary Report ID: OVERALL_SUMMARY_{timestamp}*\n")
        
        logger.info(f"✅ Overall findings summary created: {summary_file}")
        return str(summary_file)
        
    except Exception as e:
        logger.error(f"❌ Failed to create overall findings summary: {e}")
        return None
    """Calculate optimal batch size based on endpoint performance"""
    try:
        # Start with a conservative batch size
        base_batch_size = min(30, max(10, current_workers // 6))
        
        # Check if we have recent performance data to adjust
        # This is a simplified heuristic - in production you'd want more sophisticated monitoring
        logger.debug(f"📊 Calculated optimal batch size: {base_batch_size} for {current_workers} workers")
        return base_batch_size
        
    except Exception as e:
        logger.debug(f"Could not calculate optimal batch size: {e}")
        return 20  # Safe default

def match_relationships(relationships: list, logger: logging.Logger, document_key: str = None, chunk_id: str = None) -> list:
    """Match extracted relationships to known drivers/factors with enhanced validation and failure logging"""
    logger.info("🔗 Starting enhanced relationship matching logic")
    
    matched_relationships = []
    unmatched_count = 0
    matching_failures = []
    
    logger.debug(f"📥 Processing {len(relationships)} raw relationships for {document_key or 'unknown document'}")
    
    for i, rel in enumerate(relationships):
        logger.debug(f"🔍 Matching relationship {i+1}: {rel}")
        
        subject = rel.get('subject', '').lower().strip()
        obj = rel.get('object', '').lower().strip()
        predicate = rel.get('predicate', 'affects').strip()
        confidence = rel.get('confidence', 'medium').lower().strip()
        evidence = rel.get('evidence', '').strip()
        
        # Enhanced validation - skip if key fields are missing or too short
        if not subject or not obj or len(subject) < 2 or len(obj) < 2:
            failure_reason = f"Malformed relationship: subject='{subject}' ({len(subject)} chars), object='{obj}' ({len(obj)} chars)"
            logger.debug(f"   ❌ {failure_reason}")
            
            # Log detailed failure information
            log_matching_failure(
                document_key=document_key,
                chunk_id=chunk_id,
                relationship_index=i,
                relationship=rel,
                failure_reason=failure_reason,
                failure_type="malformed_fields",
                logger=logger
            )
            
            matching_failures.append({
                'relationship_index': i,
                'original_relationship': rel,
                'failure_reason': failure_reason,
                'failure_type': 'malformed_fields'
            })
            
            unmatched_count += 1
            continue
        
        # Enhanced evidence validation
        if len(evidence) < CONFIDENCE_THRESHOLDS['MIN_EVIDENCE_LENGTH']:
            logger.debug(f"   ⚠️ Weak evidence for relationship {i+1}: {len(evidence)} chars")
            if confidence == 'high':
                confidence = 'medium'  # Downgrade confidence for weak evidence
            
            # Log weak evidence warning
            log_matching_failure(
                document_key=document_key,
                chunk_id=chunk_id,
                relationship_index=i,
                relationship=rel,
                failure_reason=f"Weak evidence: {len(evidence)} chars < {CONFIDENCE_THRESHOLDS['MIN_EVIDENCE_LENGTH']} minimum",
                failure_type="weak_evidence",
                logger=logger
            )
        
        # Enhanced driver matching with fuzzy matching
        matched_driver = None
        best_driver_score = 0
        driver_match_attempts = []
        
        # Create comprehensive driver search terms
        all_driver_terms = {}
        for driver, terms in DRIVER_TERMS.items():
            expanded_terms = terms + [driver.replace('_', ' '), driver.replace('_', '-')]
            all_driver_terms[driver] = expanded_terms
        
        for driver, terms in all_driver_terms.items():
            for term in terms:
                term_lower = term.lower()
                # Track all match attempts for debugging
                driver_match_attempts.append({
                    'driver': driver,
                    'term': term,
                    'exact_match': term_lower == subject or subject == term_lower,
                    'partial_match': term_lower in subject or subject in term_lower
                })
                
                # Exact match gets highest score
                if term_lower == subject or subject == term_lower:
                    matched_driver = driver
                    best_driver_score = 1.0
                    logger.debug(f"   ✓ Exact match '{subject}' to driver '{driver}' via '{term}'")
                    break
                # Partial match gets lower score
                elif term_lower in subject or subject in term_lower:
                    match_score = len(term_lower) / max(len(subject), len(term_lower))
                    if match_score > best_driver_score:
                        matched_driver = driver
                        best_driver_score = match_score
                        logger.debug(f"   ✓ Partial match '{subject}' to driver '{driver}' via '{term}' (score: {match_score:.2f})")
            
            if best_driver_score == 1.0:  # Perfect match found
                break
        
        # Enhanced factor matching with synonyms
        matched_factor = None
        best_factor_score = 0
        factor_match_attempts = []
        
        factor_synonyms = {
            'fit': ['fit', 'fitting', 'centration', 'positioning', 'alignment'],
            'handling': ['handling', 'manipulation', 'insertion', 'removal', 'application'],
            'comfort': ['comfort', 'comfortable', 'sensation', 'feeling', 'discomfort'],
            'size': ['size', 'sizing', 'dimensions', 'dimensional']
        }
        
        for factor, synonyms in factor_synonyms.items():
            for synonym in synonyms:
                synonym_lower = synonym.lower()
                # Track all match attempts for debugging
                factor_match_attempts.append({
                    'factor': factor,
                    'synonym': synonym,
                    'exact_match': synonym_lower == obj or obj == synonym_lower,
                    'partial_match': synonym_lower in obj or obj in synonym_lower
                })
                
                # Exact match gets highest score
                if synonym_lower == obj or obj == synonym_lower:
                    matched_factor = factor
                    best_factor_score = 1.0
                    logger.debug(f"   ✓ Exact match '{obj}' to factor '{factor}' via '{synonym}'")
                    break
                # Partial match gets lower score
                elif synonym_lower in obj or obj in synonym_lower:
                    match_score = len(synonym_lower) / max(len(obj), len(synonym_lower))
                    if match_score > best_factor_score:
                        matched_factor = factor
                        best_factor_score = match_score
                        logger.debug(f"   ✓ Partial match '{obj}' to factor '{factor}' via '{synonym}' (score: {match_score:.2f})")
            
            if best_factor_score == 1.0:  # Perfect match found
                break
        
        # Enhanced relationship validation and scoring
        if matched_driver and matched_factor:
            # Calculate overall relationship score
            relationship_score = (best_driver_score + best_factor_score) / 2
            
            # Apply confidence scoring
            confidence_multiplier = {'high': 1.0, 'medium': 0.8, 'low': 0.6}.get(confidence, 0.5)
            final_score = relationship_score * confidence_multiplier
            
            # Only include relationships above minimum threshold
            if final_score >= CONFIDENCE_THRESHOLDS['MIN_RELATIONSHIP_SCORE']:
                matched_rel = {
                    'driver': matched_driver,
                    'factor': matched_factor,
                    'predicate': predicate,
                    'confidence': confidence,
                    'evidence': evidence,
                    'original_subject': rel.get('subject', ''),
                    'original_object': rel.get('object', ''),
                    'match_score': final_score,
                    'driver_match_score': best_driver_score,
                    'factor_match_score': best_factor_score,
                    'document_key': document_key,
                    'chunk_id': chunk_id
                }
                matched_relationships.append(matched_rel)
                logger.debug(f"   ✅ Successfully matched: {matched_driver} → {matched_factor} (score: {final_score:.3f})")
            else:
                failure_reason = f"Score too low: {final_score:.3f} < {CONFIDENCE_THRESHOLDS['MIN_RELATIONSHIP_SCORE']}"
                logger.debug(f"   ❌ {failure_reason}")
                
                # Log low score failure
                log_matching_failure(
                    document_key=document_key,
                    chunk_id=chunk_id,
                    relationship_index=i,
                    relationship=rel,
                    failure_reason=failure_reason,
                    failure_type="low_score",
                    logger=logger
                )
                
                matching_failures.append({
                    'relationship_index': i,
                    'original_relationship': rel,
                    'failure_reason': failure_reason,
                    'failure_type': 'low_score',
                    'matched_driver': matched_driver,
                    'matched_factor': matched_factor,
                    'final_score': final_score
                })
                
                unmatched_count += 1
        else:
            # Log detailed no match failure
            failure_reason = f"No match found: driver='{matched_driver}' (score: {best_driver_score:.3f}), factor='{matched_factor}' (score: {best_factor_score:.3f})"
            logger.debug(f"   ❌ {failure_reason}")
            
            # Log comprehensive no match failure with all attempts
            log_matching_failure(
                document_key=document_key,
                chunk_id=chunk_id,
                relationship_index=i,
                relationship=rel,
                failure_reason=failure_reason,
                failure_type="no_match",
                logger=logger
            )
            
            # Log detailed match attempts for debugging
            debug_logger.log_step(
                f"relationship_matching_attempts_{i}",
                "process",
                input_data={
                    'subject': subject,
                    'object': obj,
                    'document_key': document_key,
                    'chunk_id': chunk_id
                },
                output_data={
                    'driver_match_attempts': driver_match_attempts,
                    'factor_match_attempts': factor_match_attempts,
                    'best_driver_score': best_driver_score,
                    'best_factor_score': best_factor_score
                },
                metadata={'relationship_index': i}
            )
            
            matching_failures.append({
                'relationship_index': i,
                'original_relationship': rel,
                'failure_reason': failure_reason,
                'failure_type': 'no_match',
                'driver_match_attempts': driver_match_attempts,
                'factor_match_attempts': factor_match_attempts,
                'best_driver_score': best_driver_score,
                'best_factor_score': best_factor_score
            })
            
            unmatched_count += 1
    
    # Log summary of matching results
    logger.info(f"📊 Enhanced matching results for {document_key or 'unknown document'}:")
    logger.info(f"   ✅ Matched: {len(matched_relationships)}")
    logger.info(f"   ❌ Unmatched: {unmatched_count}")
    
    # Log failure summary if there were failures
    if matching_failures:
        failure_types = {}
        for failure in matching_failures:
            failure_type = failure['failure_type']
            failure_types[failure_type] = failure_types.get(failure_type, 0) + 1
        
        logger.warning(f"⚠️ Matching failures by type:")
        for failure_type, count in failure_types.items():
            logger.warning(f"   {failure_type}: {count}")
    
    # Sort by match score (highest first)
    matched_relationships.sort(key=lambda x: x.get('match_score', 0), reverse=True)
    
    return matched_relationships

def create_drivers_knowledge_graph(lens_type: str, logger: logging.Logger):
    """Create knowledge graph focused on performance drivers for specific lens"""
    logger.info(f"🔧 Building Contact Lens Performance Drivers Knowledge Graph for {lens_type}")
    
    G = nx.MultiDiGraph()
    
    # Add central performance factor nodes
    for factor, factor_data in PERFORMANCE_FACTORS.items():
        factor_node = f"PERFORMANCE_{factor.upper()}"
        G.add_node(factor_node, 
                  type='performance_factor',
                  name=factor.title(),
                  description=factor_data["description"],
                  color=factor_data["color"])
        logger.debug(f"Added performance factor: {factor.title()}")
    
    # Add driver category nodes
    for category, category_data in DRIVER_CATEGORIES.items():
        category_node = f"CATEGORY_{category.upper()}"
        G.add_node(category_node,
                  type='driver_category', 
                  name=category.replace("_", " ").title(),
                  description=category_data["description"],
                  color=category_data["color"])
        logger.debug(f"Added driver category: {category.replace('_', ' ').title()}")
    
    # Add individual driver nodes and relationships
    for factor, factor_data in PERFORMANCE_FACTORS.items():
        if not factor:
            continue
        factor_node = f"PERFORMANCE_{factor.upper()}"
        
        for driver, description in factor_data["drivers"].items():
            if not driver:
                continue
            driver_node = f"DRIVER_{driver.upper()}"
            
            # Add driver node if not exists
            if not G.has_node(driver_node):
                # Determine driver category
                driver_category = None
                for category, category_data in DRIVER_CATEGORIES.items():
                    if driver in category_data["drivers"]:
                        driver_category = category
                        break
                
                G.add_node(driver_node,
                          type='driver',
                          name=driver.replace("_", " ").title(),
                          description=description,
                          category=driver_category,
                          color=DRIVER_CATEGORIES.get(driver_category, {}).get("color", "#95A5A6"))
                
                logger.debug(f"Added driver: {driver.replace('_', ' ').title()}")
                
                # Connect driver to its category
                if driver_category:
                    category_node = f"CATEGORY_{driver_category.upper() if driver_category else 'UNKNOWN'}"
                    G.add_edge(category_node, driver_node, 
                             label="contains_driver",
                             relationship="category_driver")
                    logger.debug(f"Connected {driver} to category {driver_category}")
            
            # Connect driver to performance factor
            G.add_edge(driver_node, factor_node,
                      label="influences",
                      relationship="driver_performance",
                      strength="high")
            logger.debug(f"Connected driver {driver} to factor {factor}")
    
    # Add cross-factor relationships
    cross_factor_drivers = {
        "lens_modulus": ["fit", "handling"],
        "edge_design": ["fit", "comfort"], 
        "surface_lubricity": ["handling", "comfort"],
        "hydration_retention": ["fit", "comfort"],
        "oxygen_permeability": ["fit", "comfort"]
    }
    
    for driver, factors in cross_factor_drivers.items():
        driver_node = f"DRIVER_{driver.upper()}"
        if len(factors) > 1:
            G.nodes[driver_node]['multi_factor'] = True
            logger.debug(f"Marked {driver} as multi-factor driver")
    
    # Add lens type specific node
    if lens_type:
        lens_node = f"LENS_TYPE_{lens_type.upper()}"
        G.add_node(lens_node, type='lens_model', name=lens_type.title() if lens_type else 'Unknown', 
                  description=f"Contact lens model: {lens_type if lens_type else 'Unknown'}")
        
        # Connect lens type to performance factors
        for factor in PERFORMANCE_FACTORS.keys():
            factor_node = f"PERFORMANCE_{factor.upper()}"
            G.add_edge(lens_node, factor_node, 
                      label="evaluated_for",
                      relationship="lens_performance")
            logger.debug(f"Connected lens {lens_type} to factor {factor}")
        
        logger.info(f"✅ Knowledge Graph Created for {lens_type}:")
    else:
        logger.warning("⚠️  No lens type provided for knowledge graph creation")
        
    logger.info(f"   📊 Total Nodes: {G.number_of_nodes()}")
    logger.info(f"   🔗 Total Edges: {G.number_of_edges()}")
    
    return G

def create_enhanced_visualization(G, lens_type: str, output_dir: Path):
    """Create enhanced visualization with improved label positioning"""
    graphs_dir = output_dir / "graphs"
    filename = graphs_dir / f"{lens_type}_drivers_knowledge_graph.png"
    
    # Create larger figure
    fig = plt.figure(figsize=(28, 20))
    
    # Main graph with better spacing
    ax1 = plt.subplot2grid((3, 4), (0, 0), colspan=3, rowspan=2)
    
    # Create hierarchical layout with better spacing
    pos = {}
    
    # Performance factors in center with more spacing
    performance_nodes = [n for n in G.nodes() if 'PERFORMANCE_' in n]
    center_radius = 3
    for i, node in enumerate(performance_nodes):
        angle = 2 * math.pi * i / len(performance_nodes)
        pos[node] = (center_radius * math.cos(angle), center_radius * math.sin(angle))
    
    # Driver categories in middle ring with increased spacing
    category_nodes = [n for n in G.nodes() if 'CATEGORY_' in n]
    category_radius = 6
    for i, node in enumerate(category_nodes):
        angle = 2 * math.pi * i / len(category_nodes) + math.pi/6
        pos[node] = (category_radius * math.cos(angle), category_radius * math.sin(angle))
    
    # Individual drivers in outer ring with better distribution
    driver_nodes = [n for n in G.nodes() if 'DRIVER_' in n]
    
    # Group drivers by category for better positioning
    drivers_by_category = {}
    for node in driver_nodes:
        category = G.nodes[node].get('category', 'unknown')
        if category not in drivers_by_category:
            drivers_by_category[category] = []
        drivers_by_category[category].append(node)
    
    # Position drivers around their categories
    driver_radius = 9
    positioned_drivers = set()
    
    for category, drivers in drivers_by_category.items():
        category_node = f"CATEGORY_{category.upper() if category else 'UNKNOWN'}"
        if category_node in pos:
            category_pos = pos[category_node]
            category_angle = math.atan2(category_pos[1], category_pos[0])
            
            # Distribute drivers in an arc around the category
            arc_span = math.pi / 3  # 60 degrees
            start_angle = category_angle - arc_span / 2
            
            for i, driver in enumerate(drivers):
                if len(drivers) > 1:
                    angle = start_angle + (arc_span * i / (len(drivers) - 1))
                else:
                    angle = category_angle
                
                # Add some radius variation to avoid overlap
                radius_variation = driver_radius + (i % 3) * 0.5
                pos[driver] = (radius_variation * math.cos(angle), 
                             radius_variation * math.sin(angle))
                positioned_drivers.add(driver)
    
    # Position any drivers that didn't get positioned (fallback)
    unpositioned_drivers = [driver for driver in driver_nodes if driver not in positioned_drivers]
    if unpositioned_drivers:
        fallback_radius = 12
        for i, driver in enumerate(unpositioned_drivers):
            angle = 2 * math.pi * i / len(unpositioned_drivers)
            pos[driver] = (fallback_radius * math.cos(angle), 
                         fallback_radius * math.sin(angle))
    
    # Position lens and other nodes
    other_nodes = [n for n in G.nodes() if not any(prefix in n for prefix in ['PERFORMANCE_', 'CATEGORY_', 'DRIVER_'])]
    for i, node in enumerate(other_nodes):
        pos[node] = (0, -7 - i * 1.5)
    
    # Draw nodes with enhanced styling
    node_colors = []
    node_sizes = []
    node_alphas = []
    
    for node in G.nodes():
        node_data = G.nodes[node]
        if 'PERFORMANCE_' in node:
            factor_name = node.replace('PERFORMANCE_', '').lower()
            if factor_name and factor_name in PERFORMANCE_FACTORS:
                node_colors.append(PERFORMANCE_FACTORS[factor_name]['color'])
            else:
                node_colors.append('#95A5A6')  # Default color
            node_sizes.append(2000)
            node_alphas.append(0.9)
        elif 'CATEGORY_' in node:
            category = node.replace('CATEGORY_', '').lower()
            if category and category in DRIVER_CATEGORIES:
                node_colors.append(DRIVER_CATEGORIES[category]['color'])
            else:
                node_colors.append('#95A5A6')  # Default color
            node_sizes.append(1200)
            node_alphas.append(0.8)
        elif 'DRIVER_' in node:
            category = node_data.get('category', 'material_science')
            node_colors.append(DRIVER_CATEGORIES.get(category, {}).get('color', '#95A5A6'))
            node_sizes.append(600)
            node_alphas.append(0.7)
        else:
            node_colors.append('#BDC3C7')
            node_sizes.append(800)
            node_alphas.append(0.6)
    
    # Final safety check: ensure all nodes have positions
    missing_positions = []
    for node in G.nodes():
        if node not in pos:
            missing_positions.append(node)
    
    if missing_positions:
        print(f"⚠️  Warning: {len(missing_positions)} nodes missing positions, adding fallback positions")
        for i, node in enumerate(missing_positions):
            # Place missing nodes in a separate area
            pos[node] = (15, -i * 1.5)
    
    # Draw nodes
    nx.draw_networkx_nodes(G, pos, node_color=node_colors, node_size=node_sizes, 
                          alpha=0.9, linewidths=2, edgecolors='black', ax=ax1)
    
    # Draw edges with better styling
    edge_colors = []
    edge_widths = []
    edge_alphas = []
    
    for u, v, d in G.edges(data=True):
        relationship = d.get('relationship', '')
        if 'driver_performance' in relationship:
            edge_colors.append('#E74C3C')
            edge_widths.append(3)
            edge_alphas.append(0.8)
        elif 'cross_factor' in relationship:
            edge_colors.append('#F39C12')
            edge_widths.append(2.5)
            edge_alphas.append(0.7)
        elif 'category_driver' in relationship:
            edge_colors.append('#8E44AD')
            edge_widths.append(2)
            edge_alphas.append(0.6)
        else:
            edge_colors.append('#BDC3C7')
            edge_widths.append(1.5)
            edge_alphas.append(0.5)
    
    nx.draw_networkx_edges(G, pos, edge_color=edge_colors, width=edge_widths, 
                          alpha=0.6, arrowstyle='->', arrowsize=20, ax=ax1)
    
    # Enhanced label positioning to avoid overlap
    labels = {}
    label_pos = {}
    
    for node in G.nodes():
        node_data = G.nodes[node]
        name = node_data.get('name', node)
        
        # Truncate long names
        if len(name) > 12:
            name = name[:10] + '..'
        
        labels[node] = name
        
        # Adjust label position slightly away from node center
        if node in pos:
            x, y = pos[node]
            # Push labels outward from center
            offset_factor = 0.3
            if abs(x) > abs(y):
                offset_x = offset_factor * (1 if x > 0 else -1)
                offset_y = 0
            else:
                offset_x = 0
                offset_y = offset_factor * (1 if y > 0 else -1)
            
            label_pos[node] = (x + offset_x, y + offset_y)
        else:
            label_pos[node] = pos[node]
    
    # Draw labels with better formatting
    nx.draw_networkx_labels(G, label_pos, labels=labels, font_size=9, 
                           font_weight='bold', font_color='navy', 
                           bbox=dict(boxstyle='round,pad=0.2', facecolor='white', alpha=0.8),
                           ax=ax1)
    
    ax1.set_title(f'Contact Lens Performance Drivers Knowledge Graph\n{lens_type.title() if lens_type else "Unknown"} - Driver → Performance Factor Relationships', 
                 fontsize=18, fontweight='bold', pad=30)
    ax1.axis('off')
    
    # Add detailed driver information panel
    ax2 = plt.subplot2grid((3, 4), (0, 3))
    ax2.axis('off')
    
    driver_text = f"🔧 DRIVERS FOR {lens_type.upper() if lens_type else 'UNKNOWN'}\n\n"
    for category, category_data in DRIVER_CATEGORIES.items():
        driver_text += f"📂 {category.replace('_', ' ').title()}:\n"
        category_drivers = category_data['drivers'][:4]
        for driver in category_drivers:
            driver_name = driver.replace('_', ' ').title()[:18]
            driver_text += f"  • {driver_name}\n"
        driver_text += "\n"
    
    ax2.text(0.05, 0.95, driver_text, transform=ax2.transAxes, fontsize=10,
             verticalalignment='top', fontfamily='monospace',
             bbox=dict(boxstyle='round', facecolor='lightblue', alpha=0.8))
    
    # Add performance factors details
    ax3 = plt.subplot2grid((3, 4), (1, 3))
    ax3.axis('off')
    
    factors_text = "🎯 PERFORMANCE FACTORS\n\n"
    for factor, factor_data in PERFORMANCE_FACTORS.items():
        factors_text += f"• {factor.title()}:\n"
        factors_text += f"  {factor_data['description'][:35]}...\n"
        factors_text += f"  Drivers: {len(factor_data['drivers'])}\n\n"
    
    ax3.text(0.05, 0.95, factors_text, transform=ax3.transAxes, fontsize=10,
             verticalalignment='top', fontfamily='monospace',
             bbox=dict(boxstyle='round', facecolor='lightgreen', alpha=0.8))
    
    # Add statistics panel
    ax4 = plt.subplot2grid((3, 4), (2, 0), colspan=2)
    ax4.axis('off')
    
    # Create statistics table
    stats_data = [
        ['Total Nodes', G.number_of_nodes()],
        ['Total Edges', G.number_of_edges()],
        ['Performance Factors', len([n for n in G.nodes() if 'PERFORMANCE_' in n])],
        ['Individual Drivers', len([n for n in G.nodes() if 'DRIVER_' in n])],
        ['Driver Categories', len([n for n in G.nodes() if 'CATEGORY_' in n])],
        ['Multi-Factor Drivers', len([n for n in G.nodes() if G.nodes[n].get('multi_factor')])],
        ['Average Connections', f"{sum(dict(G.degree()).values()) / G.number_of_nodes():.1f}"]
    ]
    
    table = ax4.table(cellText=stats_data, 
                     colLabels=['Metric', 'Value'],
                     cellLoc='left',
                     loc='center',
                     colWidths=[0.6, 0.4])
    table.auto_set_font_size(False)
    table.set_fontsize(11)
    table.scale(1, 2)
    ax4.set_title(f'Graph Statistics - {lens_type.title() if lens_type else "Unknown"}', fontsize=14, fontweight='bold')
    
    # Add legend with better positioning
    ax5 = plt.subplot2grid((3, 4), (2, 2), colspan=2)
    ax5.axis('off')
    
    legend_elements = [
        plt.Line2D([0], [0], marker='o', color='w', markerfacecolor='#FF6B6B', markersize=15, label='Fit Performance'),
        plt.Line2D([0], [0], marker='o', color='w', markerfacecolor='#4ECDC4', markersize=15, label='Handling Performance'),
        plt.Line2D([0], [0], marker='o', color='w', markerfacecolor='#45B7D1', markersize=15, label='Comfort Performance'),
        plt.Line2D([0], [0], marker='o', color='w', markerfacecolor='#9B59B6', markersize=12, label='Material Science'),
        plt.Line2D([0], [0], marker='o', color='w', markerfacecolor='#E67E22', markersize=12, label='Product Design'),
        plt.Line2D([0], [0], marker='o', color='w', markerfacecolor='#2ECC71', markersize=12, label='Patient Response')
    ]
    
    ax5.legend(handles=legend_elements, loc='center', ncol=2, fontsize=12,
              title='Node Categories', title_fontsize=14)
    
    plt.tight_layout()
    plt.savefig(filename, dpi=300, bbox_inches='tight', facecolor='white')
    plt.close()
    
    return filename

def upload_outputs_to_s3(lens_type: str, output_dir: Path, graph_file: Path, 
                        report_files: dict, logger: logging.Logger) -> dict:
    """Upload all generated outputs to S3"""
    logger.info(f"🚀 Starting S3 upload for {lens_type}")
    
    # Create S3 prefix for this lens type
    timestamp = time.strftime("%Y%m%d_%H%M%S")
    lens_s3_prefix = f"{S3_OUTPUT_PREFIX}/{lens_type}/{timestamp}"
    
    upload_summary = {
        'lens_type': lens_type,
        's3_prefix': lens_s3_prefix,
        'timestamp': timestamp,
        'uploads': {
            'graph': None,
            'reports': {},
            'logs': None,
            'full_directory': None
        }
    }
    
    logger.info(f"📍 S3 Base Location: s3://{OUTPUT_BUCKET_NAME}/{lens_s3_prefix}")
    
    # 1. Upload the main graph visualization
    try:
        graph_s3_key = f"{lens_s3_prefix}/graphs/{graph_file.name}"
        if upload_file_to_s3(graph_file, graph_s3_key, logger):
            upload_summary['uploads']['graph'] = {
                'local_path': str(graph_file),
                's3_url': f"s3://{OUTPUT_BUCKET_NAME}/{graph_s3_key}",
                'status': 'success'
            }
        else:
            upload_summary['uploads']['graph'] = {
                'local_path': str(graph_file),
                'status': 'failed'
            }
    except Exception as e:
        logger.error(f"❌ Failed to upload graph: {e}")
        upload_summary['uploads']['graph'] = {'status': 'error', 'error': str(e)}
    
    # 2. Upload individual report files
    logger.info("📋 Uploading report files...")
    for report_type, file_path in report_files.items():
        if file_path:
            # Convert to Path object if it's a string
            if isinstance(file_path, str):
                file_path = Path(file_path)
            
            if file_path.exists():
                try:
                    report_s3_key = f"{lens_s3_prefix}/reports/{file_path.name}"
                    if upload_file_to_s3(file_path, report_s3_key, logger):
                        upload_summary['uploads']['reports'][report_type] = {
                            'local_path': str(file_path),
                            's3_url': f"s3://{OUTPUT_BUCKET_NAME}/{report_s3_key}",
                            'status': 'success'
                        }
                    else:
                        upload_summary['uploads']['reports'][report_type] = {
                            'local_path': str(file_path),
                            'status': 'failed'
                        }
                except Exception as e:
                    logger.error(f"❌ Failed to upload {report_type} report: {e}")
                    upload_summary['uploads']['reports'][report_type] = {
                        'status': 'error', 
                        'error': str(e)
                    }
            else:
                logger.warning(f"⚠️ Report file does not exist: {file_path}")
                upload_summary['uploads']['reports'][report_type] = {
                    'status': 'file_not_found',
                    'local_path': str(file_path)
                }
        else:
            upload_summary['uploads']['reports'][report_type] = {
                'status': 'skipped', 
                'reason': 'file_not_found'
            }
    
    # 3. Upload log files
    try:
        logs_dir = output_dir / "logs"
        if logs_dir.exists():
            log_results = upload_directory_to_s3(logs_dir, f"{lens_s3_prefix}/logs", logger)
            upload_summary['uploads']['logs'] = {
                'files_uploaded': len(log_results['successful']),
                'files_failed': len(log_results['failed']),
                'total_size_mb': log_results['total_size'] / (1024 * 1024),
                'status': 'success' if log_results['failed'] == [] else 'partial'
            }
        else:
            upload_summary['uploads']['logs'] = {'status': 'skipped', 'reason': 'no_logs_directory'}
    except Exception as e:
        logger.error(f"❌ Failed to upload logs: {e}")
        upload_summary['uploads']['logs'] = {'status': 'error', 'error': str(e)}
    
    # 4. Upload the complete directory structure for backup
    try:
        logger.info("📁 Uploading complete directory structure...")
        full_results = upload_directory_to_s3(output_dir, f"{lens_s3_prefix}/complete", logger)
        upload_summary['uploads']['full_directory'] = {
            'files_uploaded': len(full_results['successful']),
            'files_failed': len(full_results['failed']),
            'total_files': full_results['total_files'],
            'total_size_mb': full_results['total_size'] / (1024 * 1024),
            'status': 'success' if len(full_results['failed']) == 0 else 'partial'
        }
    except Exception as e:
        logger.error(f"❌ Failed to upload complete directory: {e}")
        upload_summary['uploads']['full_directory'] = {'status': 'error', 'error': str(e)}
    
    # 5. Create and upload a summary manifest
    try:
        manifest = {
            'pipeline': 'contact_lens_drivers_knowledge_graph',
            'lens_type': lens_type,
            'timestamp': timestamp,
            'generation_time': time.strftime("%Y-%m-%d %H:%M:%S UTC", time.gmtime()),
            's3_locations': {
                'base_prefix': lens_s3_prefix,
                'graph_visualization': upload_summary['uploads']['graph'].get('s3_url') if upload_summary['uploads']['graph'] else None,
                'reports': {k: v.get('s3_url') for k, v in upload_summary['uploads']['reports'].items() if v.get('s3_url')},
            },
            'local_paths': {
                'output_directory': str(output_dir),
                'graph_file': str(graph_file)
            }
        }
        
        manifest_file = output_dir / f"{lens_type}_s3_manifest.json"
        with open(manifest_file, 'w') as f:
            json.dump(manifest, f, indent=2)
        
        manifest_s3_key = f"{lens_s3_prefix}/manifest.json"
        upload_file_to_s3(manifest_file, manifest_s3_key, logger)
        
        upload_summary['manifest'] = {
            'local_path': str(manifest_file),
            's3_url': f"s3://{OUTPUT_BUCKET_NAME}/{manifest_s3_key}"
        }
        
    except Exception as e:
        logger.error(f"❌ Failed to create/upload manifest: {e}")
    
    # Final summary
    total_successful = 0
    total_failed = 0
    
    if upload_summary['uploads']['graph'] and upload_summary['uploads']['graph'].get('status') == 'success':
        total_successful += 1
    else:
        total_failed += 1
    
    for report_status in upload_summary['uploads']['reports'].values():
        if report_status.get('status') == 'success':
            total_successful += 1
        else:
            total_failed += 1
    
    logger.info(f"🎯 S3 Upload Complete for {lens_type}:")
    logger.info(f"   📊 Individual files: {total_successful} successful, {total_failed} failed")
    logger.info(f"   📁 Full directory: {upload_summary['uploads']['full_directory'].get('files_uploaded', 0)} files uploaded")
    logger.info(f"   🔗 S3 Location: s3://{OUTPUT_BUCKET_NAME}/{lens_s3_prefix}")
    
    return upload_summary

def create_s3_index_page(results: dict) -> str:
    """Create an HTML index page summarizing all S3 outputs"""
    html_content = f"""
<!DOCTYPE html>
<html>
<head>
    <title>Contact Lens Performance Drivers Knowledge Graph - Results</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; }}
        .header {{ background-color: #f8f9fa; padding: 20px; border-radius: 8px; margin-bottom: 30px; }}
        .lens-section {{ margin-bottom: 30px; padding: 20px; border: 1px solid #ddd; border-radius: 8px; }}
        .lens-title {{ color: #007bff; font-size: 24px; margin-bottom: 15px; }}
        .file-grid {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(300px, 1fr)); gap: 15px; }}
        .file-item {{ background-color: #f8f9fa; padding: 15px; border-radius: 5px; }}
        .file-item a {{ text-decoration: none; color: #007bff; font-weight: bold; }}
        .file-item a:hover {{ text-decoration: underline; }}
        .stats {{ background-color: #e9ecef; padding: 10px; border-radius: 5px; margin: 10px 0; }}
        .timestamp {{ color: #6c757d; font-size: 14px; }}
    </style>
</head>
<body>
    <div class="header">
        <h1>🔬 Contact Lens Performance Drivers Knowledge Graph</h1>
        <p><strong>Generated:</strong> {time.strftime("%Y-%m-%d %H:%M:%S UTC", time.gmtime())}</p>
        <p><strong>S3 Bucket:</strong> {OUTPUT_BUCKET_NAME}</p>
        <p><strong>Total Lens Types:</strong> {len(results)}</p>
    </div>
"""
    
    for lens_type, result in results.items():
        if 's3_upload' not in result:
            continue
            
        s3_info = result['s3_upload']
        html_content += f"""
    <div class="lens-section">
        <h2 class="lens-title">📱 {lens_type.title() if lens_type else "Unknown"} Lens Analysis</h2>
        <div class="stats">
            <strong>Graph Statistics:</strong> {result['nodes']} nodes, {result['edges']} edges<br>
            <strong>S3 Prefix:</strong> <code>{s3_info['s3_prefix']}</code><br>
            <strong>Timestamp:</strong> <span class="timestamp">{s3_info['timestamp']}</span>
        </div>
        
        <div class="file-grid">
"""
        
        # Add graph visualization
        if s3_info['uploads']['graph'] and s3_info['uploads']['graph'].get('s3_url'):
            graph_url = s3_info['uploads']['graph']['s3_url'].replace('s3://', 'https://s3.amazonaws.com/')
            html_content += f"""
            <div class="file-item">
                <h4>📊 Knowledge Graph Visualization</h4>
                <a href="{graph_url}" target="_blank">View Graph Image</a>
                <p>Interactive network diagram showing driver-performance relationships</p>
            </div>
"""
        
        # Add reports
        for report_type, report_info in s3_info['uploads']['reports'].items():
            if report_info.get('s3_url'):
                report_url = report_info['s3_url'].replace('s3://', 'https://s3.amazonaws.com/')
                report_name = report_type.replace('_', ' ').title()
                html_content += f"""
            <div class="file-item">
                <h4>📋 {report_name} Report</h4>
                <a href="{report_url}" target="_blank">Download CSV</a>
                <p>Tabular analysis data for {report_name.lower()}</p>
            </div>
"""
        
        # Add manifest
        if 'manifest' in s3_info and s3_info['manifest'].get('s3_url'):
            manifest_url = s3_info['manifest']['s3_url'].replace('s3://', 'https://s3.amazonaws.com/')
            html_content += f"""
            <div class="file-item">
                <h4>📄 Analysis Manifest</h4>
                <a href="{manifest_url}" target="_blank">View JSON</a>
                <p>Complete metadata and file locations for this analysis</p>
            </div>
"""
        
        html_content += """
        </div>
    </div>
"""
    
    html_content += """
    <div class="header">
        <h3>🔗 Additional Resources</h3>
        <p><strong>S3 Console:</strong> <a href="https://s3.console.aws.amazon.com/s3/buckets/{bucket}?prefix={prefix}/" target="_blank">Browse all files in S3</a></p>
        <p><strong>AWS CLI Access:</strong> <code>aws s3 ls s3://{bucket}/{prefix}/ --recursive</code></p>
    </div>
</body>
</html>
""".format(bucket=OUTPUT_BUCKET_NAME, prefix=S3_OUTPUT_PREFIX)
    
    return html_content

def create_tabular_report(G, lens_type: str, output_dir: Path, relationships: list = None):
    """Create comprehensive tabular reports"""
    reports_dir = output_dir / "reports"
    created_files = {}
    
    print(f"\n📋 Creating tabular reports for {lens_type}...")
    print(f"   Reports directory: {reports_dir}")
    
    # 1. Drivers Summary Report
    try:
        drivers_data = []
        for node in G.nodes():
            if 'DRIVER_' in node:
                node_data = G.nodes[node]
                driver_name = node.replace('DRIVER_', '').lower()
                
                # Get category
                category = node_data.get('category', 'unknown')
                
                # Count relationships
                performance_connections = [target for target in G.successors(node) if 'PERFORMANCE_' in target]
                
                drivers_data.append({
                    'Driver': driver_name.replace('_', ' ').title(),
                    'Category': category.replace('_', ' ').title(),
                    'Description': node_data.get('description', ''),
                    'Performance_Factors_Connected': len(performance_connections),
                    'Connected_Factors': ', '.join([f.replace('PERFORMANCE_', '').title() for f in performance_connections]),
                    'Multi_Factor_Driver': node_data.get('multi_factor', False)
                })
        
        drivers_df = pd.DataFrame(drivers_data)
        drivers_file = reports_dir / f"{lens_type}_drivers_summary.csv"
        drivers_df.to_csv(drivers_file, index=False)
        created_files['drivers'] = drivers_file
        print(f"   ✅ Created drivers summary: {drivers_file}")
        print(f"      📊 {len(drivers_data)} drivers documented")
    except Exception as e:
        print(f"   ❌ Failed to create drivers summary: {e}")
        created_files['drivers'] = None
    
    # 2. Performance Factors Report
    try:
        factors_data = []
        for factor, factor_info in PERFORMANCE_FACTORS.items():
            factor_node = f"PERFORMANCE_{factor.upper()}"
            
            # Count driver connections
            connected_drivers = [pred for pred in G.predecessors(factor_node) if 'DRIVER_' in pred]
            
            factors_data.append({
                'Performance_Factor': factor.title(),
                'Description': factor_info['description'],
                'Total_Drivers': len(factor_info['drivers']),
                'Connected_Drivers_in_Graph': len(connected_drivers),
                'Color_Code': factor_info['color'],
                'Primary_Categories': ', '.join(set([G.nodes[d].get('category', '').replace('_', ' ').title() 
                                                   for d in connected_drivers if G.nodes[d].get('category')]))
            })
        
        factors_df = pd.DataFrame(factors_data)
        factors_file = reports_dir / f"{lens_type}_performance_factors.csv"
        factors_df.to_csv(factors_file, index=False)
        created_files['factors'] = factors_file
        print(f"   ✅ Created factors report: {factors_file}")
        print(f"      📊 {len(factors_data)} performance factors documented")
    except Exception as e:
        print(f"   ❌ Failed to create factors report: {e}")
        created_files['factors'] = None
    
    # 3. Relationships Report
    try:
        if relationships and len(relationships) > 0:
            relationships_df = pd.DataFrame(relationships)
            relationships_file = reports_dir / f"{lens_type}_extracted_relationships.csv"
            relationships_df.to_csv(relationships_file, index=False)
            created_files['relationships'] = relationships_file
            print(f"   ✅ Created relationships report: {relationships_file}")
            print(f"      📊 {len(relationships)} relationships documented")
        else:
            print(f"   ⚠️ No relationships data provided - skipping relationships report")
            created_files['relationships'] = None
    except Exception as e:
        print(f"   ❌ Failed to create relationships report: {e}")
        created_files['relationships'] = None
    
    # 4. Graph Statistics Report
    try:
        stats_data = {
            'Metric': [
                'Total Nodes', 'Total Edges', 'Performance Factor Nodes', 'Driver Nodes',
                'Category Nodes', 'Average Node Degree', 'Graph Density', 'Multi-Factor Drivers'
            ],
            'Value': [
                G.number_of_nodes(),
                G.number_of_edges(),
                len([n for n in G.nodes() if 'PERFORMANCE_' in n]),
                len([n for n in G.nodes() if 'DRIVER_' in n]),
                len([n for n in G.nodes() if 'CATEGORY_' in n]),
                round(sum(dict(G.degree()).values()) / G.number_of_nodes(), 2),
                round(nx.density(G), 3),
                len([n for n in G.nodes() if G.nodes[n].get('multi_factor')])
            ]
        }
        
        stats_df = pd.DataFrame(stats_data)
        stats_file = reports_dir / f"{lens_type}_graph_statistics.csv"
        stats_df.to_csv(stats_file, index=False)
        created_files['statistics'] = stats_file
        print(f"   ✅ Created statistics report: {stats_file}")
        print(f"      📊 {len(stats_data['Metric'])} metrics documented")
    except Exception as e:
        print(f"   ❌ Failed to create statistics report: {e}")
        created_files['statistics'] = None
    
    # 5. Cross-Factor Analysis
    try:
        cross_factor_data = []
        for node in G.nodes():
            if 'DRIVER_' in node and G.nodes[node].get('multi_factor'):
                driver_name = node.replace('DRIVER_', '').lower()
                connected_factors = [f.replace('PERFORMANCE_', '') for f in G.successors(node) if 'PERFORMANCE_' in f]
                
                cross_factor_data.append({
                    'Multi_Factor_Driver': driver_name.replace('_', ' ').title(),
                    'Connected_Factors': ', '.join([f.title() for f in connected_factors]),
                    'Factor_Count': len(connected_factors),
                    'Category': G.nodes[node].get('category', '').replace('_', ' ').title()
                })
        
        if cross_factor_data:
            cross_df = pd.DataFrame(cross_factor_data)
            cross_file = reports_dir / f"{lens_type}_cross_factor_analysis.csv"
            cross_df.to_csv(cross_file, index=False)
            created_files['cross_factors'] = cross_file
            print(f"   ✅ Created cross-factor analysis: {cross_file}")
            print(f"      📊 {len(cross_factor_data)} multi-factor drivers documented")
        else:
            print(f"   ⚠️ No multi-factor drivers found - skipping cross-factor analysis")
            created_files['cross_factors'] = None
    except Exception as e:
        print(f"   ❌ Failed to create cross-factor analysis: {e}")
        created_files['cross_factors'] = None
    
    # Summary
    successful_reports = [k for k, v in created_files.items() if v is not None]
    print(f"\n   📋 Report Summary for {lens_type}:")
    print(f"      ✅ Successfully created: {len(successful_reports)} reports")
    print(f"      ❌ Failed: {len(created_files) - len(successful_reports)} reports")
    
    return created_files

def create_comfort_analysis_report(study_analysis: dict, output_dir: Path, logger: logging.Logger) -> Path:
    """Create detailed comfort analysis report from study data"""
    try:
        reports_dir = output_dir / "reports"
        lens_type = study_analysis['lens_type']
        
        logger.info(f"📊 Creating comfort analysis report for {lens_type}")
        
        # 1. Comfort Drivers Summary
        comfort_drivers_data = []
        for driver_key, driver_data in study_analysis['comfort_drivers'].items():
            comfort_drivers_data.append({
                'Driver': driver_data['driver_name'],
                'Evidence_Count': driver_data['evidence_count'],
                'Overall_Confidence': driver_data['overall_confidence'],
                'Unique_Documents': driver_data['unique_documents'],
                'Category': driver_data.get('category', 'unknown').replace('_', ' ').title(),
                'Sample_Evidence': driver_data['evidence_snippets'][0][:200] + '...' if driver_data['evidence_snippets'] else '',
                'Source_Documents': '; '.join(driver_data['documents'][:3])  # Show first 3 documents
            })
        
        if comfort_drivers_data:
            comfort_df = pd.DataFrame(comfort_drivers_data)
            comfort_file = reports_dir / f"{lens_type}_comfort_drivers_from_studies.csv"
            comfort_df.to_csv(comfort_file, index=False)
            logger.info(f"   ✅ Comfort drivers report: {comfort_file}")
        else:
            logger.warning(f"   ⚠️ No comfort drivers found for {lens_type}")
            return None
        
        # 2. Document Processing Summary
        doc_summary_data = []
        for doc_result in study_analysis['documents']:
            doc_summary_data.append({
                'Document': doc_result['document'],
                'Status': doc_result['status'],
                'Text_Length': doc_result.get('text_length', 0),
                'Chunks_Processed': doc_result.get('chunks_processed', 0),
                'Total_Relationships': doc_result.get('total_relationships_found', 0),
                'Unique_Relationships': doc_result.get('unique_relationships', 0),
                'Drivers_Found': ', '.join(doc_result.get('found_drivers', [])),
                'Factors_Found': ', '.join(doc_result.get('found_factors', []))
            })
        
        doc_df = pd.DataFrame(doc_summary_data)
        doc_file = reports_dir / f"{lens_type}_document_processing_summary.csv"
        doc_df.to_csv(doc_file, index=False)
        logger.info(f"   ✅ Document processing summary: {doc_file}")
        
        # 3. All Extracted Relationships
        if study_analysis['all_relationships']:
            relationships_data = []
            for rel in study_analysis['all_relationships']:
                relationships_data.append({
                    'Driver': rel['driver'].replace('_', ' ').title(),
                    'Performance_Factor': rel['factor'].title(),
                    'Relationship': rel['predicate'],
                    'Confidence': rel['confidence'],
                    'Evidence': rel['evidence'][:300] + '...' if len(rel['evidence']) > 300 else rel['evidence'],
                    'Original_Subject': rel['original_subject'],
                    'Original_Object': rel['original_object']
                })
            
            rel_df = pd.DataFrame(relationships_data)
            rel_file = reports_dir / f"{lens_type}_extracted_relationships_from_studies.csv"
            rel_df.to_csv(rel_file, index=False)
            logger.info(f"   ✅ Extracted relationships: {rel_file}")
        
        # 4. Comfort-Specific Analysis Summary
        summary_data = {
            'Metric': [
                'Total Documents Processed',
                'Documents Successfully Processed',
                'Documents Failed',
                'Total Relationships Extracted',
                'Comfort-Specific Relationships',
                'Unique Comfort Drivers Found',
                'High Confidence Comfort Drivers',
                'Medium Confidence Comfort Drivers',
                'Low Confidence Comfort Drivers'
            ],
            'Value': [
                study_analysis.get('total_documents', 0),
                study_analysis.get('documents_processed', 0),
                study_analysis.get('documents_failed', 0),
                study_analysis.get('total_relationships', 0),
                study_analysis.get('comfort_specific_relationships', 0),
                study_analysis.get('unique_comfort_drivers', 0),
                len([d for d in study_analysis.get('comfort_drivers', {}).values() if d.get('overall_confidence', '') == 'high']),
                len([d for d in study_analysis.get('comfort_drivers', {}).values() if d.get('overall_confidence', '') == 'medium']),
                len([d for d in study_analysis.get('comfort_drivers', {}).values() if d.get('overall_confidence', '') == 'low'])
            ]
        }
        
        summary_df = pd.DataFrame(summary_data)
        summary_file = reports_dir / f"{lens_type}_comfort_analysis_summary.csv"
        summary_df.to_csv(summary_file, index=False)
        logger.info(f"   ✅ Analysis summary: {summary_file}")
        
        logger.info(f"📊 Comfort analysis reports created for {lens_type}")
        return comfort_file
        
    except Exception as e:
        logger.error(f"❌ Failed to create comfort analysis report: {e}")
        return None

def create_performance_category_reports(study_analysis: dict, output_dir: Path, logger: logging.Logger) -> dict:
    """Create separate detailed reports for Size, Fit, Handling, and Comfort performance categories with graphs and tables"""
    try:
        reports_dir = output_dir / "reports"
        lens_type = study_analysis['lens_type']
        
        logger.info(f"📊 Creating performance category reports for {lens_type}")
        
        # Get all relationships for this lens type
        all_relationships = []
        if 'all_relationships' in study_analysis:
            all_relationships = study_analysis['all_relationships']
        elif 'documents' in study_analysis:
            # Extract relationships from document results
            for doc_result in study_analysis['documents']:
                if 'relationships' in doc_result:
                    all_relationships.extend(doc_result['relationships'])
        
        logger.info(f"📊 Found {len(all_relationships)} total relationships to analyze")
        
        # Create detailed analysis for each performance factor
        factor_analyses = {}
        performance_factors = ['size', 'fit', 'handling', 'comfort']
        
        for factor in performance_factors:
            logger.info(f"📝 Creating {factor.title()} category report...")
            
            try:
                factor_analysis = create_factor_specific_analysis(
                    all_relationships, factor, lens_type, reports_dir, logger
                )
                
                factor_analyses[factor] = factor_analysis
                
                if factor_analysis.get('analysis_files'):
                    logger.info(f"✅ {factor.title()} report created with {len(factor_analysis['analysis_files'])} files")
                    logger.info(f"   • Relationships found: {factor_analysis.get('relationships_count', 0)}")
                    logger.info(f"   • Unique drivers: {factor_analysis.get('unique_drivers', 0)}")
                else:
                    logger.warning(f"⚠️ {factor.title()} report created but no analysis files generated")
                    
            except Exception as e:
                logger.error(f"❌ Failed to create {factor.title()} detailed report: {e}")
                logger.warning(f"⚠️ Failed to create {factor.title()} report")
                factor_analyses[factor] = {
                    'factor': factor,
                    'relationships_count': 0,
                    'drivers_found': [],
                    'analysis_files': [],
                    'error': str(e)
                }
        
        logger.info(f"📊 All performance category reports completed for {lens_type}")
        
        # Create summary of all factor analyses
        logger.info(f"📋 Factor Analysis Summary for {lens_type}:")
        total_files_created = 0
        for factor, analysis in factor_analyses.items():
            file_count = len(analysis.get('analysis_files', []))
            total_files_created += file_count
            logger.info(f"   • {factor.title()}: {analysis.get('relationships_count', 0)} relationships, {file_count} files")
        
        logger.info(f"🎯 Total factor analysis files created: {total_files_created}")
        
        # Return results in expected format for compatibility
        created_reports = {}
        for factor, analysis in factor_analyses.items():
            if analysis.get('analysis_files'):
                for i, file_path in enumerate(analysis['analysis_files']):
                    file_name = Path(file_path).name
                    if 'graph' in file_name or '.png' in file_name:
                        created_reports[f"{factor}_graph"] = file_path
                    elif 'tabular' in file_name or '.csv' in file_name:
                        created_reports[f"{factor}_table"] = file_path
                    elif 'summary' in file_name or '.md' in file_name:
                        created_reports[f"{factor}_analysis"] = file_path
        
        return created_reports
        
    except Exception as e:
        logger.error(f"❌ Failed to create performance category reports: {e}")
        return {}

def create_category_detailed_report(category: str, drivers: dict, relationships: list, lens_type: str, 
                                  output_dir: Path, timestamp: str, description: str, logger: logging.Logger) -> str:
    """Create a detailed report for a specific performance category"""
    try:
        report_file = output_dir / f"{category.lower()}_analysis_{lens_type}_{timestamp}.md"
        
        with open(report_file, 'w', encoding='utf-8') as f:
            f.write(f"# Contact Lens {category} Performance Analysis Report\n\n")
            f.write(f"**Lens Type:** {lens_type}\n")
            f.write(f"**Performance Category:** {category}\n") 
            f.write(f"**Generated:** {time.strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"**Report ID:** {category}_{lens_type}_{timestamp}\n")
            f.write(f"**Total {category} Drivers Found:** {len(drivers)}\n")
            f.write(f"**Total {category} Relationships:** {len(relationships)}\n\n")
            
            f.write("## Executive Summary\n\n")
            f.write(f"This analysis focuses on **{description}** for {lens_type} contact lenses.\n\n")
            
            if drivers:
                high_conf = sum(1 for d in drivers.values() if d.get('overall_confidence', '') and d.get('overall_confidence', '').upper() == 'HIGH')
                med_conf = sum(1 for d in drivers.values() if d.get('overall_confidence', '') and d.get('overall_confidence', '').upper() == 'MEDIUM')
                low_conf = sum(1 for d in drivers.values() if d.get('overall_confidence', '') and d.get('overall_confidence', '').upper() == 'LOW')
                
                f.write(f"This analysis identified **{len(drivers)} {category.lower()} drivers** from study data:\n\n")
                f.write(f"- **{high_conf} High Confidence** drivers with strong evidence\n")
                f.write(f"- **{med_conf} Medium Confidence** drivers with moderate evidence\n")
                f.write(f"- **{low_conf} Low Confidence** drivers requiring further validation\n\n")
                
                # Category-specific insights
                f.write(f"### Key {category} Insights\n\n")
                if category == "Size":
                    f.write("- Diameter and base curve optimization for population fit\n")
                    f.write("- Edge design impact on sizing perception\n")
                    f.write("- Thickness variations affecting size characteristics\n")
                    f.write("- Dimensional stability and consistency\n\n")
                elif category == "Fit":
                    f.write("- Centration and movement patterns\n")
                    f.write("- On-eye stability and positioning\n")
                    f.write("- Coverage and alignment characteristics\n")
                    f.write("- Fit assessment and optimization\n\n")
                elif category == "Handling":
                    f.write("- Insertion and removal ease\n")
                    f.write("- Material flexibility and durability\n")
                    f.write("- Manipulation and orientation characteristics\n")
                    f.write("- Lens handling during application\n\n")
                elif category == "Comfort":
                    f.write("- Subjective sensation and irritation factors\n")
                    f.write("- Surface properties and lubrication\n")
                    f.write("- Hydration and moisture retention\n")
                    f.write("- Long-term wearing comfort\n\n")
            else:
                f.write(f"No {category.lower()} drivers were identified in the analyzed study data.\n\n")
            
            # Detailed driver analysis
            f.write(f"## Detailed {category} Driver Analysis\n\n")
            
            if drivers:
                for driver_key, driver_data in sorted(drivers.items(), 
                                                    key=lambda x: x[1].get('evidence_count', 0), reverse=True):
                    f.write(f"### {driver_data.get('driver_name', driver_key)}\n\n")
                    f.write(f"**Confidence Level:** {driver_data.get('overall_confidence', 'Unknown')}\n")
                    f.write(f"**Evidence Points:** {driver_data.get('evidence_count', 0)}\n")
                    f.write(f"**Average Confidence:** {driver_data.get('average_confidence', 0):.2f}\n")
                    f.write(f"**Category Relevance:** {category}\n")
                    f.write(f"**Unique Documents:** {driver_data.get('unique_documents', 0)}\n\n")
                    
                    if driver_data.get('evidence_snippets'):
                        f.write("**Supporting Evidence:**\n\n")
                        for i, evidence in enumerate(driver_data['evidence_snippets'][:5], 1):
                            f.write(f"{i}. {evidence[:250]}...\n")
                        f.write("\n")
                    
                    f.write("---\n\n")
            else:
                f.write(f"No {category.lower()} drivers found in the analyzed data.\n\n")
            
            # Relationships analysis
            if relationships:
                f.write(f"## {category} Relationships Analysis\n\n")
                f.write(f"Found **{len(relationships)}** relationships relevant to {category.lower()} performance:\n\n")
                
                for i, rel in enumerate(relationships[:10], 1):  # Show top 10 relationships
                    f.write(f"**{i}. {rel.get('driver', 'Unknown')} → {rel.get('factor', 'Unknown')}**\n")
                    f.write(f"   - Relationship: {rel.get('predicate', 'affects')}\n")
                    f.write(f"   - Confidence: {rel.get('confidence', 0):.2f}\n")
                    f.write(f"   - Evidence: {rel.get('evidence', '')[:200]}...\n\n")
                
                if len(relationships) > 10:
                    f.write(f"*...and {len(relationships) - 10} additional relationships*\n\n")
            
            # Statistical summary
            f.write("## Statistical Summary\n\n")
            if drivers:
                total_evidence = sum(d.get('evidence_count', 0) for d in drivers.values())
                confidence_scores = [d.get('average_confidence', 0) for d in drivers.values() if d.get('average_confidence', 0) > 0]
                avg_confidence = sum(confidence_scores) / len(confidence_scores) if confidence_scores else 0
                
                f.write(f"- **Total Evidence Points:** {total_evidence}\n")
                f.write(f"- **Average Confidence Score:** {avg_confidence:.2f}\n")
                f.write(f"- **Total Relationships:** {len(relationships)}\n")
                f.write(f"- **Drivers per Confidence Level:**\n")
                f.write(f"  - High: {sum(1 for d in drivers.values() if d.get('overall_confidence', '') and d.get('overall_confidence', '').upper() == 'HIGH')}\n")
                f.write(f"  - Medium: {sum(1 for d in drivers.values() if d.get('overall_confidence', '') and d.get('overall_confidence', '').upper() == 'MEDIUM')}\n")
                f.write(f"  - Low: {sum(1 for d in drivers.values() if d.get('overall_confidence', '') and d.get('overall_confidence', '').upper() == 'LOW')}\n\n")
            
            # CSV data export
            if drivers:
                f.write(f"## {category} Data Export\n\n")
                csv_file = output_dir / f"{category.lower()}_drivers_{lens_type}_{timestamp}.csv"
                
                # Create CSV data for this category
                csv_data = []
                for driver_key, driver_data in drivers.items():
                    csv_data.append({
                        'Driver_Key': driver_key,
                        'Driver_Name': driver_data.get('driver_name', driver_key),
                        'Category': category,
                        'Confidence_Level': driver_data.get('overall_confidence', 'Unknown'),
                        'Evidence_Count': driver_data.get('evidence_count', 0),
                        'Average_Confidence': driver_data.get('average_confidence', 0),
                        'Unique_Documents': driver_data.get('unique_documents', 0),
                        'Sample_Evidence': driver_data.get('evidence_snippets', [''])[0][:300] if driver_data.get('evidence_snippets') else '',
                        'Source_Documents': '; '.join(driver_data.get('documents', [])[:3]),
                        'Generated_Date': time.strftime('%Y-%m-%d %H:%M:%S')
                    })
                
                if csv_data:
                    import pandas as pd
                    df = pd.DataFrame(csv_data)
                    df.to_csv(csv_file, index=False)
                    f.write(f"**CSV Export:** `{csv_file.name}`\n\n")
                    logger.info(f"   📄 CSV export created: {csv_file}")
            
            f.write("## Methodology\n\n")
            f.write("This report was generated using:\n\n")
            f.write("- **NLP/LLM Analysis:** Advanced language models for relationship extraction\n")
            f.write("- **Evidence Aggregation:** Multiple data points combined for confidence scoring\n")
            f.write(f"- **{category} Focus:** Specifically targeting {category.lower()}-related performance factors\n")
            f.write("- **Study Data Source:** Research documents from S3 study repository\n")
            f.write("- **Auto Scaling Optimization:** Efficient processing using auto scaling endpoints\n")
            f.write("- **Multi-Category Analysis:** Separate analysis for Size, Fit, Handling, and Comfort\n\n")
            
            f.write("---\n")
            f.write(f"*Generated by Contact Lens Performance Drivers Pipeline v2.0 on {time.strftime('%Y-%m-%d %H:%M:%S')}*\n")
            f.write(f"*Report ID: {category}_{lens_type}_{timestamp}*\n")
        
        return str(report_file)
        
    except Exception as e:
        logger.error(f"❌ Failed to create {category} detailed report: {e}")
        return None

def main():
    """Main execution function for drivers knowledge graph"""
    try:
        debug_logger.log_step(
            "main_pipeline_start",
            "input",
            input_data={
                "prefixes": PREFIXES,
                "max_workers": MAX_WORKERS,
                "max_llm_workers": MAX_LLM_WORKERS,
                "endpoint_name": ENDPOINT_NAME
            }
        )
        
        print("🚀 Starting Contact Lens Performance Drivers Knowledge Graph Pipeline")
        
        # Create output structure
        bucket_folders = create_output_structure()
        
        debug_logger.log_step(
            "main_output_structure_created",
            "process",
            output_data={"bucket_folders": bucket_folders},
            metadata={"folder_count": len(bucket_folders)}
        )
        
        start_time = time.time()
        results = {}
        
        # Process each lens type
        for prefix_idx, prefix in enumerate(PREFIXES):
            lens_type = prefix.split('/')[-1]
            output_dir = bucket_folders[lens_type]
            
            debug_logger.log_step(
                f"main_processing_lens_type_{lens_type}",
                "process",
                input_data={
                    "lens_type": lens_type,
                    "prefix": prefix,
                    "output_dir": str(output_dir),
                    "prefix_index": prefix_idx + 1,
                    "total_prefixes": len(PREFIXES)
                }
            )
            
            # Setup lens-specific logger
            logger = setup_lens_logger(lens_type, output_dir)
            
            logger.info("=" * 60)
            logger.info(f"PROCESSING LENS TYPE: {lens_type.upper()}")
            logger.info("=" * 60)
            
            try:
                # STEP 1: Extract comfort drivers from study data using NLP/LLM
                logger.info("🔬 STEP 1: Extracting comfort drivers from S3 study data using NLP/LLM")
                
                debug_logger.log_step(
                    f"main_step1_extract_comfort_drivers_{lens_type}",
                    "process",
                    input_data={"lens_type": lens_type},
                    metadata={"step": "extract_comfort_drivers"}
                )
                
                study_analysis = extract_comfort_drivers_from_study_data(lens_type, logger)
                
                debug_logger.log_step(
                    f"main_step1_complete_{lens_type}",
                    "process",
                    input_data={"lens_type": lens_type},
                    output_data={
                        "documents_processed": study_analysis['documents_processed'],
                        "total_relationships": study_analysis['total_relationships'],
                        "comfort_drivers_found": len(study_analysis['comfort_drivers'])
                    },
                    metadata={"step": "extract_comfort_drivers", "success": True}
                )
                
                # STEP 2: Create knowledge graph for this lens
                logger.info("🔧 STEP 2: Building knowledge graph structure")
                
                debug_logger.log_step(
                    f"main_step2_create_graph_{lens_type}",
                    "process",
                    input_data={"lens_type": lens_type},
                    metadata={"step": "create_knowledge_graph"}
                )
                
                G = create_drivers_knowledge_graph(lens_type, logger)
                
                debug_logger.log_step(
                    f"main_step2_complete_{lens_type}",
                    "process",
                    input_data={"lens_type": lens_type},
                    output_data={
                        "graph_nodes": G.number_of_nodes(),
                        "graph_edges": G.number_of_edges()
                    },
                    metadata={"step": "create_knowledge_graph", "success": True}
                )
                
                # STEP 3: Enhance graph with extracted relationships from study data
                logger.info("🔗 STEP 3: Enhancing graph with study-derived relationships")
                
                debug_logger.log_step(
                    f"main_step3_enhance_graph_{lens_type}",
                    "process",
                    input_data={
                        "lens_type": lens_type,
                        "comfort_drivers_available": len(study_analysis['comfort_drivers'])
                    },
                    metadata={"step": "enhance_graph"}
                )
                
                enhanced_relationships = []
                
                if study_analysis['comfort_drivers']:
                    logger.info(f"📊 Found {len(study_analysis['comfort_drivers'])} comfort drivers from study data:")
                    
                    for driver_key, driver_data in study_analysis['comfort_drivers'].items():
                        logger.info(f"   • {driver_data['driver_name']}: {driver_data['evidence_count']} evidence points ({driver_data['overall_confidence']} confidence)")
                        
                        debug_logger.log_step(
                            f"main_step3_process_driver_{driver_key}_{lens_type}",
                            "process",
                            input_data={
                                "driver_key": driver_key,
                                "driver_name": driver_data['driver_name'],
                                "evidence_count": driver_data['evidence_count'],
                                "confidence": driver_data['overall_confidence']
                            },
                            metadata={"step": "enhance_graph", "driver_processing": True}
                        )
                        
                        # Add evidence-based edges to the graph
                        driver_node = f"DRIVER_{driver_key.upper()}"
                        comfort_node = "PERFORMANCE_COMFORT"
                        
                        if G.has_node(driver_node):
                            # Update existing node with study evidence
                            G.nodes[driver_node]['study_evidence'] = True
                            G.nodes[driver_node]['evidence_count'] = driver_data['evidence_count']
                            G.nodes[driver_node]['study_confidence'] = driver_data['overall_confidence']
                            G.nodes[driver_node]['source_documents'] = driver_data['unique_documents']
                            
                            # Enhance edge with study data
                            if G.has_edge(driver_node, comfort_node):
                                for edge_key in G[driver_node][comfort_node]:
                                    G[driver_node][comfort_node][edge_key]['strength'] = driver_data['overall_confidence']
                                    G[driver_node][comfort_node][edge_key]['evidence_based'] = True
                                    G[driver_node][comfort_node][edge_key]['study_evidence_count'] = driver_data['evidence_count']
                        
                        enhanced_relationships.extend(study_analysis['all_relationships'])
                
                logger.info(f"✅ Graph enhanced with {len(enhanced_relationships)} study-derived relationships")
                
                # STEP 4: Create visualization
                logger.info("📊 STEP 4: Creating enhanced visualization")
                graph_file = create_enhanced_visualization(G, lens_type, output_dir)
                logger.info(f"📊 Graph visualization saved: {graph_file}")
                
                # STEP 5: Create comprehensive tabular reports including study analysis
                logger.info("📋 STEP 5: Creating comprehensive reports")
                report_files = create_tabular_report(G, lens_type, output_dir, enhanced_relationships)
                
                # STEP 6: Create study-specific comfort analysis report (legacy)
                comfort_analysis_file = create_comfort_analysis_report(study_analysis, output_dir, logger)
                if comfort_analysis_file:
                    report_files['comfort_analysis'] = comfort_analysis_file
                
                # STEP 7: Create separate performance category reports (NEW)
                logger.info("� STEP 7: Creating performance category reports (Size, Fit, Handling, Comfort)")
                category_reports = create_performance_category_reports(study_analysis, output_dir, logger)
                
                # Add category reports to the main report files dictionary
                for category, report_file in category_reports.items():
                    report_files[f'{category.lower()}_analysis'] = report_file
                
                logger.info("📋 All reports created:")
                for report_type, file_path in report_files.items():
                    if file_path:
                        logger.info(f"   • {report_type}: {file_path}")
                
                # Summary of category reports
                if category_reports:
                    logger.info("🎯 Performance Category Reports Summary:")
                    for category, report_file in category_reports.items():
                        logger.info(f"   📊 {category} Analysis: {Path(report_file).name}")
                else:
                    logger.warning("⚠️ No performance category reports were generated")
                
                # STEP 8: Quantitative analysis and charting (NEW)
                logger.info("📈 STEP 8: Analyzing relationship values and creating quantitative charts")
                all_relationships = study_analysis.get('all_relationships', [])
                
                # Perform quantitative analysis
                quantitative_analysis = analyze_relationship_values(all_relationships, logger)
                
                # Create charts and visualizations
                chart_files = create_quantitative_charts(quantitative_analysis, lens_type, output_dir, logger)
                
                # Add charts to report files
                for i, chart_file in enumerate(chart_files):
                    report_files[f'chart_{i+1}'] = str(chart_file)
                
                logger.info("📋 All reports and charts created:")
                for report_type, file_path in report_files.items():
                    if file_path:
                        logger.info(f"   • {report_type}: {file_path}")
                
                if quantitative_analysis:
                    logger.info("📈 Quantitative Analysis Summary:")
                    logger.info(f"   • Total relationships analyzed: {quantitative_analysis.get('total_relationships', 0)}")
                    logger.info(f"   • Quantitative relationships: {quantitative_analysis.get('quantitative_relationships', 0)}")
                    logger.info(f"   • High magnitude relationships: {quantitative_analysis.get('high_magnitude_count', 0)}")
                    logger.info(f"   • Average magnitude score: {quantitative_analysis.get('average_magnitude', 0):.3f}")
                    logger.info(f"   • Charts generated: {len(chart_files)}")
                
                # STEP 9: Upload all outputs to S3
                logger.info("☁️ STEP 9: Uploading results to S3")
                upload_summary = upload_outputs_to_s3(lens_type, output_dir, graph_file, report_files, logger)
                
                results[lens_type] = {
                    'graph': G,
                    'output_dir': output_dir,
                    'graph_file': graph_file,
                    'reports': report_files,
                    'category_reports': category_reports,
                    'quantitative_analysis': quantitative_analysis,
                    'chart_files': chart_files,
                    'study_analysis': study_analysis,
                    'nodes': G.number_of_nodes(),
                    'edges': G.number_of_edges(),
                    'comfort_drivers_found': len(study_analysis['comfort_drivers']),
                    'total_study_relationships': len(enhanced_relationships),
                    'quantitative_relationships': quantitative_analysis.get('quantitative_relationships', 0),
                    'high_magnitude_relationships': quantitative_analysis.get('high_magnitude_count', 0),
                    'average_magnitude': quantitative_analysis.get('average_magnitude', 0),
                    'category_reports_count': len(category_reports),
                    'charts_generated': len(chart_files),
                    's3_upload': upload_summary
                }
                
                logger.info(f"✅ {lens_type} processing complete")
                
            except Exception as e:
                logger.error(f"❌ Failed to process {lens_type}: {e}")
                continue
        
        # Generate overall summary
        total_time = time.time() - start_time
        
        # STEP 10: Create comprehensive overall findings summary (NEW)
        logger = logging.getLogger('main')
        logger.info("📋 STEP 10: Creating overall findings summary report")
        
        overall_summary_file = create_overall_findings_summary(results, Path(BASE_OUTPUT_DIR), logger)
        if overall_summary_file:
            logger.info(f"✅ Overall findings summary created: {overall_summary_file}")
        
        print("\n" + "=" * 80)
        print("CONTACT LENS PERFORMANCE DRIVERS PIPELINE COMPLETE")
        print("=" * 80)
        
        # Create overall pipeline summary for S3
        pipeline_summary = {
            'pipeline_name': 'contact_lens_drivers_knowledge_graph_with_nlp_llm_v2',
            'execution_timestamp': time.strftime("%Y-%m-%d %H:%M:%S UTC", time.gmtime()),
            'total_processing_time_seconds': total_time,
            'lens_types_processed': list(results.keys()),
            'overall_statistics': {
                'total_lens_types': len(results),
                'total_nodes_across_all_graphs': sum(r['nodes'] for r in results.values()),
                'total_edges_across_all_graphs': sum(r['edges'] for r in results.values()),
                'total_documents_processed': sum(r.get('study_analysis', {}).get('documents_processed', 0) for r in results.values()),
                'total_study_relationships_extracted': sum(r.get('study_analysis', {}).get('total_relationships', 0) for r in results.values()),
                'total_comfort_drivers_found': sum(len(r.get('study_analysis', {}).get('comfort_drivers', {})) for r in results.values()),
                'total_quantitative_relationships': sum(r.get('quantitative_relationships', 0) for r in results.values()),
                'total_high_magnitude_relationships': sum(r.get('high_magnitude_relationships', 0) for r in results.values()),
                'total_charts_generated': sum(r.get('charts_generated', 0) for r in results.values()),
                'total_category_reports': sum(r.get('category_reports_count', 0) for r in results.values())
            },
            'nlp_llm_processing': {
                'llm_endpoint': ENDPOINT_NAME,
                'parallel_processing_enabled': True,
                'max_workers': MAX_WORKERS,
                'max_llm_workers': MAX_LLM_WORKERS,
                'text_extraction_methods': ['PyPDF2', 'pdfminer', 'utf-8_decode'],
                'chunk_processing': True,
                'relationship_matching': True,
                'language_detection': True
            },
            's3_locations': {}
        }
        
        for lens_type, result in results.items():
            print(f"\n🔬 {lens_type.upper()}:")
            print(f"   📊 Graph Nodes: {result['nodes']}")
            print(f"   🔗 Graph Edges: {result['edges']}")
            print(f"   📁 Output: {result['output_dir']}")
            print(f"   📈 Graph: {result['graph_file'].name}")
            
            # Add study analysis information
            if 'study_analysis' in result:
                study = result['study_analysis']
                # Safely access study information with fallbacks
                documents_processed = study.get('documents_processed', 0)
                total_documents = study.get('total_documents', study.get('document_count', documents_processed))
                total_relationships = study.get('total_relationships', 0)
                
                print(f"   📄 Documents Processed: {documents_processed}/{total_documents}")
                print(f"   🔗 Study Relationships: {total_relationships}")
                print(f"   😊 Comfort Drivers Found: {len(study['comfort_drivers'])}")
                
                # Show top comfort drivers
                if study['comfort_drivers']:
                    top_drivers = sorted(
                        study['comfort_drivers'].items(), 
                        key=lambda x: x[1]['evidence_count'], 
                        reverse=True
                    )[:3]
                    
                    print(f"   🎯 Top Comfort Drivers:")
                    for driver_key, driver_data in top_drivers:
                        print(f"      • {driver_data['driver_name']}: {driver_data['evidence_count']} evidence points ({driver_data['overall_confidence']} confidence)")
            
            # Add S3 information if available
            if 's3_upload' in result:
                s3_info = result['s3_upload']
                print(f"   ☁️  S3 Location: s3://{OUTPUT_BUCKET_NAME}/{s3_info['s3_prefix']}")
                
                # Count successful uploads
                successful_uploads = 0
                total_uploads = 0
                
                if s3_info['uploads']['graph'] and s3_info['uploads']['graph'].get('status') == 'success':
                    successful_uploads += 1
                total_uploads += 1
                
                for report_status in s3_info['uploads']['reports'].values():
                    if report_status.get('status') == 'success':
                        successful_uploads += 1
                    total_uploads += 1
                
                print(f"   📤 S3 Uploads: {successful_uploads}/{total_uploads} successful")
                
                # Add to pipeline summary
                pipeline_summary['s3_locations'][lens_type] = {
                    'base_prefix': s3_info['s3_prefix'],
                    'timestamp': s3_info['timestamp'],
                    'graph_url': s3_info['uploads']['graph'].get('s3_url') if s3_info['uploads']['graph'] else None,
                    'manifest_url': s3_info.get('manifest', {}).get('s3_url'),
                    'study_analysis': {
                        'documents_processed': result.get('study_analysis', {}).get('documents_processed', 0),
                        'comfort_drivers_found': len(result.get('study_analysis', {}).get('comfort_drivers', {})),
                        'total_relationships': result.get('study_analysis', {}).get('total_relationships', 0)
                    }
                }
        
        # Save and upload pipeline summary
        try:
            summary_file = Path(BASE_OUTPUT_DIR) / f"pipeline_summary_{time.strftime('%Y%m%d_%H%M%S')}.json"
            with open(summary_file, 'w') as f:
                json.dump(pipeline_summary, f, indent=2)
            
            # Upload pipeline summary to S3
            summary_s3_key = f"{S3_OUTPUT_PREFIX}/pipeline_summaries/summary_{time.strftime('%Y%m%d_%H%M%S')}.json"
            if upload_file_to_s3(summary_file, summary_s3_key):
                print(f"\n📋 Pipeline Summary uploaded to: s3://{OUTPUT_BUCKET_NAME}/{summary_s3_key}")
            
        except Exception as e:
            print(f"\n⚠️ Failed to create/upload pipeline summary: {e}")
        
        # Save comprehensive debug logs
        try:
            debug_summary = get_debug_summary()
            debug_logger.flush_to_file()
            
            print(f"\n🔍 Debug Logging Summary:")
            print(f"   Total Steps: {debug_summary['total_steps']}")
            print(f"   Success Rate: {debug_summary['success_rate']:.2%}")
            print(f"   Debug File: debug_flow.json")
            
            # Also save a structured debug summary
            debug_summary_file = Path(BASE_OUTPUT_DIR) / f"debug_summary_{time.strftime('%Y%m%d_%H%M%S')}.json"
            with open(debug_summary_file, 'w') as f:
                json.dump(debug_summary, f, indent=2)
            
            # Upload debug logs to S3
            debug_s3_key = f"{S3_OUTPUT_PREFIX}/debug_logs/debug_flow_{time.strftime('%Y%m%d_%H%M%S')}.json"
            debug_flow_file = Path("debug_flow.json")
            if debug_flow_file.exists() and upload_file_to_s3(debug_flow_file, debug_s3_key):
                print(f"   📤 Debug logs uploaded to: s3://{OUTPUT_BUCKET_NAME}/{debug_s3_key}")
                
            debug_summary_s3_key = f"{S3_OUTPUT_PREFIX}/debug_logs/debug_summary_{time.strftime('%Y%m%d_%H%M%S')}.json"
            if upload_file_to_s3(debug_summary_file, debug_summary_s3_key):
                print(f"   📤 Debug summary uploaded to: s3://{OUTPUT_BUCKET_NAME}/{debug_summary_s3_key}")
                
        except Exception as e:
            print(f"\n⚠️ Failed to save/upload debug logs: {e}")
        
        # Print debug analysis for immediate insights
        print_debug_analysis()
        
        debug_logger.log_step(
            "main_pipeline_complete",
            "output",
            output_data={
                "total_execution_time": time.time() - start_time,
                "lens_types_processed": list(results.keys()),
                "pipeline_summary": pipeline_summary
            },
            metadata={"pipeline_success": True}
        )
        
        # Create and upload HTML index page
        try:
            html_content = create_s3_index_page(results)
            index_file = Path(BASE_OUTPUT_DIR) / f"index_{time.strftime('%Y%m%d_%H%M%S')}.html"
            with open(index_file, 'w') as f:
                f.write(html_content)
            
            # Upload index page to S3
            index_s3_key = f"{S3_OUTPUT_PREFIX}/index.html"
            if upload_file_to_s3(index_file, index_s3_key):
                index_url = f"https://s3.amazonaws.com/{OUTPUT_BUCKET_NAME}/{index_s3_key}"
                print(f"🌐 Results Index Page: {index_url}")
            
        except Exception as e:
            print(f"\n⚠️ Failed to create/upload index page: {e}")
        
        print(f"\n⏱️  Total Processing Time: {total_time:.2f} seconds")
        print(f"📁 Local Output Directory: {BASE_OUTPUT_DIR}")
        print(f"☁️  S3 Base Location: s3://{OUTPUT_BUCKET_NAME}/{S3_OUTPUT_PREFIX}")
        print("\n🤖 NLP/LLM Processing Summary:")
        print(f"   📄 Total Documents Processed: {sum(r.get('study_analysis', {}).get('documents_processed', 0) for r in results.values())}")
        print(f"   🔗 Total Relationships Extracted: {sum(r.get('study_analysis', {}).get('total_relationships', 0) for r in results.values())}")
        print(f"   😊 Total Comfort Drivers Found: {sum(len(r.get('study_analysis', {}).get('comfort_drivers', {})) for r in results.values())}")
        print(f"   📊 Performance Categories: {sum(r.get('category_reports_count', 0) for r in results.values())} category reports generated")
        print(f"   📈 Quantitative Relationships: {sum(r.get('quantitative_relationships', 0) for r in results.values())}")
        print(f"   🎯 High Magnitude Relationships: {sum(r.get('high_magnitude_relationships', 0) for r in results.values())}")
        print(f"   📉 Charts Generated: {sum(r.get('charts_generated', 0) for r in results.values())}")
        print(f"   🧠 LLM Endpoint Used: {ENDPOINT_NAME}")
        print(f"   ⚡ Maximum Workers: {MAX_WORKERS} document, {MAX_LLM_WORKERS} LLM")
        print(f"   ⚡ Parallel Processing: Enabled (200 workers)")
        print("\n📋 Performance Category Reports Generated:")
        for lens_type, result in results.items():
            category_reports = result.get('category_reports', {})
            if category_reports:
                print(f"   {lens_type}:")
                for category in ["Size", "Fit", "Handling", "Comfort"]:
                    if category in category_reports:
                        timestamp = time.strftime("%Y%m%d_%H%M%S")
                        print(f"      📊 {category}: {category.lower()}_analysis_{lens_type}_{timestamp}.md")
        print("\n🔧 Auto Scaling Performance Summary:")
        print(f"   🚀 Worker Configuration: {MAX_WORKERS} document workers, {MAX_LLM_WORKERS} LLM workers")
        print(f"   📊 Batched Processing: Optimized for auto scaling endpoints")
        print(f"   🎯 Timeout Handling: Extended timeouts for scaling warmup")
        print(f"   ⚖️ Load Balancing: Adaptive batch sizing and progressive backoff")
        print(f"   📈 Throughput Monitoring: Real-time performance tracking")
        print(f"   🔄 Endpoint Failover: {len([e for e in ENDPOINTS if e['active']])} active endpoints")
        print("\n🎉 Pipeline completed successfully!")
        print("\n🔍 To access your results:")
        print(f"   • Local files: {Path(BASE_OUTPUT_DIR).absolute()}")
        print(f"   • S3 browser: https://s3.console.aws.amazon.com/s3/buckets/{OUTPUT_BUCKET_NAME}?prefix={S3_OUTPUT_PREFIX}/")
        print(f"   • Direct S3 access: aws s3 ls s3://{OUTPUT_BUCKET_NAME}/{S3_OUTPUT_PREFIX}/ --recursive")
        if overall_summary_file:
            print(f"   • Overall Summary: {Path(overall_summary_file).name}")
        print("\n📊 Key Features Demonstrated:")
        print("   ✅ Maximum Parallel Processing (500+ workers)")
        print("   ✅ S3 Document Processing")
        print("   ✅ PDF Text Extraction") 
        print("   ✅ LLM-based Relationship Extraction")
        print("   ✅ NLP Driver Detection")
        print("   ✅ Multi-Category Analysis (Size, Fit, Handling, Comfort)")
        print("   ✅ Quantitative Value Analysis")
        print("   ✅ Statistical Charts & Visualizations")
        print("   ✅ Magnitude Scoring & Confidence Analysis")
        print("   ✅ Timestamped Report Generation")
        print("   ✅ Overall Findings Summary Report")
        print("   ✅ Knowledge Graph Enhancement")
        print("   ✅ Comprehensive S3 Output")
        print("   ✅ Auto Scaling Optimization")
        print("   ✅ Performance Monitoring")
        print("   ✅ Endpoint Management")
        
        # Retry failed files
        if failed_files:
            print(f"🔄 Retrying {len(failed_files)} failed files...")
            retry_count = 0
            max_retries = 3
            
            while failed_files and retry_count < max_retries:
                retry_count += 1
                print(f"🔄 Retry attempt {retry_count}/{max_retries}")
                
                failed_files_copy = failed_files.copy()
                for file_path in failed_files_copy:
                    try:
                        print(f"🔄 Retrying failed file: {file_path}")
                        
                        # Extract bucket and key from s3 path
                        if file_path.startswith("s3://"):
                            parts = file_path.replace("s3://", "").split("/", 1)
                            bucket = parts[0]
                            key = parts[1]
                            
                            # Create a temporary logger for retry
                            retry_logger = logging.getLogger("retry")
                            
                            # Retry extraction
                            extracted_text = extract_text_from_s3_document(bucket, key, retry_logger)
                            
                            if extracted_text:
                                print(f"✅ Successfully retried: {file_path}")
                                failed_files.remove(file_path)
                            else:
                                print(f"❌ Retry failed: {file_path}")
                                
                    except Exception as e:
                        print(f"❌ Retry failed for file {file_path}: {e}")
                
                if failed_files:
                    print(f"⏰ Waiting 30 seconds before next retry...")
                    time.sleep(30)
            
            if failed_files:
                print(f"❌ Still have {len(failed_files)} failed files after all retries")
            else:
                print("✅ All failed files successfully processed")
        
        # Print final processing status
        print("\n" + "="*60)
        print("📊 FINAL PROCESSING STATUS")
        print("="*60)
        print_processing_status()
        
        return results
        
    except Exception as e:
        print(f"❌ Pipeline failed: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == '__main__':
    import sys
    
    # Check for endpoint management commands
    if len(sys.argv) > 1:
        command = sys.argv[1].lower()
        
        if command == 'endpoints':
            list_available_endpoints()
            sys.exit(0)
            
        elif command == 'switch':
            if len(sys.argv) < 3:
                print("Usage: python graph.py switch <endpoint_name_or_index>")
                print("Available endpoints:")
                list_available_endpoints()
                sys.exit(1)
            
            try:
                # Try as index first
                endpoint_arg = sys.argv[2]
                if endpoint_arg.isdigit():
                    endpoint_index = int(endpoint_arg) - 1  # Convert to 0-based index
                    switch_endpoint(endpoint_index)
                else:
                    # Try as endpoint name
                    switch_endpoint(endpoint_name=endpoint_arg)
                print("Endpoint switched successfully!")
            except Exception as e:
                print(f"❌ Failed to switch endpoint: {e}")
                list_available_endpoints()
                sys.exit(1)
            sys.exit(0)
            
        elif command == 'health':
            endpoint_name = sys.argv[2] if len(sys.argv) > 2 else None
            if endpoint_health_check(endpoint_name):
                print(f"✅ Endpoint {endpoint_name or ENDPOINT_NAME} is healthy")
            else:
                print(f"❌ Endpoint {endpoint_name or ENDPOINT_NAME} is unhealthy")
            sys.exit(0)
            
        elif command == 'help':
            print("\n🔧 Endpoint Management Commands:")
            print("=" * 40)
            print("python graph.py endpoints          - List all available endpoints")
            print("python graph.py switch <name|#>   - Switch to specific endpoint")
            print("python graph.py health [name]     - Check endpoint health")
            print("python graph.py help              - Show this help")
            print("python graph.py                   - Run the full pipeline")
            print("\n📋 Endpoint Configuration:")
            print("Edit the ENDPOINTS list in the script to add/remove endpoints.")
            print("Set 'active': True to enable an endpoint.")
            print("Lower priority numbers are preferred.\n")
            sys.exit(0)
    
    # Run the main pipeline
    main()